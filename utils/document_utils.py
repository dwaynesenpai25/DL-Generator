import pandas as pd
import gspread
from google.oauth2.service_account import Credentials
from docx import Document
from docx.shared import Inches, Pt
import re
from pathlib import Path
from tempfile import TemporaryDirectory
from utils.config import logger
import qrcode
from barcode import Code128
from barcode.writer import ImageWriter
from io import BytesIO

def get_sheet_data(service_account_json_path, spreadsheet_id, sheet_name="LetterHeads"):
    try:
        # Updated scopes to match exactly what gspread expects
        scopes = [
            'https://spreadsheets.google.com/feeds',
            'https://www.googleapis.com/auth/spreadsheets',
            'https://www.googleapis.com/auth/spreadsheets',
            'https://www.googleapis.com/auth/drive'
        ]
        
        credentials = Credentials.from_service_account_file(
            service_account_json_path,
            scopes=scopes
        )
        
        client = gspread.authorize(credentials)
        spreadsheet = client.open_by_key(spreadsheet_id)
        worksheet = spreadsheet.worksheet(sheet_name)
        sheet_data = worksheet.get_all_records()
        df = pd.DataFrame(sheet_data)
        required_columns = ["CAMPAIGN", "DL TYPE", "FILE"]
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            logger.error(f"Missing columns in spreadsheet: {missing_columns}")
            return pd.DataFrame()
        df = df[["CAMPAIGN", "DL TYPE", "FILE"]].dropna()
        logger.info(f"Retrieved sheet data with {len(df)} valid rows")
        return df
    except Exception as e:
        logger.error(f"Failed to access Google Sheets: {e}")
        return pd.DataFrame()

def combine_templates(header_footer_path, content_path, signature_img_path, word_app):
    try:
        header_footer_doc = Document(header_footer_path)
        content_doc = Document(content_path)
        while header_footer_doc.paragraphs:
            header_footer_doc.paragraphs[0]._element.getparent().remove(header_footer_doc.paragraphs[0]._element)

        for elem in content_doc.element.body:
            header_footer_doc.element.body.append(elem)

        for section in header_footer_doc.sections:
            section.header_distance = Inches(0.1)
            section.footer_distance = Inches(0.1)

        if signature_img_path and Path(signature_img_path).exists():
            with TemporaryDirectory() as temp_dir:
                temp_doc_path = Path(temp_dir) / "temp_combined_doc.docx"
                header_footer_doc.save(temp_doc_path)
                temp_doc_path = replace_in_text_boxes(header_footer_doc, "«IMAGE_SIGNATURE»", signature_img_path, word_app, temp_doc_path)
                header_footer_doc = Document(temp_doc_path)
        
        return header_footer_doc
    except Exception as e:
        logger.error(f"Failed to combine templates or replace signature: {e}")
        return None

def replace_in_text_boxes(doc, find_str, replace_with_image_path, word_app, temp_doc_path):
    try:
        word_doc = word_app.Documents.Open(str(temp_doc_path.absolute()))
        modified = False
        for shape in word_doc.Shapes:
            if shape.TextFrame.HasText:
                text_range = shape.TextFrame.TextRange
                find = text_range.Find
                find.Text = find_str
                find.Forward = True
                find.Wrap = 1
                while find.Execute():
                    found_range = text_range.Duplicate
                    found_range.Find.Execute(FindText=find_str)
                    found_range.Text = ""
                    inline_shape = found_range.InlineShapes.AddPicture(
                        FileName=str(Path(replace_with_image_path).absolute()),
                        LinkToFile=False,
                        SaveWithDocument=True
                    )
                    inline_shape.Width = 110
                    inline_shape.Height = 50
                    modified = True
        word_doc.Save()
        word_doc.Close(SaveChanges=True)
        if modified:
            logger.info(f"Replaced '{find_str}' with image in shapes.")
        return temp_doc_path
    except Exception as e:
        logger.error(f"Error replacing text in shapes: {e}")
        try:
            word_doc.Close(SaveChanges=False)
        except:
            pass
        return temp_doc_path

def extract_placeholders(doc):
    placeholders = set()
    for para in doc.paragraphs:
        if "«" in para.text:
            matches = re.findall(r"«(.*?)»", para.text)
            placeholders.update(["«" + m.strip() + "»" for m in matches])
    for node in doc._element.iter():
        if node.tag.endswith("}t") and node.text and "«" in node.text:
            matches = re.findall(r"«(.*?)»", node.text)
            placeholders.update(["«" + m.strip() + "»" for m in matches])
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header is not None:
                for para in header.paragraphs:
                    if "«" in para.text:
                        matches = re.findall(r"«(.*?)»", para.text)
                        placeholders.update(["«" + m.strip() + "»" for m in matches])
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if "«" in para.text:
                                    matches = re.findall(r"«(.*?)»", para.text)
                                    placeholders.update(["«" + m.strip() + "»" for m in matches])
                for node in header._element.iter():
                    if node.tag.endswith("}t") and node.text and "«" in node.text:
                        matches = re.findall(r"«(.*?)»", node.text)
                        placeholders.update(["«" + m.strip() + "»" for m in matches])
        for footer in [section.footer, section.first_page_footer, section.even_page_header]:
            if footer is not None:
                for para in footer.paragraphs:
                    if "«" in para.text:
                        matches = re.findall(r"«(.*?)»", para.text)
                        placeholders.update(["«" + m.strip() + "»" for m in matches])
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if "«" in para.text:
                                    matches = re.findall(r"«(.*?)»", para.text)
                                    placeholders.update(["«" + m.strip() + "»" for m in matches])
                for node in footer._element.iter():
                    if node.tag.endswith("}t") and node.text and "«" in node.text:
                        matches = re.findall(r"«(.*?)»", node.text)
                        placeholders.update(["«" + m.strip() + "»" for m in matches])
    return sorted(placeholders)

def extract_transmittal_placeholders(template_path):
    """
    Extract placeholders from transmittal templates that contain nested tables.
    Based on the structure: 1 main table with nested tables in each cell.
    
    Args:
        template_path (str): Path to the transmittal template document
        
    Returns:
        list: Sorted list of unique placeholders found in the template
    """
    try:
        # Open the document
        doc = Document(template_path)
        placeholders = set()
        
        def extract_from_text(text):
            """Helper function to extract placeholders from text"""
            if text and "«" in text:
                matches = re.findall(r"«(.*?)»", text)
                return ["«" + m.strip() + "»" for m in matches]
            return []
        
        def extract_from_paragraphs(paragraphs):
            """Helper function to extract placeholders from paragraphs"""
            for para in paragraphs:
                placeholders.update(extract_from_text(para.text))
        
        def extract_from_inner_table(inner_table):
            """Extract placeholders from inner/nested tables"""
            for row in inner_table.rows:
                for cell in row.cells:
                    # Extract from cell paragraphs
                    extract_from_paragraphs(cell.paragraphs)
                    
                    # Check for further nested tables
                    for nested_table in cell.tables:
                        extract_from_inner_table(nested_table)
        
        # Extract from main document paragraphs
        extract_from_paragraphs(doc.paragraphs)
        
        # Extract from the main table structure (based on your fill_transmittal_template logic)
        if doc.tables:
            main_table = doc.tables[0]  # The main table containing nested tables
            
            for row in main_table.rows:
                for cell in row.cells:
                    # Extract from cell paragraphs
                    extract_from_paragraphs(cell.paragraphs)
                    
                    # Extract from nested tables in each cell (this is the key part)
                    if cell.tables:
                        for inner_table in cell.tables:
                            extract_from_inner_table(inner_table)
        
        # Extract from any additional tables at document level
        for table_idx, table in enumerate(doc.tables):
            if table_idx > 0:  # Skip the main table we already processed
                for row in table.rows:
                    for cell in row.cells:
                        extract_from_paragraphs(cell.paragraphs)
                        for nested_table in cell.tables:
                            extract_from_inner_table(nested_table)
        
        # Extract from XML nodes (for text that might not be in paragraphs)
        for node in doc._element.iter():
            if node.tag.endswith("}t") and node.text:
                placeholders.update(extract_from_text(node.text))
        
        # Extract from headers and footers
        for section in doc.sections:
            # Process headers
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header is not None:
                    # Header paragraphs
                    extract_from_paragraphs(header.paragraphs)
                    
                    # Header tables (including nested ones)
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                extract_from_paragraphs(cell.paragraphs)
                                for nested_table in cell.tables:
                                    extract_from_inner_table(nested_table)
                    
                    # Header XML nodes
                    for node in header._element.iter():
                        if node.tag.endswith("}t") and node.text:
                            placeholders.update(extract_from_text(node.text))
            
            # Process footers
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer is not None:
                    # Footer paragraphs
                    extract_from_paragraphs(footer.paragraphs)
                    
                    # Footer tables (including nested ones)
                    for table in footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                extract_from_paragraphs(cell.paragraphs)
                                for nested_table in cell.tables:
                                    extract_from_inner_table(nested_table)
                    
                    # Footer XML nodes
                    for node in footer._element.iter():
                        if node.tag.endswith("}t") and node.text:
                            placeholders.update(extract_from_text(node.text))
        
        return sorted(placeholders)
        
    except Exception as e:
        print(f"Error extracting transmittal placeholders: {e}")
        return []


def generate_barcode(barcode_value):
    try:
        barcode = Code128(barcode_value, writer=ImageWriter())
        buffer = BytesIO()
        barcode.write(buffer, options={"write_text": False, "module_width": 1, "module_height": 8, "quiet_zone": 2.0})
        buffer.seek(0)
        return buffer
    except Exception as e:
        logger.error(f"Failed to generate barcode for {barcode_value}: {e}")
        return None

def generate_qrcode(data):
    try:
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        qr.add_data(data)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        buffer = BytesIO()
        img.save(buffer, format="PNG")
        buffer.seek(0)
        return buffer
    except Exception as e:
        logger.error(f"Failed to generate QR code for {data}: {e}")
        return None

def number_to_words(number):
    if number == 0:
        return "ZERO"
    units = ["", "ONE", "TWO", "THREE", "FOUR", "FIVE", "SIX", "SEVEN", "EIGHT", "NINE"]
    teens = ["", "ELEVEN", "TWELVE", "THIRTEEN", "FOURTEEN", "FIFTEEN", "SIXTEEN", "SEVENTEEN", "EIGHTEEN", "NINETEEN"]
    tens = ["", "TEN", "TWENTY", "THIRTY", "FORTY", "FIFTY", "SIXTY", "SEVENTY", "EIGHTY", "NINETY"]
    thousands = ["", "THOUSAND", "MILLION"]
    def convert_hundreds(num):
        if num == 0:
            return ""
        result = []
        if num >= 100:
            result.append(units[num // 100] + " HUNDRED")
            num %= 100
            if num > 0:
                result.append("AND")
        if num >= 20:
            result.append(tens[num // 10])
            num %= 10
            if num > 0:
                result.append(units[num])
        elif num >= 11:
            result.append(teens[num - 10])
        elif num == 10:
            result.append(tens[1])
        elif num > 0:
            result.append(units[num])
        return " ".join(result)
    result = []
    i = 0
    while number > 0:
        chunk = number % 1000
        if chunk > 0:
            chunk_words = convert_hundreds(chunk)
            if thousands[i]:
                chunk_words += f" {thousands[i]}"
            result.insert(0, chunk_words)
        number //= 1000
        i += 1
    return ", ".join(result).replace(", AND", " AND").strip()

def amount_to_words(amount_str):
    try:
        amount = float(amount_str.replace(",", ""))
        pesos = int(amount)
        cents = int(round((amount - pesos) * 100))
        pesos_words = number_to_words(pesos)
        if pesos == 0:
            pesos_words = "ZERO"
        cents_words = number_to_words(cents) if cents > 0 else "ZERO"
        result = f"{pesos_words} PESOS"
        if cents > 0:
            result += f", AND {cents_words} CENTS"
        else:
            result += ", AND ZERO CENTS"
        return result
    except Exception as e:
        logger.error(f"Failed to convert amount to words: {amount_str}, error: {e}")
        return "ERROR CONVERTING AMOUNT"
