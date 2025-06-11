import pandas as pd
import asyncio
import json
import time
import uuid
import shutil
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory
from fastapi import HTTPException
from fastapi.responses import FileResponse
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from PyPDF2 import PdfMerger
from utils.config import logger
from utils.session_management import get_user_session_state
from utils.document_utils import generate_barcode, generate_qrcode, amount_to_words
from utils.pdf_conversion import batch_convert_libreoffice
from utils.database import add_audit_entry, add_processed_accounts
import os
import re
def get_raw_file(file):
    try:
        contents = file.file.read()
        df = pd.read_excel(BytesIO(contents), dtype=str)
        return df
    except Exception as e:
        logger.error(f"Error reading file: {e}")
        return pd.DataFrame([])

def process_excel_file(file):
    if not file.filename.endswith('.xlsx'):
        raise HTTPException(status_code=400, detail="Invalid file format. Please upload an .xlsx file")
    df = get_raw_file(file)
    if df.empty:
        raise HTTPException(status_code=500, detail="Failed to read Excel file")
    if 'FINAL_AREA' not in df.columns or 'DL_CODE' not in df.columns:
        raise HTTPException(status_code=400, detail="Excel file must contain FINAL_AREA and DL_CODE columns")
    return {"data": df.to_dict(orient='records')}

def fill_template(doc, mapping, barcode_buffer=None):
    def replace_in_text(text):
        for k, v in mapping.items():
            if k == "«IMAGE_BARCODE»" and v:
                continue
            text = text.replace(k, str(v))
        return text
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = replace_in_text(run.text)
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = replace_in_text(run.text)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if "«IMAGE_BARCODE»" in para.text and barcode_buffer:
                                    para.clear()
                                    cell.paragraphs[0].paragraph_format.left_indent = Pt(-20)
                                    run = para.add_run()
                                    run.add_picture(barcode_buffer, width=Inches(3.0), height=Inches(0.35))
                                else:
                                    for run in para.runs:
                                        if run.text:
                                            run.text = replace_in_text(run.text)
        for footer in [section.footer, section.first_page_footer, section.even_page_header]:
            if footer:
                for para in footer.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = replace_in_text(run.text)
    return doc

def clear_placeholders(inner_table):
    def replace_in_text(text):
        return re.sub(r'«[^»]+»', "", text)
    for row in inner_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text:
                        run.text = replace_in_text(run.text)

def fill_inner_table(inner_table, mapping, qrcode_buffer=None):
    def replace_in_text(text):
        for k, v in mapping.items():
            if k == "«IMAGE_QRCODE»":
                continue
            text = text.replace(k, str(v))
        return text
    for row in inner_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if "«IMAGE_QRCODE»" in para.text and qrcode_buffer:
                    para.clear()
                    para.paragraph_format.left_indent = Pt(-5)
                    run = para.add_run()
                    run.add_picture(qrcode_buffer, width=Inches(1), height=Inches(1))
                else:
                    for run in para.runs:
                        if run.text:
                            run.text = replace_in_text(run.text)

def fill_transmittal_template(template_path, group_df):
    try:
        temp_doc = Document(template_path)
        if not temp_doc.tables:
            logger.error("No tables found in the transmittal template.")
            return None
        total_groups = len(group_df) // 4 + (1 if len(group_df) % 4 else 0)
        temp_docs = []
        for group_idx in range(total_groups):
            temp_group_doc = Document(template_path)
            current_group_table = temp_group_doc.tables[0]
            page_rows = group_df.iloc[group_idx * 4:(group_idx + 1) * 4]
            for row_idx, row in enumerate(current_group_table.rows):
                if row_idx < len(page_rows):
                    row_data = page_rows.iloc[row_idx]
                    mapping = {f"«{col.upper()}»": str(row_data[col]) for col in group_df.columns if pd.notnull(row_data[col])}
                    qrcode_buffer = None
                    if dl_code := row_data.get('DL_CODE', ''):
                        qrcode_buffer = generate_qrcode(dl_code)
                    for cell in row.cells:
                        if cell.tables:
                            inner_table = cell.tables[0]
                            fill_inner_table(inner_table, mapping, qrcode_buffer)
                else:
                    for cell in row.cells:
                        if cell.tables:
                            inner_table = cell.tables[0]
                            clear_placeholders(inner_table)
            temp_docs.append(temp_group_doc)
        return temp_docs
    except Exception as e:
        logger.error(f"Failed to fill transmittal template: {e}")
        return None

def format_time_duration(seconds):
    """Convert seconds to human-readable format"""
    if seconds < 60:
        return f"{seconds:.1f} seconds"
    elif seconds < 3600:
        minutes = seconds / 60
        return f"{minutes:.1f} minutes"
    else:
        hours = seconds / 3600
        return f"{hours:.1f} hours"

async def generate_pdfs_stream(uploaded_file, dataframe: pd.DataFrame, user_info: dict):
    user_email = user_info.get("email", "")
    session_state = get_user_session_state(user_email)
    
    # Use processing lock to prevent concurrent processing for the same user
    if not session_state['processing_lock'].acquire(blocking=False):
        yield json.dumps({'error': 'Another processing task is already running for this user'}) + '\n'
        return
    
    zipf = None
    try:
        if not session_state.get('selected_mode'):
            yield json.dumps({'error': 'No processing mode selected'}) + '\n'
            return

        # Check for Excel validation errors
        if session_state.get('excel_errors'):
            error_msg = f"Excel validation errors found: {'; '.join(session_state['excel_errors'])}"
            yield json.dumps({'error': error_msg}) + '\n'
            return

        # Check for signature errors
        if session_state.get('signature_error'):
            yield json.dumps({'error': session_state['signature_error']}) + '\n'
            return

        if not session_state.get('base_template') and session_state.get('selected_mode') in ["DL Only", "DL w/ Transmittal"]:
            yield json.dumps({'error': 'DL templates not loaded'}) + '\n'
            return
        if not session_state.get('transmittal_template_path') and session_state.get('selected_mode') in ["DL w/ Transmittal", "Transmittal Only"]:
            yield json.dumps({'error': 'Transmittal template not loaded'}) + '\n'
            return

        logger.debug(f"Excel file loaded with {len(dataframe)} rows for user {user_email}")
        from datetime import datetime
        today_date = datetime.now().strftime("%B %d, %Y")
        dataframe['DL_ADDRESS'] = dataframe['DL_ADDRESS'].str.upper()
        dataframe['LEADS_NEW_OB'] = dataframe['LEADS_NEW_OB'].apply(lambda x: f"{float(x):,.2f}")
        valid_rows = dataframe[dataframe['LEADS_CHNAME'].notna()]
        total_records = len(valid_rows)
        if total_records == 0:
            yield json.dumps({'error': 'No valid rows found (LEADS_CHNAME missing)'}) + '\n'
            return

        # Log available FINAL_AREA values for debugging
        available_areas = valid_rows['FINAL_AREA'].unique().tolist()
        logger.info(f"Available FINAL_AREA values for user {user_email}: {available_areas}")
        if 'NCR' not in available_areas:
            logger.warning(f"No rows found with FINAL_AREA 'NCR' in the input data for user {user_email}")

        logger.info(f"Processing {total_records} records across {len(valid_rows.groupby('FINAL_AREA'))} FINAL_AREA groups for user {user_email}")
        
        # Get user info for audit trail
        user_name = user_info.get("name", user_email)
        selected_folder = session_state.get('selected_folder', 'Unknown')
        selected_mode = session_state.get('selected_mode', 'Unknown')
        
        # Create audit trail entry and get its ID for linking processed accounts
        audit_id = add_audit_entry(selected_folder, user_name, total_records, selected_mode, selected_folder, session_state.get('selected_dl_type', 'Unknown'))

        # Get selected output format
        output_format = session_state.get('output_format', 'zip')  # Default to zip if not set
        
        # Use user-specific output directory
        user_output_dir = session_state['user_output_dir']

        # Log placeholders for debugging
        placeholders = session_state.get('placeholders', [])
        logger.info(f"Placeholders in template for user {user_email}: {placeholders}")

        with TemporaryDirectory() as temp_zip_dir:
            zip_path = Path(temp_zip_dir) / f"final_area_pdfs_{user_email.replace('@', '_').replace('.', '_')}.zip"
            if output_format == "zip":
                zipf = zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED)
            
            # Enhanced progress callback for detailed reporting
            def progress_callback(progress_data):
                if progress_data['type'] == 'conversion_start':
                    yield json.dumps({
                        'progress': 85,
                        'message': f"Starting PDF conversion: {progress_data['total_files']} files in {progress_data['total_batches']} batches",
                        'conversion_details': {
                            'phase': 'conversion_start',
                            'total_files': progress_data['total_files'],
                            'total_batches': progress_data['total_batches'],
                            'batch_size': progress_data['batch_size']
                        }
                    }) + '\n'
                elif progress_data['type'] == 'batch_start':
                    yield json.dumps({
                        'progress': 85 + (progress_data['batch_id'] / progress_data.get('total_batches', 1)) * 10,
                        'message': f"Converting batch {progress_data['batch_id']}: {progress_data['batch_size']} files (Attempt {progress_data['attempt']})",
                        'conversion_details': {
                            'phase': 'batch_processing',
                            'batch_id': progress_data['batch_id'],
                            'batch_size': progress_data['batch_size'],
                            'attempt': progress_data['attempt']
                        }
                    }) + '\n'
                elif progress_data['type'] == 'batch_complete':
                    yield json.dumps({
                        'progress': 85 + (progress_data['batch_id'] / progress_data.get('total_batches', 1)) * 10,
                        'message': f"Batch {progress_data['batch_id']} complete: {progress_data['successful']}/{progress_data['successful'] + progress_data['failed']} files ({progress_data['success_rate']:.1f}% success)",
                        'conversion_details': {
                            'phase': 'batch_complete',
                            'batch_id': progress_data['batch_id'],
                            'successful': progress_data['successful'],
                            'failed': progress_data['failed'],
                            'success_rate': progress_data['success_rate']
                        }
                    }) + '\n'
                elif progress_data['type'] == 'conversion_summary':
                    yield json.dumps({
                        'progress': 95,
                        'message': f"PDF conversion complete: {progress_data['successful']}/{progress_data['total_files']} files in {progress_data['total_time_formatted']} ({progress_data['conversion_rate']:.1f} files/sec)",
                        'conversion_summary': {
                            'total_files': progress_data['total_files'],
                            'successful': progress_data['successful'],
                            'failed': progress_data['failed'],
                            'success_rate': progress_data['success_rate'],
                            'total_time': progress_data['total_time'],
                            'total_time_formatted': progress_data['total_time_formatted'],
                            'conversion_rate': progress_data['conversion_rate'],
                            'total_batches': progress_data['total_batches']
                        }
                    }) + '\n'
            
            try:
                with TemporaryDirectory() as temp_dir:
                    temp_base_path = Path(temp_dir) / "base_template.docx"
                    if session_state.get('selected_mode') in ["DL Only", "DL w/ Transmittal"]:
                        session_state['base_template'].save(temp_base_path)
                        session_state['files_to_cleanup'].append(temp_base_path)
                    temp_transmittal_path = Path(temp_dir) / "transmittal_template.docx"
                    if session_state.get('selected_mode') in ["DL w/ Transmittal", "Transmittal Only"]:
                        Document(session_state['transmittal_template_path']).save(temp_transmittal_path)
                        session_state['files_to_cleanup'].append(temp_transmittal_path)

                    processed_records = 0
                    start_time = time.time()
                    
                    # Store processed accounts in batches
                    account_batch = []
                    batch_size = 100
                    # Store DOCX files for print format
                    area_docx_files = {}  # Dictionary to track DOCX files by FINAL_AREA
                    
                    # Enhanced progress tracking for Transmittal Only mode
                    total_areas = len(valid_rows.groupby('FINAL_AREA'))
                    current_area_index = 0
                    
                    for final_area, group_df in valid_rows.groupby('FINAL_AREA'):
                        current_area_index += 1
                        logger.debug(f"Processing FINAL_AREA: {final_area} ({len(group_df)} records) for user {user_email}")
                        area_dir = user_output_dir / f"{final_area}"
                        os.makedirs(area_dir, exist_ok=True)
                        docx_files = []
                        temp_files = []
                        dl_pdf_merger = PdfMerger() if output_format == "zip" and session_state.get('selected_mode') in ["DL Only", "DL w/ Transmittal"] else None
                        transmittal_pdf_merger = PdfMerger() if output_format == "zip" and session_state.get('selected_mode') in ["DL w/ Transmittal", "Transmittal Only"] else None
                        record_count = len(group_df)
                        
                        # Enhanced progress reporting for Transmittal Only mode
                        if session_state.get('selected_mode') == "Transmittal Only":
                            base_progress = (current_area_index - 1) / total_areas * 80
                            area_progress = current_area_index / total_areas * 80
                            yield json.dumps({
                                'progress': base_progress,
                                'message': f"Processing Transmittal for {final_area}: {record_count} records ({current_area_index}/{total_areas} areas)"
                            }) + '\n'
                        
                        try:
                            if session_state.get('selected_mode') in ["DL Only", "DL w/ Transmittal"]:
                                for idx, row in group_df.iterrows():
                                    # Store account data
                                    account_batch.append((
                                        audit_id,
                                        row.get('DL_CODE', ''),
                                        row.get('LEADS_CHNAME', ''),
                                        row.get('DL_ADDRESS', ''),
                                        row.get('FINAL_AREA', '')
                                    ))
                                    
                                    # Process batch if it reaches the batch size
                                    if len(account_batch) >= batch_size:
                                        add_processed_accounts(audit_id, account_batch)
                                        account_batch = []
                                    
                                    # Create mapping with normalized placeholders
                                    barcode_buffer = None
                                    if barcode_value := row.get('DL_CODE', ''):
                                        barcode_buffer = generate_barcode(barcode_value)
                                    amount_words = amount_to_words(row.get('amount', '0.00'))
                                    mapping = {
                                        f"«{col.upper()}»": str(row[col]) for col in dataframe.columns if pd.notnull(row[col])
                                    }
                                    mapping.update({
                                        "«IMAGE_BARCODE»": barcode_buffer or "",
                                        "«DL_DATE»": today_date,
                                        "«AMOUNT_ABBR»": amount_words,
                                        "«IMAGE_SIGNATURE»": session_state.get('files_to_cleanup', [])[0] or ""
                                    })
                                    
                                    # Log mapping for debugging
                                    logger.debug(f"Mapping for row {idx} in {final_area} for user {user_email}: {mapping}")

                                    filled_doc = Document(temp_base_path)
                                    filled_doc = fill_template(filled_doc, mapping, barcode_buffer)
                                    if filled_doc:
                                        unique_name = f"dl_{final_area}_{idx}_{uuid.uuid4().hex[:6]}"
                                        docx_output = area_dir / f"{unique_name}.docx"
                                        filled_doc.save(docx_output)
                                        docx_files.append(str(docx_output))
                                        if output_format == "print":
                                            if final_area not in area_docx_files:
                                                area_docx_files[final_area] = []
                                            area_docx_files[final_area].append(str(docx_output))
                                        else:
                                            temp_files.append(str(docx_output))
                                    else:
                                        logger.warning(f"Failed to fill template for row {idx} in {final_area} for user {user_email}")
                                    processed_records += 1
                                    progress = (processed_records / total_records) * 80  # Reserve 20% for conversion
                                    yield json.dumps({
                                        'progress': progress,
                                        'message': f"Processing record {processed_records}/{total_records} (FINAL_AREA: {final_area})"
                                    }) + '\n'
                                    await asyncio.sleep(0)

                            if session_state.get('selected_mode') in ["DL w/ Transmittal", "Transmittal Only"]:
                                if session_state.get('selected_mode') == "Transmittal Only":
                                    for _, row in group_df.iterrows():
                                        account_batch.append((
                                            audit_id,
                                            row.get('DL_CODE', ''),
                                            row.get('LEADS_CHNAME', ''),
                                            row.get('DL_ADDRESS', ''),
                                            row.get('FINAL_AREA', '')
                                        ))
                                        if len(account_batch) >= batch_size:
                                            add_processed_accounts(audit_id, account_batch)
                                            account_batch = []
                                
                                transmittal_docs = fill_transmittal_template(temp_transmittal_path, group_df)
                                if transmittal_docs:
                                    for doc_idx, transmittal_doc in enumerate(transmittal_docs):
                                        unique_name = f"transmittal_{final_area}_{doc_idx}_{uuid.uuid4().hex[:6]}"
                                        transmittal_docx_output = area_dir / f"{unique_name}.docx"
                                        transmittal_doc.save(transmittal_docx_output)
                                        docx_files.append(str(transmittal_docx_output))
                                        if output_format == "print":
                                            if final_area not in area_docx_files:
                                                area_docx_files[final_area] = []
                                            area_docx_files[final_area].append(str(transmittal_docx_output))
                                        else:
                                            temp_files.append(str(transmittal_docx_output))
                                
                                # Enhanced progress for Transmittal Only mode
                                if session_state.get('selected_mode') == "Transmittal Only":
                                    yield json.dumps({
                                        'progress': area_progress,
                                        'message': f"Completed Transmittal for {final_area}: {len(transmittal_docs) if transmittal_docs else 0} documents created"
                                    }) + '\n'

                            if docx_files:
                                if output_format == "zip":
                                    logger.debug(f"Converting {len(docx_files)} DOCX files for {final_area} for user {user_email}")
                                    yield json.dumps({
                                        'progress': 80,
                                        'message': f"Starting PDF conversion for {final_area}: {len(docx_files)} documents..."
                                    }) + '\n'
                                    await asyncio.sleep(0)
                                    
                                    # Enhanced conversion with progress reporting
                                    def area_progress_callback(progress_data):
                                        # Store total batches for progress calculation
                                        if progress_data['type'] == 'conversion_start':
                                            progress_data['total_batches'] = progress_data.get('total_batches', 1)
                                        return progress_callback(progress_data)
                                    
                                    pdf_files = batch_convert_libreoffice(docx_files, area_dir, batch_size=300, progress_callback=area_progress_callback)
                                    
                                    if pdf_files:
                                        for pdf_file in pdf_files:
                                            if "transmittal" in Path(pdf_file).name and transmittal_pdf_merger:
                                                transmittal_pdf_merger.append(pdf_file)
                                            elif "dl" in Path(pdf_file).name and dl_pdf_merger:
                                                dl_pdf_merger.append(pdf_file)
                                        
                                        yield json.dumps({
                                            'progress': 96,
                                            'message': f"Creating merged PDFs for {final_area}..."
                                        }) + '\n'
                                        await asyncio.sleep(0)
                                        
                                        if dl_pdf_merger and session_state.get('selected_mode') in ["DL Only", "DL w/ Transmittal"]:
                                            dl_merged_path = area_dir / f"{final_area}_DL.pdf"
                                            with open(dl_merged_path, 'wb') as output_file:
                                                dl_pdf_merger.write(output_file)
                                            zipf.write(dl_merged_path, f"{final_area}_DL.pdf")
                                            logger.info(f"Created DL PDF for {final_area} for user {user_email}: {dl_merged_path}")
                                        if transmittal_pdf_merger and session_state.get('selected_mode') in ["DL w/ Transmittal", "Transmittal Only"]:
                                            transmittal_merged_path = area_dir / f"{final_area}_Transmittal.pdf"
                                            with open(transmittal_merged_path, 'wb') as output_file:
                                                transmittal_pdf_merger.write(output_file)
                                            zipf.write(transmittal_merged_path, f"{final_area}_Transmittal.pdf")
                                            logger.info(f"Created Transmittal PDF for {final_area} for user {user_email}: {transmittal_merged_path}")
                                    else:
                                        logger.warning(f"No PDFs generated for {final_area} for user {user_email}")
                                    from utils.session_management import cleanup_files
                                    cleanup_files(temp_files)
                                else:  # output_format == "print"
                                    logger.debug(f"Prepared {len(docx_files)} DOCX files for printing in {final_area} for user {user_email}")
                                    yield json.dumps({
                                        'progress': (processed_records / total_records) * 100,
                                        'message': f"Prepared {len(docx_files)} DOCX files for printing in {final_area}"
                                    }) + '\n'
                                    await asyncio.sleep(0)
                                    # Do not clean up DOCX files here for print format
                        finally:
                            if dl_pdf_merger:
                                dl_pdf_merger.close()
                            if transmittal_pdf_merger:
                                transmittal_pdf_merger.close()
                    
                    # Process any remaining accounts in the batch
                    if account_batch:
                        add_processed_accounts(audit_id, account_batch)
                    
                    if output_format == "zip":
                        yield json.dumps({
                            'progress': 98,
                            'message': "Finalizing ZIP file..."
                        }) + '\n'
                        await asyncio.sleep(0)
                        
                        zipf.close()
                        zipf = None
                        final_zip_path = user_output_dir / f"final_area_pdfs_{user_email.replace('@', '_').replace('.', '_')}.zip"
                        shutil.move(str(zip_path), str(final_zip_path))
                        session_state['zip_path'] = str(final_zip_path)
                        session_state['files_to_cleanup'].append(str(final_zip_path))
                    
                    total_time = time.time() - start_time
                    total_time_formatted = format_time_duration(total_time)
                    
                    if output_format == "zip":
                        yield json.dumps({
                            'progress': 100,
                            'message': f"Processing complete! Generated PDFs for {len(valid_rows.groupby('FINAL_AREA'))} FINAL_AREA groups in {total_time_formatted}",
                            'download_ready': True,
                            'processing_summary': {
                                'total_records': total_records,
                                'total_areas': len(valid_rows.groupby('FINAL_AREA')),
                                'total_time': total_time,
                                'total_time_formatted': total_time_formatted,
                                'processing_rate': total_records / total_time if total_time > 0 else 0
                            }
                        }) + '\n'
                    else:  # print format
                        yield json.dumps({
                            'progress': 100,
                            'message': f"Processing complete! Prepared {sum(len(files) for files in area_docx_files.values())} DOCX files for {len(valid_rows.groupby('FINAL_AREA'))} FINAL_AREA groups in {total_time_formatted}",
                            'print_ready': True,
                            'areas': list(valid_rows.groupby('FINAL_AREA').groups.keys()),
                            'docx_files': {area: files for area, files in area_docx_files.items()},
                            'processing_summary': {
                                'total_records': total_records,
                                'total_areas': len(valid_rows.groupby('FINAL_AREA')),
                                'total_files': sum(len(files) for files in area_docx_files.values()),
                                'total_time': total_time,
                                'total_time_formatted': total_time_formatted,
                                'processing_rate': total_records / total_time if total_time > 0 else 0
                            }
                        }) + '\n'
            except Exception as e:
                logger.error(f"Error during processing for user {user_email}: {e}")
                yield json.dumps({'error': f'Processing failed: {str(e)}'}) + '\n'
            finally:
                if zipf:
                    zipf.close()
    except Exception as e:
        logger.error(f"Error in generate_pdfs_stream for user {user_email}: {e}")
        yield json.dumps({'error': f'Processing failed: {str(e)}'}) + '\n'
    finally:
        session_state['processing_lock'].release()

def get_zip_for_download(user_email: str):
    """Get ZIP file for download"""
    session_state = get_user_session_state(user_email)
    zip_path = session_state.get('zip_path')
    if not zip_path or not Path(zip_path).exists():
        raise HTTPException(status_code=404, detail="ZIP file not found")
    
    # Enhanced cleanup after download
    def cleanup_after_download():
        try:
            # Clean up the ZIP file after a delay to ensure download completes
            import threading
            import time
            def delayed_cleanup():
                time.sleep(30)  # Wait 30 seconds before cleanup
                if Path(zip_path).exists():
                    Path(zip_path).unlink()
                    logger.info(f"Cleaned up ZIP file after download: {zip_path}")
            
            cleanup_thread = threading.Thread(target=delayed_cleanup)
            cleanup_thread.daemon = True
            cleanup_thread.start()
        except Exception as e:
            logger.warning(f"Failed to schedule ZIP cleanup: {e}")
    
    cleanup_after_download()
    return FileResponse(zip_path, filename=f"final_area_pdfs_{user_email.replace('@', '_').replace('.', '_')}.zip", media_type="application/zip")
