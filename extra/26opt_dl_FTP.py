import streamlit as st
import pandas as pd
from docx import Document
import re
import os
import uuid
from tempfile import NamedTemporaryFile, TemporaryDirectory
from docx.shared import Inches, Pt
from io import BytesIO
from PyPDF2 import PdfMerger, PdfReader
from barcode import Code128
from barcode.writer import ImageWriter
import logging
from ftplib import FTP
from datetime import datetime
from dotenv import load_dotenv
import gspread
from google.oauth2.service_account import Credentials
from pathlib import Path
import win32com.client
import pythoncom
import zipfile
import time

# Initialize COM
pythoncom.CoInitializeEx(pythoncom.COINIT_APARTMENTTHREADED)

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
env_path = os.path.join(os.path.dirname(__file__), "config", ".env")
load_dotenv(env_path)

# Directories
OUTPUT_DIR = Path("output").absolute()
BARCODE_DIR = Path("barcode_images").absolute()
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(BARCODE_DIR, exist_ok=True)

# FTP configuration
FTP_CONFIG = {
    "hostname": os.getenv("OMKT_FTP_HOSTNAME"),
    "port": int(os.getenv("OMKT_FTP_PORT", 21)),
    "username": os.getenv("OMKT_FTP_USERNAME"),
    "password": os.getenv("OMKT_FTP_PASSWORD")
}

# Google Sheets configuration
SERVICE_ACCOUNT_JSON = "config/dl_automation_sheet.json"
SPREADSHEET_ID = "1M0Vmmf9HfPRB0oSeJR_xUAZPPpTJ4xsR3gYDQMvnu5k"
SHEET_NAME = "LetterHeads"

# FTP Connection Manager
class FTPConnection:
    def __init__(self, hostname, port, username, password):
        self.ftp = None
        self.hostname = hostname
        self.port = port
        self.username = username
        self.password = password

    def connect(self):
        try:
            self.ftp = FTP()
            self.ftp.connect(self.hostname, self.port)
            self.ftp.login(self.username, self.password)
            logger.info("Connected to FTP server")
            return self.ftp
        except Exception as e:
            logger.error(f"Failed to connect to FTP: {e}")
            st.error(f"Failed to connect to FTP: {e}")
            return None

    def close(self):
        if self.ftp:
            try:
                self.ftp.quit()
                logger.info("Closed FTP connection")
            except Exception as e:
                logger.warning(f"Error closing FTP connection: {e}")
            self.ftp = None

    def __enter__(self):
        self.connect()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()

# Cached Google Sheets data fetching
@st.cache_data
def get_sheet_data(_service_account_json_path, _spreadsheet_id, sheet_name="LetterHeads"):
    try:
        credentials = Credentials.from_service_account_file(
            _service_account_json_path,
            scopes=["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
        )
        client = gspread.authorize(credentials)
        spreadsheet = client.open_by_key(_spreadsheet_id)
        worksheet = spreadsheet.worksheet(sheet_name)
        sheet_data = worksheet.get_all_records()
        df = pd.DataFrame(sheet_data)
        required_columns = ["CAMPAIGN", "DL TYPE", "FILE"]
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            st.error(f"Missing columns in spreadsheet: {missing_columns}")
            return pd.DataFrame()
        df = df[["CAMPAIGN", "DL TYPE", "FILE"]].dropna()
        if df.empty:
            st.warning("No valid CAMPAIGN, DL TYPE, or FILE values found.")
            return pd.DataFrame()
        logger.info(f"Retrieved sheet data with {len(df)} valid rows")
        return df
    except Exception as e:
        st.error(f"Failed to access Google Sheets: {e}")
        logger.error(f"Failed to access Google Sheets: {e}")
        return pd.DataFrame()

# Cached FTP folder retrieval
@st.cache_data
def get_ftp_folders(_ftp):
    try:
        ftp_path = "/DL AUTOMATION/Template DL/Content"
        _ftp.cwd(ftp_path)
        items = _ftp.nlst()
        folders = []
        current_dir = _ftp.pwd()
        for item in items:
            try:
                _ftp.cwd(item)
                folders.append(item)
                _ftp.cwd(current_dir)
            except:
                continue
        return sorted(folders)
    except Exception as e:
        logger.error(f"Failed to retrieve folders from FTP: {e}")
        st.error(f"Failed to retrieve folders from FTP: {e}")
        return []

# Cached FTP template retrieval
@st.cache_data
def get_ftp_templates(_ftp, folder_name):
    try:
        ftp_path = f"/DL AUTOMATION/Template DL/Content/{folder_name}"
        _ftp.cwd(ftp_path)
        templates = [item for item in _ftp.nlst() if item.lower().endswith('.docx')]
        return sorted(templates)
    except Exception as e:
        logger.error(f"Failed to retrieve templates from FTP folder {folder_name}: {e}")
        st.error(f"Failed to retrieve templates from FTP folder {folder_name}: {e}")
        return []

# Download template from FTP
def download_ftp_template(ftp, folder_name, template_name, is_header_footer=False):
    try:
        ftp_path = "/DL AUTOMATION/Template DL/Letter Head" if is_header_footer else f"/DL AUTOMATION/Template DL/Content/{folder_name}"
        ftp.cwd(ftp_path)
        with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            ftp.retrbinary(f"RETR {template_name}", tmp.write)
            tmp_path = tmp.name
        if os.path.exists(tmp_path):
            logger.info(f"Downloaded template {template_name} to: {tmp_path}")
            return tmp_path
        raise Exception(f"Failed to download template {template_name}")
    except Exception as e:
        logger.error(f"Failed to download template {template_name}: {e}")
        st.error(f"Failed to download template {template_name}: {e}")
        return None

# Fetch signature image from FTP
def fetch_signature_from_ftp(ftp):
    try:
        ftp_path = "field/DL/ATTY SIGNATURE/05-24-2025"
        ftp.cwd(ftp_path)
        with NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            ftp.retrbinary("RETR attySignature.PNG", tmp.write)
            tmp_path = tmp.name
        if os.path.exists(tmp_path):
            return tmp_path
        st.error("Signature image not found on FTP.")
        return None
    except Exception as e:
        st.error(f"Failed to fetch signature from FTP: {e}")
        logger.error(f"Failed to fetch signature from FTP: {e}")
        return None

# Combine templates
def combine_templates(header_footer_path, content_path, signature_img_path, word_app):
    try:
        header_footer_doc = Document(header_footer_path)
        content_doc = Document(content_path)
        while header_footer_doc.paragraphs:
            header_footer_doc.paragraphs[0]._element.getparent().remove(header_footer_doc.paragraphs[0]._element)
        for elem in content_doc.element.body:
            header_footer_doc.element.body.append(elem)
        for para in header_footer_doc.paragraphs:
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
        if signature_img_path and Path(signature_img_path).exists():
            with TemporaryDirectory() as temp_dir:
                temp_doc_path = Path(temp_dir) / "temp_combined_doc.docx"
                header_footer_doc.save(temp_doc_path)
                temp_doc_path = replace_in_text_boxes("«IMAGE_SIGNATURE»", signature_img_path, word_app, temp_doc_path)
                header_footer_doc = Document(temp_doc_path)
        logger.info("Combined content into header/footer template and replaced signature")
        return header_footer_doc
    except Exception as e:
        logger.error(f"Failed to combine templates or replace signature: {e}")
        st.error(f"Failed to combine templates or replace signature: {e}")
        return None

# Number to words conversion
def number_to_words(number):
    if number == 0:
        return "ZERO"
    units = ["", "ONE", "TWO", "THREE", "FOUR", "FIVE", "SIX", "SEVEN", "EIGHT", "NINE"]
    teens = ["", "ELEVEN", "TWELVE", "THIRTEEN", "FOURTEEN", "FIFTEEN", "SIXTEEN", "SEVENTEEN", "EIGHTEEN", "NINETEEN"]
    tens = ["", "TEN", "TWENTY", "THIRTY", "FORTY", "FIFTY", "SIXTY", "SEVENTY", "EIGHTY", "NINETY"]
    thousands = ["", "THOUSAND", "MILLION"]
    def convert_hundreds(num):
        if num == 0:
            return ""
        result = []
        if num >= 100:
            result.append(units[num // 100] + " HUNDRED")
            num %= 100
            if num > 0:
                result.append("AND")
        if num >= 20:
            result.append(tens[num // 10])
            num %= 10
            if num > 0:
                result.append(units[num])
        elif num >= 11:
            result.append(teens[num - 10])
        elif num == 10:
            result.append(tens[1])
        elif num > 0:
            result.append(units[num])
        return " ".join(result)
    result = []
    i = 0
    while number > 0:
        chunk = number % 1000
        if chunk > 0:
            chunk_words = convert_hundreds(chunk)
            if thousands[i]:
                chunk_words += f" {thousands[i]}"
            result.insert(0, chunk_words)
        number //= 1000
        i += 1
    return ", ".join(result).replace(", AND", " AND").strip()

# Convert amount to words
def amount_to_words(amount_str):
    try:
        amount = float(amount_str.replace(",", ""))
        pesos = int(amount)
        cents = int(round((amount - pesos) * 100))
        pesos_words = number_to_words(pesos) if pesos > 0 else "ZERO"
        cents_words = number_to_words(cents) if cents > 0 else "ZERO"
        result = f"{pesos_words} PESOS"
        if cents > 0:
            result += f", AND {cents_words} CENTS"
        else:
            result += ", AND ZERO CENTS"
        return result
    except Exception as e:
        logger.error(f"Failed to convert amount to words: {amount_str}, error: {e}")
        st.error(f"Failed to convert amount to words: {amount_str}")
        return "ERROR CONVERTING AMOUNT"

# Generate barcode image
def generate_barcode(barcode_value):
    try:
        barcode = Code128(barcode_value, writer=ImageWriter())
        buffer = BytesIO()
        barcode.write(buffer, options={"write_text": False, "module_width": 1, "module_height": 8, "quiet_zone": 2.0})
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Failed to generate barcode for {barcode_value}: {e}")
        return None

# Extract placeholders from .docx
def extract_placeholders(doc):
    placeholders = set()
    def extract_from_text(text):
        if text and "«" in text:
            matches = re.findall(r"«(.*?)»", text)
            placeholders.update(["«" + m.strip() + "»" for m in matches])
    for para in doc.paragraphs:
        extract_from_text(para.text)
    for node in doc._element.iter():
        if node.tag.endswith("}t") and node.text:
            extract_from_text(node.text)
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    extract_from_text(para.text)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                extract_from_text(para.text)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    extract_from_text(para.text)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                extract_from_text(para.text)
    return sorted(placeholders)

# Replace text in text boxes
def replace_in_text_boxes(find_str, replace_with_image_path, word_app, temp_doc_path):
    try:
        word_doc = word_app.Documents.Open(str(temp_doc_path.absolute()))
        modified = False
        for shape in word_doc.Shapes:
            if shape.TextFrame.HasText:
                text_range = shape.TextFrame.TextRange
                find = text_range.Find
                find.Text = find_str
                find.Forward = True
                find.Wrap = 1
                while find.Execute():
                    found_range = text_range.Duplicate
                    found_range.Find.Execute(FindText=find_str)
                    found_range.Text = ""
                    inline_shape = found_range.InlineShapes.AddPicture(
                        FileName=str(Path(replace_with_image_path).absolute()),
                        LinkToFile=False,
                        SaveWithDocument=True
                    )
                    inline_shape.Width = 110
                    inline_shape.Height = 50
                    modified = True
        word_doc.Save()
        word_doc.Close(SaveChanges=True)
        if modified:
            logger.info(f"Replaced '{find_str}' with image in shapes.")
        else:
            logger.info("No placeholders found in shapes.")
        return temp_doc_path
    except Exception as e:
        logger.error(f"Error replacing text in shapes: {e}")
        st.error(f"Error replacing text in shapes: {e}")
        try:
            word_doc.Close(SaveChanges=False)
        except:
            pass
        return temp_doc_path

# Fill template with data
def fill_template(doc, mapping, barcode_buffer=None):
    def replace_in_text(text):
        for k, v in mapping.items():
            if k == "«IMAGE_BARCODE»" and v:
                continue
            text = text.replace(k, str(v))
        return text
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = replace_in_text(run.text)
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = replace_in_text(run.text)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if "«IMAGE_BARCODE»" in para.text and barcode_buffer:
                                    para.clear()
                                    cell.paragraphs[0].paragraph_format.left_indent = Pt(-20)
                                    run = para.add_run()
                                    run.add_picture(barcode_buffer, width=Inches(3.0), height=Inches(0.35))
                                else:
                                    for run in para.runs:
                                        if run.text:
                                            run.text = replace_in_text(run.text)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = replace_in_text(run.text)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    if run.text:
                                        run.text = replace_in_text(run.text)
    return doc

# Read Excel file
def get_raw_file(file, sheet_name=None, engine=None):
    try:
        if sheet_name is None:
            return pd.read_excel(file, dtype=str, engine=engine)
        return pd.read_excel(file, sheet_name=sheet_name, engine=engine, dtype=str)
    except Exception as e:
        st.error(f"Error reading file: {e}")
        logger.error(f"Error reading file: {e}")
        return pd.DataFrame()

# Cleanup individual files
def cleanup_files(file_paths, output_dir=OUTPUT_DIR, retries=3, delay=1):
    os.system("taskkill /IM WINWORD.EXE /F >nul 2>&1")
    # Delete specified file paths
    for file_path in file_paths:
        if not file_path or not os.path.exists(file_path):
            continue
        for attempt in range(retries):
            try:
                if os.path.isfile(file_path):
                    os.remove(file_path)
                    logger.info(f"Deleted file: {file_path}")
                break
            except PermissionError:
                logger.warning(f"Permission denied for {file_path}, retrying {attempt + 1}/{retries}...")
                time.sleep(delay)
            except Exception as e:
                logger.warning(f"Failed to delete file {file_path}: {e}")
                break
    # Delete all files in output_dir
    if os.path.exists(output_dir):
        for file_path in output_dir.glob('*'):
            if file_path.is_file():
                for attempt in range(retries):
                    try:
                        file_path.unlink()
                        logger.info(f"Deleted file in output directory: {file_path}")
                        break
                    except PermissionError:
                        logger.warning(f"Permission denied for {file_path}, retrying {attempt + 1}/{retries}...")
                        time.sleep(delay)
                    except Exception as e:
                        logger.warning(f"Failed to delete file {file_path}: {e}")
                        break
    # Ensure output_dir exists
    os.makedirs(output_dir, exist_ok=True)
    logger.info(f"Ensured output directory exists: {output_dir}")

# Batch convert DOCX to PDF
def convert_batch_with_word_retry(batch_files, output_dir, batch_id, word_app, max_retries=2):
    wdFormatPDF = 17
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    for attempt in range(max_retries + 1):
        logger.info(f"Batch {batch_id} (Attempt {attempt + 1}): Converting {len(batch_files)} files...")
        batch_pdfs = []
        failed_files = []
        for docx_path in batch_files:
            try:
                docx_path = Path(docx_path)
                pdf_path = output_dir / f"{docx_path.stem}.pdf"
                doc = word_app.Documents.Open(str(docx_path), ReadOnly=True)
                doc.SaveAs(str(pdf_path), FileFormat=wdFormatPDF)
                doc.Close(False)
                batch_pdfs.append(str(pdf_path))
                time.sleep(0.1)
            except Exception as e:
                logger.error(f"Failed: {docx_path} - Open.SaveAs: {e}")
                failed_files.append(str(docx_path))
        success_rate = len(batch_pdfs) / len(batch_files) * 100
        logger.info(f"Batch {batch_id} result: {len(batch_pdfs)}/{len(batch_files)} successful ({success_rate:.1f}%)")
        if failed_files and attempt < max_retries:
            logger.info(f"Retrying {len(failed_files)} failed files...")
            batch_files = failed_files
            time.sleep(1)
            continue
        return batch_pdfs, failed_files
    return [], batch_files

def batch_convert_word(docx_files, output_dir, batch_size=350):
    if not docx_files:
        logger.info("No DOCX files provided for conversion.")
        return []
    pdf_files = []
    total_failed = []
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    batches = [docx_files[i:i + batch_size] for i in range(0, len(docx_files), batch_size)]
    logger.info(f"Starting conversion of {len(docx_files)} DOCX files in {len(batches)} batches (size: {batch_size}).")
    start_time = time.time()
    for batch_id, batch in enumerate(batches, 1):
        logger.info(f"Processing batch {batch_id}/{len(batches)} with {len(batch)} files.")
        pythoncom.CoInitialize()
        word_app = None
        try:
            word_app = win32com.client.DispatchEx("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = False
            batch_pdfs, batch_failed = convert_batch_with_word_retry(batch, output_dir, batch_id, word_app)
            pdf_files.extend(batch_pdfs)
            total_failed.extend(batch_failed)
        except Exception as e:
            logger.error(f"Error in batch {batch_id}: {e}")
            total_failed.extend(batch)
        finally:
            if word_app:
                try:
                    word_app.Quit()
                    logger.info(f"Word application closed for batch {batch_id}.")
                except Exception as e:
                    logger.error(f"Failed to quit Word for batch {batch_id}: {e}")
            pythoncom.CoUninitialize()
            os.system("taskkill /IM WINWORD.EXE /F >nul 2>&1")
        completed_files = len(pdf_files) + len(total_failed)
        progress = completed_files / len(docx_files) * 100
        elapsed = time.time() - start_time
        rate = completed_files / elapsed if elapsed > 0 else 0
        logger.info(f"Batch {batch_id} progress: {completed_files}/{len(docx_files)} ({progress:.1f}%), Success: {len(pdf_files)}, Failed: {len(total_failed)}")
    total_time = time.time() - start_time
    success_rate = len(pdf_files) / len(docx_files) * 100
    logger.info(f"Conversion summary: Total {len(docx_files)}, Successful {len(pdf_files)} ({success_rate:.1f}%), Failed {len(total_failed)}, Time {total_time:.1f}s")
    if total_failed:
        logger.warning(f"Failed files: {total_failed[:10]}{' and more' if len(total_failed) > 10 else ''}")
    return pdf_files

# Main Streamlit app
def main():
    st.title("📄 DL GENERATOR")
    files_to_cleanup = []

    # Initialize session state
    if "selected_folder" not in st.session_state:
        st.session_state.selected_folder = ""
    if "selected_dl_type" not in st.session_state:
        st.session_state.selected_dl_type = ""
    if "selected_template" not in st.session_state:
        st.session_state.selected_template = ""
    if "base_template" not in st.session_state:
        st.session_state.base_template = None
    if "template_path" not in st.session_state:
        st.session_state.template_path = None
    if "header_footer_template_path" not in st.session_state:
        st.session_state.header_footer_template_path = None
    if "placeholders" not in st.session_state:
        st.session_state.placeholders = []
    if "download_completed" not in st.session_state:
        st.session_state.download_completed = False
    if "files_to_cleanup" not in st.session_state:
        st.session_state.files_to_cleanup = []

    # Validate FTP configuration
    if not all([FTP_CONFIG["hostname"], FTP_CONFIG["username"], FTP_CONFIG["password"]]):
        st.error("FTP configuration is incomplete. Check your .env file.")
        cleanup_files(files_to_cleanup)
        return

    # Initialize FTP connection
    with FTPConnection(FTP_CONFIG["hostname"], FTP_CONFIG["port"], FTP_CONFIG["username"], FTP_CONFIG["password"]) as ftp_conn:
        ftp = ftp_conn.connect()
        if not ftp:
            st.error("Failed to establish FTP connection.")
            cleanup_files(files_to_cleanup)
            return

        # Fetch signature image
        signature_img_path = fetch_signature_from_ftp(ftp)
        if signature_img_path:
            files_to_cleanup.append(signature_img_path)
            st.session_state.files_to_cleanup.append(signature_img_path)
        else:
            st.error("Signature image not found. Check FTP connection.")
            cleanup_files(files_to_cleanup)
            return

        # Fetch Google Sheets data
        sheet_df = get_sheet_data(SERVICE_ACCOUNT_JSON, SPREADSHEET_ID, SHEET_NAME)
        if sheet_df.empty:
            st.error("Failed to retrieve data from Google Sheets.")
            cleanup_files(files_to_cleanup)
            return

        # Fetch folder names from FTP
        st.subheader("Select Template Folder")
        folders = get_ftp_folders(ftp)
        if not folders:
            st.error("No folders found in FTP directory.")
            cleanup_files(files_to_cleanup)
            return

        # Pre-select folder
        campaign_values = sheet_df["CAMPAIGN"].dropna().unique().tolist()
        default_folder = campaign_values[0] if campaign_values else None
        default_index = folders.index(default_folder) + 1 if default_folder in folders else 0
        if default_folder and default_index == 0:
            st.warning(f"CAMPAIGN folder '{default_folder}' not found in FTP.")
        selected_folder = st.selectbox("Choose a folder:", [""] + folders, index=default_index, key="folder_select")
        if selected_folder != st.session_state.selected_folder:
            st.session_state.selected_folder = selected_folder
            st.session_state.selected_dl_type = ""
            st.session_state.selected_template = ""
            st.session_state.base_template = None
            st.session_state.template_path = None
            st.session_state.header_footer_template_path = None
            st.session_state.placeholders = []
            st.session_state.download_completed = False
            st.session_state.files_to_cleanup = [signature_img_path] if signature_img_path else []

        # Select DL TYPE
        if selected_folder:
            dl_types = sorted(sheet_df[sheet_df["CAMPAIGN"] == selected_folder]["DL TYPE"].dropna().unique().tolist())
            if not dl_types:
                st.warning(f"No DL TYPE values found for CAMPAIGN '{selected_folder}'.")
            else:
                st.subheader("Select DL TYPE")
                selected_dl_type = st.selectbox("Choose a DL TYPE:", [""] + dl_types, key="dl_type_select")
                if selected_dl_type != st.session_state.selected_dl_type:
                    st.session_state.selected_dl_type = selected_dl_type
                    st.session_state.selected_template = ""
                    st.session_state.base_template = None
                    st.session_state.template_path = None
                    st.session_state.header_footer_template_path = None
                    st.session_state.placeholders = []
                    st.session_state.download_completed = False
                    st.session_state.files_to_cleanup = [signature_img_path] if signature_img_path else []

        # Template selection
        if selected_folder and selected_dl_type:
            templates = get_ftp_templates(ftp, selected_folder)
            if not templates:
                st.error(f"No .docx templates found in folder '{selected_folder}'.")
                cleanup_files(st.session_state.files_to_cleanup)
                return
            st.subheader("Select Content File")
            selected_template = st.selectbox("Choose a template:", [""] + templates, key="template_select")
            if selected_template != st.session_state.selected_template or not st.session_state.base_template:
                st.session_state.selected_template = selected_template
                st.session_state.base_template = None
                st.session_state.template_path = None
                st.session_state.header_footer_template_path = None
                st.session_state.placeholders = []
                st.session_state.download_completed = False
                st.session_state.files_to_cleanup = [signature_img_path] if signature_img_path else []
                if selected_template:
                    pythoncom.CoInitialize()
                    word_app = None
                    try:
                        word_app = win32com.client.DispatchEx("Word.Application")
                        word_app.Visible = False
                        word_app.DisplayAlerts = False
                        template_path = download_ftp_template(ftp, selected_folder, selected_template, is_header_footer=False)
                        if not template_path:
                            st.error("Failed to download content template.")
                            cleanup_files(st.session_state.files_to_cleanup)
                            return
                        files_to_cleanup.append(template_path)
                        st.session_state.files_to_cleanup.append(template_path)
                        st.session_state.template_path = template_path
                        matching_row = sheet_df[(sheet_df["CAMPAIGN"] == selected_folder) & (sheet_df["DL TYPE"] == selected_dl_type)]
                        if matching_row.empty:
                            st.error(f"No header/footer template found for CAMPAIGN '{selected_folder}' and DL TYPE '{selected_dl_type}'.")
                            cleanup_files(st.session_state.files_to_cleanup)
                            return
                        header_footer_filename = matching_row["FILE"].iloc[0]
                        if not header_footer_filename.lower().endswith('.docx'):
                            header_footer_filename += '.docx'
                        header_footer_template_path = download_ftp_template(ftp, None, header_footer_filename, is_header_footer=True)
                        if not header_footer_template_path:
                            st.error(f"Failed to download header/footer template '{header_footer_filename}'.")
                            cleanup_files(st.session_state.files_to_cleanup)
                            return
                        files_to_cleanup.append(header_footer_template_path)
                        st.session_state.files_to_cleanup.append(header_footer_template_path)
                        st.session_state.header_footer_template_path = header_footer_template_path
                        base_template = combine_templates(header_footer_template_path, template_path, signature_img_path, word_app)
                        if not base_template:
                            st.error("Failed to combine templates.")
                            cleanup_files(st.session_state.files_to_cleanup)
                            return
                        st.session_state.base_template = base_template
                        st.session_state.placeholders = extract_placeholders(base_template)
                        st.success(f"Loaded content template '{selected_template}' and header/footer template '{header_footer_filename}'.")
                        st.write("🔍 Detected placeholders:", st.session_state.placeholders)
                    except Exception as e:
                        st.error(f"Failed to process templates: {e}")
                        logger.error(f"Failed to process templates: {e}")
                        cleanup_files(st.session_state.files_to_cleanup)
                        return
                    finally:
                        if word_app:
                            try:
                                word_app.Quit()
                                logger.info("Word application closed after template processing.")
                            except Exception as e:
                                logger.error(f"Failed to quit Word application: {e}")
                        pythoncom.CoUninitialize()
                        os.system("taskkill /IM WINWORD.EXE /F >nul 2>&1")

        # Process Excel file and generate PDFs
        if st.session_state.base_template and st.session_state.template_path and st.session_state.header_footer_template_path:
            uploaded_excel = st.file_uploader("Upload Excel file with matching column headers", type=["xlsx"], key="excel_uploader")
            if uploaded_excel:
                df = get_raw_file(uploaded_excel)
                st.write("📋 Data preview:")
                st.write(df)
                if 'FINAL_AREA' not in df.columns:
                    st.error("Excel file must contain a 'FINAL_AREA' column.")
                    cleanup_files(st.session_state.files_to_cleanup)
                    return
                if st.button("🔄 Generate PDFs by FINAL_AREA"):
                    try:
                        os.remove(r"C:\Users\SPM\Desktop\ONLY SAVE FILE HERE\mail\output")
                    except:     
                        pass
                    st.info("Processing...")
                    st.session_state.download_completed = False
                    st.session_state.files_to_cleanup = [signature_img_path, st.session_state.template_path, st.session_state.header_footer_template_path] if signature_img_path else [st.session_state.template_path, st.session_state.header_footer_template_path]
                    today_date = datetime.now().strftime("%B %d, %Y")
                    valid_rows = df[df['LEADS_CHNAME'].notna()]
                    total_rows = len(valid_rows)
                    if total_rows == 0:
                        st.error("No valid rows found (LEADS_CHNAME missing).")
                        cleanup_files(st.session_state.files_to_cleanup)
                        return
                    try:
                        progress_bar = st.progress(0)
                        progress_text = st.empty()
                        row_counter = 0
                        pdf_paths = []
                        start_time = time.time()
                        with TemporaryDirectory() as temp_dir:
                            temp_doc_path = Path(temp_dir) / "base_template.docx"
                            st.session_state.base_template.save(temp_doc_path)
                            st.session_state.files_to_cleanup.append(temp_doc_path)
                            for final_area, group_df in valid_rows.groupby('FINAL_AREA'):
                                st.write(f"Processing FINAL_AREA: {final_area}")
                                docx_files = []
                                temp_files = []
                                pdf_merger = PdfMerger()
                                try:
                                    for idx, row in group_df.iterrows():
                                        row_counter += 1
                                        progress_bar.progress(row_counter / total_rows)
                                        elapsed = time.time() - start_time
                                        rate = row_counter / elapsed if elapsed > 0 else 0
                                        eta = (total_rows - row_counter) / rate if rate > 0 else 0
                                        progress_text.text(f"Creating DOCX {row_counter}/{total_rows} | Rate: {rate:.1f}/sec | ETA: {eta/60:.1f}min")
                                        barcode_buffer = None
                                        if barcode_value := row.get('DL_CODE', ''):
                                            barcode_buffer = generate_barcode(barcode_value)
                                        amount_words = amount_to_words(row.get('amount', '0.00'))
                                        mapping = {f"«{col.upper()}»": row[col] for col in df.columns if pd.notnull(row[col])}
                                        mapping.update({
                                            "«IMAGE_BARCODE»": barcode_buffer or "",
                                            "«DL_DATE»": today_date,
                                            "«AMOUNT_ABBR»": amount_words,
                                            "«IMAGE_SIGNATURE»": signature_img_path or ""
                                        })
                                        filled_doc = Document(temp_doc_path)
                                        filled_doc = fill_template(filled_doc, mapping, barcode_buffer)
                                        if filled_doc:
                                            unique_name = f"doc_{final_area}_{row_counter}_{uuid.uuid4().hex[:6]}"
                                            docx_output = OUTPUT_DIR / f"{unique_name}.docx"
                                            filled_doc.save(docx_output)
                                            docx_files.append(str(docx_output))
                                            temp_files.append(str(docx_output))
                                    if docx_files:
                                        st.write(f"Converting {len(docx_files)} DOCX files to PDF for {final_area}...")
                                        conversion_start = time.time()
                                        pdf_files = batch_convert_word(docx_files, OUTPUT_DIR)
                                        conversion_time = time.time() - conversion_start
                                        if pdf_files:
                                            conversion_rate = len(pdf_files) / conversion_time
                                            st.write(f"✅ Converted {len(pdf_files)} PDFs in {conversion_time:.1f}s ({conversion_rate:.1f} PDFs/sec)")
                                            for pdf_file in pdf_files:
                                                pdf_merger.append(pdf_file)
                                                temp_files.append(pdf_file)
                                        merged_pdf_path = OUTPUT_DIR / f"{final_area}.pdf"
                                        with open(merged_pdf_path, "wb") as f:
                                            pdf_merger.write(f)
                                        if merged_pdf_path.exists():
                                            with open(merged_pdf_path, "rb") as f:
                                                pdf_reader = PdfReader(f)
                                                st.write(f"📄 Final PDF for '{final_area}': {len(pdf_reader.pages)} pages")
                                            pdf_paths.append(str(merged_pdf_path))
                                            temp_files.append(str(merged_pdf_path))
                                finally:
                                    pdf_merger.close()
                                    st.session_state.files_to_cleanup.extend(temp_files)
                            if pdf_paths:
                                zip_path = OUTPUT_DIR / "final_area_pdfs.zip"
                                with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                                    for pdf_path in pdf_paths:
                                        zipf.write(pdf_path, os.path.basename(pdf_path))
                                st.session_state.files_to_cleanup.append(str(zip_path))
                                if zip_path.exists():
                                    total_time = time.time() - start_time
                                    overall_rate = total_rows / total_time
                                    st.success(f"🎉 COMPLETED! {total_rows} documents in {total_time/60:.1f} minutes ({overall_rate:.1f} docs/sec)")
                                    with open(zip_path, "rb") as f:
                                        zip_data = f.read()
                                    st.download_button(
                                        label="⬇️ Download ZIP with all FINAL_AREA PDFs",
                                        data=zip_data,
                                        file_name="final_area_pdfs.zip",
                                        mime="application/zip",
                                        key="download_zip"
                                    )
                                    st.session_state.download_completed = True
                            progress_bar.empty()
                            progress_text.empty()
                        # Cleanup button appears after download is available
                        if st.session_state.download_completed and st.button("🗑️ Cleanup Files", key="cleanup_button"):
                            # cleanup_files(st.session_state.files_to_cleanup)
                            st.session_state.download_completed = False
                            st.session_state.files_to_cleanup = [signature_img_path] if signature_img_path else []
                            st.success("Files cleaned up successfully!")
                    except Exception as e:
                        st.error(f"Error during processing: {e}")
                        logger.error(f"Error during processing: {e}")
                        # cleanup_files(st.session_state.files_to_cleanup)
                        st.session_state.download_completed = False
                        st.session_state.files_to_cleanup = [signature_img_path] if signature_img_path else []
                    finally:
                        os.system("taskkill /IM WINWORD.EXE /F >nul 2>&1")
                    try:
                        os.remove(r"C:\Users\SPM\Desktop\ONLY SAVE FILE HERE\mail\output")
                    except:     
                        pass
        else:
            # cleanup_files(st.session_state.files_to_cleanup)
            logger.error(f"Error during processing:")
            # os.remove(OUTPUT_DIR)


if __name__ == "__main__":
    main()