from fastapi import FastAPI, HTTPException, UploadFile, File, Depends, Response
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional, Dict
import pandas as pd
from docx import Document
import re
import os
import uuid
from pathlib import Path
from tempfile import NamedTemporaryFile, TemporaryDirectory
from docx.shared import Inches, Pt
from io import BytesIO
from PyPDF2 import PdfMerger
from barcode import Code128
from barcode.writer import ImageWriter
import shutil
import logging
from datetime import datetime
from dotenv import load_dotenv
import gspread
from google.oauth2.service_account import Credentials
import win32com.client
import pythoncom
import zipfile
from ftplib import FTP
import time
import psutil
import qrcode
import json
import asyncio
import subprocess
# Initialize FastAPI app
app = FastAPI(title="DL Generator API")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load environment variables
env_path = Path(__file__).parent / "config" / ".env"
load_dotenv(env_path)

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Create output and barcode directories
OUTPUT_DIR = Path("output").absolute()
BARCODE_DIR = Path("barcode_images").absolute()
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(BARCODE_DIR, exist_ok=True)

FTP_CONFIG = {
    "hostname": os.getenv("OMKT_FTP_HOSTNAME"),
    "port": int(os.getenv("OMKT_FTP_PORT", 21)),
    "username": os.getenv("OMKT_FTP_USERNAME"),
    "password": os.getenv("OMKT_FTP_PASSWORD")
}

SERVICE_ACCOUNT_JSON = "config/dl_automation_sheet.json"
SPREADSHEET_ID = "1M0Vmmf9HfPRB0oSeJR_xUAZPPpTJ4xsR3gYDQMvnu5k"
SHEET_NAME = "LetterHeads"

# Pydantic Models
class LoginRequest(BaseModel):
    username: str
    password: str

class User(BaseModel):
    username: str
    role: str

class FolderRequest(BaseModel):
    folder: str

class TemplateRequest(BaseModel):
    folder: str
    dl_type: str
    template: str

class ModeRequest(BaseModel):
    mode: str

# Mock user database
USERS = {
    'admin': {'password': 'admin123', 'role': 'Admin'},
    'user': {'password': 'user123', 'role': 'User'}
}

# Session state simulation
SESSION_STATE = {
    'selected_mode': '',
    'selected_folder': '',
    'selected_dl_type': '',
    'selected_template': '',
    'base_template': None,
    'template_path': None,
    'header_footer_template_path': None,
    'transmittal_template_path': None,
    'placeholders': [],
    'download_completed': False,
    'files_to_cleanup': [],
    'zip_path': None
}

class FTPConnection:
    def __init__(self, hostname, port, username, password):
        self.ftp = None
        self.hostname = hostname
        self.port = port
        self.username = username
        self.password = password

    def connect(self):
        try:
            self.ftp = FTP()
            self.ftp.connect(self.hostname, self.port)
            self.ftp.login(self.username, self.password)
            logger.info("Connected to FTP server")
            return self.ftp
        except Exception as e:
            logger.error(f"Failed to connect to FTP: {e}")
            return None

    def close(self):
        if self.ftp:
            try:
                self.ftp.quit()
                logger.info("Closed FTP connection")
            except Exception as e:
                logger.warning(f"Error closing FTP connection: {e}")
            self.ftp = None

    def __enter__(self):
        self.connect()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()

def get_ftp_folders(ftp):
    try:
        ftp_path = "/DL AUTOMATION/Template DL V2/Content"
        ftp.cwd(ftp_path)
        items = ftp.nlst()
        folders = []
        current_dir = ftp.pwd()
        for item in items:
            try:
                ftp.cwd(item)
                folders.append(item)
                ftp.cwd(current_dir)
            except:
                continue
        return sorted(folders)
    except Exception as e:
        logger.error(f"Failed to retrieve folders from FTP: {e}")
        return []

def get_ftp_templates(ftp, folder_name):
    try:
        ftp_path = f"/DL AUTOMATION/Template DL V2/Content/{folder_name}"
        ftp.cwd(ftp_path)
        templates = [item for item in ftp.nlst() if item.lower().endswith('.docx')]
        return sorted(templates)
    except Exception as e:
        logger.error(f"Failed to retrieve templates from FTP folder {folder_name}: {e}")
        return []

def download_ftp_template(ftp, folder_name, template_name, is_header_footer=False, is_transmittal=False):
    try:
        if is_transmittal:
            ftp_path = "/DL AUTOMATION/Template Transmittal V2"
        elif is_header_footer:
            ftp_path = "/DL AUTOMATION/Template DL V2/Letter Head"
        else:
            ftp_path = f"/DL AUTOMATION/Template DL V2/Content/{folder_name}"
        ftp.cwd(ftp_path)
        with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            ftp.retrbinary(f"RETR {template_name}", tmp.write)
            tmp_path = tmp.name
        if os.path.exists(tmp_path):
            logger.info(f"Downloaded template {template_name} to: {tmp_path}")
            return tmp_path
        raise Exception(f"Failed to download template {template_name}")
    except Exception as e:
        logger.error(f"Failed to download template {template_name}: {e}")
        return None

def fetch_signature_from_ftp(ftp):
    try:
        ftp_path = f"field/DL/ATTY SIGNATURE/05-29-2025"
        ftp.cwd(ftp_path)
        with NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            ftp.retrbinary("RETR attySignature.PNG", tmp.write)
            tmp_path = tmp.name
        if os.path.exists(tmp_path):
            return tmp_path
        return None
    except Exception as e:
        logger.error(f"Failed to fetch signature from FTP: {e}")
        return None

def cleanup_files(file_paths):
    for file_path in file_paths:
        try:
            if file_path and Path(file_path).exists():
                Path(file_path).unlink()
                logger.info(f"Deleted file: {file_path}")
        except Exception as e:
            logger.warning(f"Failed to delete file {file_path}: {e}")

def get_sheet_data(service_account_json_path, spreadsheet_id, sheet_name="LetterHeads"):
    try:
        credentials = Credentials.from_service_account_file(
            service_account_json_path,
            scopes=[
                "https://www.googleapis.com/auth/spreadsheets",
                "https://www.googleapis.com/auth/drive"
            ]
        )
        client = gspread.authorize(credentials)
        spreadsheet = client.open_by_key(spreadsheet_id)
        worksheet = spreadsheet.worksheet(sheet_name)
        sheet_data = worksheet.get_all_records()
        df = pd.DataFrame(sheet_data)
        required_columns = ["CAMPAIGN", "DL TYPE", "FILE"]
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            logger.error(f"Missing columns in spreadsheet: {missing_columns}")
            return pd.DataFrame()
        df = df[["CAMPAIGN", "DL TYPE", "FILE"]].dropna()
        logger.info(f"Retrieved sheet data with {len(df)} valid rows")
        return df
    except Exception as e:
        logger.error(f"Failed to access Google Sheets: {e}")
        return pd.DataFrame()

def combine_templates(header_footer_path, content_path, signature_img_path, word_app):
    try:
        header_footer_doc = Document(header_footer_path)
        content_doc = Document(content_path)
        while header_footer_doc.paragraphs:
            header_footer_doc.paragraphs[0]._element.getparent().remove(header_footer_doc.paragraphs[0]._element)
        for elem in content_doc.element.body:
            header_footer_doc.element.body.append(elem)
        for para in header_footer_doc.paragraphs:
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
        if signature_img_path and Path(signature_img_path).exists():
            with TemporaryDirectory() as temp_dir:
                temp_doc_path = Path(temp_dir) / "temp_combined_doc.docx"
                header_footer_doc.save(temp_doc_path)
                temp_doc_path = replace_in_text_boxes(header_footer_doc, "«IMAGE_SIGNATURE»", signature_img_path, word_app, temp_doc_path)
                header_footer_doc = Document(temp_doc_path)
        return header_footer_doc
    except Exception as e:
        logger.error(f"Failed to combine templates or replace signature: {e}")
        return None

def number_to_words(number):
    if number == 0:
        return "ZERO"
    units = ["", "ONE", "TWO", "THREE", "FOUR", "FIVE", "SIX", "SEVEN", "EIGHT", "NINE"]
    teens = ["", "ELEVEN", "TWELVE", "THIRTEEN", "FOURTEEN", "FIFTEEN", "SIXTEEN", "SEVENTEEN", "EIGHTEEN", "NINETEEN"]
    tens = ["", "TEN", "TWENTY", "THIRTY", "FORTY", "FIFTY", "SIXTY", "SEVENTY", "EIGHTY", "NINETY"]
    thousands = ["", "THOUSAND", "MILLION"]
    def convert_hundreds(num):
        if num == 0:
            return ""
        result = []
        if num >= 100:
            result.append(units[num // 100] + " HUNDRED")
            num %= 100
            if num > 0:
                result.append("AND")
        if num >= 20:
            result.append(tens[num // 10])
            num %= 10
            if num > 0:
                result.append(units[num])
        elif num >= 11:
            result.append(teens[num - 10])
        elif num == 10:
            result.append(tens[1])
        elif num > 0:
            result.append(units[num])
        return " ".join(result)
    result = []
    i = 0
    while number > 0:
        chunk = number % 1000
        if chunk > 0:
            chunk_words = convert_hundreds(chunk)
            if thousands[i]:
                chunk_words += f" {thousands[i]}"
            result.insert(0, chunk_words)
        number //= 1000
        i += 1
    return ", ".join(result).replace(", AND", " AND").strip()

def amount_to_words(amount_str):
    try:
        amount = float(amount_str.replace(",", ""))
        pesos = int(amount)
        cents = int(round((amount - pesos) * 100))
        pesos_words = number_to_words(pesos)
        if pesos == 0:
            pesos_words = "ZERO"
        cents_words = number_to_words(cents) if cents > 0 else "ZERO"
        result = f"{pesos_words} PESOS"
        if cents > 0:
            result += f", AND {cents_words} CENTS"
        else:
            result += ", AND ZERO CENTS"
        return result
    except Exception as e:
        logger.error(f"Failed to convert amount to words: {amount_str}, error: {e}")
        return "ERROR CONVERTING AMOUNT"

def generate_barcode(barcode_value):
    try:
        barcode = Code128(barcode_value, writer=ImageWriter())
        buffer = BytesIO()
        barcode.write(buffer, options={"write_text": False, "module_width": 1, "module_height": 8, "quiet_zone": 2.0})
        buffer.seek(0)
        return buffer
    except Exception as e:
        logger.error(f"Failed to generate barcode for {barcode_value}: {e}")
        return None

def generate_qrcode(data):
    try:
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        qr.add_data(data)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        buffer = BytesIO()
        img.save(buffer, format="PNG")
        buffer.seek(0)
        return buffer
    except Exception as e:
        logger.error(f"Failed to generate QR code for {data}: {e}")
        return None

def extract_placeholders(doc):
    placeholders = set()
    for para in doc.paragraphs:
        if "«" in para.text:
            matches = re.findall(r"«(.*?)»", para.text)
            placeholders.update(["«" + m.strip() + "»" for m in matches])
    for node in doc._element.iter():
        if node.tag.endswith("}t") and node.text and "«" in node.text:
            matches = re.findall(r"«(.*?)»", node.text)
            placeholders.update(["«" + m.strip() + "»" for m in matches])
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header is not None:
                for para in header.paragraphs:
                    if "«" in para.text:
                        matches = re.findall(r"«(.*?)»", para.text)
                        placeholders.update(["«" + m.strip() + "»" for m in matches])
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if "«" in para.text:
                                    matches = re.findall(r"«(.*?)»", para.text)
                                    placeholders.update(["«" + m.strip() + "»" for m in matches])
                for node in header._element.iter():
                    if node.tag.endswith("}t") and node.text and "«" in node.text:
                        matches = re.findall(r"«(.*?)»", node.text)
                        placeholders.update(["«" + m.strip() + "»" for m in matches])
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer is not None:
                for para in footer.paragraphs:
                    if "«" in para.text:
                        matches = re.findall(r"«(.*?)»", para.text)
                        placeholders.update(["«" + m.strip() + "»" for m in matches])
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if "«" in para.text:
                                    matches = re.findall(r"«(.*?)»", para.text)
                                    placeholders.update(["«" + m.strip() + "»" for m in matches])
                for node in footer._element.iter():
                    if node.tag.endswith("}t") and node.text and "«" in node.text:
                        matches = re.findall(r"«(.*?)»", node.text)
                        placeholders.update(["«" + m.strip() + "»" for m in matches])
    return sorted(placeholders)

def replace_in_text_boxes(doc, find_str, replace_with_image_path, word_app, temp_doc_path):
    try:
        word_doc = word_app.Documents.Open(str(temp_doc_path.absolute()))
        modified = False
        for shape in word_doc.Shapes:
            if shape.TextFrame.HasText:
                text_range = shape.TextFrame.TextRange
                find = text_range.Find
                find.Text = find_str
                find.Forward = True
                find.Wrap = 1
                while find.Execute():
                    found_range = text_range.Duplicate
                    found_range.Find.Execute(FindText=find_str)
                    found_range.Text = ""
                    inline_shape = found_range.InlineShapes.AddPicture(
                        FileName=str(Path(replace_with_image_path).absolute()),
                        LinkToFile=False,
                        SaveWithDocument=True
                    )
                    inline_shape.Width = 110
                    inline_shape.Height = 50
                    modified = True
        word_doc.Save()
        word_doc.Close(SaveChanges=True)
        if modified:
            logger.info(f"Replaced '{find_str}' with image in shapes.")
        return temp_doc_path
    except Exception as e:
        logger.error(f"Error replacing text in shapes: {e}")
        try:
            word_doc.Close(SaveChanges=False)
        except:
            pass
        return temp_doc_path

def fill_template(doc, mapping, barcode_buffer=None):
    def replace_in_text(text):
        for k, v in mapping.items():
            if k == "«IMAGE_BARCODE»" and v:
                continue
            text = text.replace(k, str(v))
        return text
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = replace_in_text(run.text)
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = replace_in_text(run.text)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if "«IMAGE_BARCODE»" in para.text and barcode_buffer:
                                    para.clear()
                                    cell.paragraphs[0].paragraph_format.left_indent = Pt(-20)
                                    run = para.add_run()
                                    run.add_picture(barcode_buffer, width=Inches(3.0), height=Inches(0.35))
                                else:
                                    for run in para.runs:
                                        if run.text:
                                            run.text = replace_in_text(run.text)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = replace_in_text(run.text)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    if run.text:
                                        run.text = replace_in_text(run.text)
    return doc

def clear_placeholders(inner_table):
    def replace_in_text(text):
        return re.sub(r'«[^»]+»', "", text)
    for row in inner_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text:
                        run.text = replace_in_text(run.text)

def fill_inner_table(inner_table, mapping, qrcode_buffer=None):
    def replace_in_text(text):
        for k, v in mapping.items():
            if k == "«IMAGE_QRCODE»":
                continue
            text = text.replace(k, str(v))
        return text
    for row in inner_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if "«IMAGE_QRCODE»" in para.text and qrcode_buffer:
                    para.clear()
                    para.paragraph_format.left_indent = Pt(-5)
                    run = para.add_run()
                    run.add_picture(qrcode_buffer, width=Inches(1), height=Inches(1))
                else:
                    for run in para.runs:
                        if run.text:
                            run.text = replace_in_text(run.text)

def kill_libreoffice_processes():
    killed_count = 0
    try:
        for proc in psutil.process_iter(['pid', 'name']):
            if 'soffice' in proc.info['name'].lower():
                try:
                    proc.kill()
                    killed_count += 1
                except psutil.NoSuchProcess:
                    continue
        if killed_count > 0:
            logger.info(f"Killed {killed_count} existing LibreOffice processes")
        time.sleep(2)
    except Exception as e:
        logger.error(f"Error killing LibreOffice processes: {e}")

def convert_batch_with_retry(batch_files, output_dir, batch_id, timeout=180):
    max_retries = 3
    output_dir = Path(output_dir)
    for attempt in range(max_retries):
        try:
            with TemporaryDirectory() as temp_batch_dir:
                temp_output = Path(temp_batch_dir)
                logger.debug(f"Batch {batch_id} (Attempt {attempt + 1}): Converting {len(batch_files)} files...")
                cmd = [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    "--headless",
                    "--invisible",
                    "--nodefault",
                    "--nolockcheck",
                    "--nologo",
                    "--norestore",
                    "--convert-to", "pdf",
                    "--outdir", str(temp_output)
                ] + [str(Path(f)) for f in batch_files]
                process = subprocess.Popen(
                    cmd,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    creationflags=subprocess.CREATE_NO_WINDOW if os.name == 'nt' else 0,
                    text=True
                )
                try:
                    stdout, stderr = process.communicate(timeout=timeout)
                    if process.returncode != 0:
                        logger.error(f"Batch {batch_id} LibreOffice error (code {process.returncode}): {stderr}")
                        if attempt < max_retries - 1:
                            logger.info(f"Retrying batch {batch_id}...")
                            kill_libreoffice_processes()
                            time.sleep(2)
                            continue
                        return [], batch_files
                except subprocess.TimeoutExpired:
                    process.kill()
                    logger.error(f"Batch {batch_id} timeout after {timeout} seconds (attempt {attempt + 1})")
                    if attempt < max_retries - 1:
                        logger.info(f"Retrying batch {batch_id}...")
                        kill_libreoffice_processes()
                        time.sleep(2)
                        continue
                    return [], batch_files
                batch_pdfs = []
                failed_files = []
                for docx_path in batch_files:
                    docx_name = Path(docx_path).stem
                    temp_pdf = temp_output / f"{docx_name}.pdf"
                    final_pdf = output_dir / f"{docx_name}.pdf"
                    if temp_pdf.exists():
                        shutil.move(str(temp_pdf), str(final_pdf))
                        batch_pdfs.append(str(final_pdf))
                    else:
                        failed_files.append(docx_path)
                        logger.warning(f"Batch {batch_id}: Failed to convert {Path(docx_path).name}")
                success_rate = len(batch_pdfs) / len(batch_files) * 100 if batch_files else 0
                logger.info(f"Batch {batch_id} result: {len(batch_pdfs)}/{len(batch_files)} successful ({success_rate:.1f}%)")
                return batch_pdfs, failed_files
        except Exception as e:
            logger.error(f"Batch {batch_id} conversion error (attempt {attempt + 1}): {e}")
            if attempt < max_retries - 1:
                kill_libreoffice_processes()
                time.sleep(2)
                continue
            return [], batch_files
    return [], batch_files

def batch_convert_libreoffice(docx_files, output_dir, batch_size=100):
    if not docx_files:
        return []
    pdf_files = []
    total_failed = []
    output_dir = Path(output_dir)
    kill_libreoffice_processes()
    batches = [docx_files[i:i + batch_size] for i in range(0, len(docx_files), batch_size)]
    logger.info(f"Converting {len(docx_files)} DOCX files in {len(batches)} batches (size: {batch_size})...")
    start_time = time.time()
    for batch_id, batch in enumerate(batches, 1):
        batch_pdfs, batch_failed = convert_batch_with_retry(batch, output_dir, batch_id)
        pdf_files.extend(batch_pdfs)
        total_failed.extend(batch_failed)
    kill_libreoffice_processes()
    total_time = time.time() - start_time
    success_rate = len(pdf_files) / len(docx_files) * 100 if docx_files else 0
    logger.info(f"\n=== CONVERSION SUMMARY ===")
    logger.info(f"Total files: {len(docx_files)}")
    logger.info(f"Successful: {len(pdf_files)} ({success_rate:.1f}%)")
    logger.info(f"Failed: {len(total_failed)} ({100-success_rate:.1f}%)")
    logger.info(f"Time: {total_time:.1f}s | Rate: {len(pdf_files)/total_time:.1f} PDFs/sec" if total_time > 0 else "Time: 0s")
    return pdf_files

def get_raw_file(file: UploadFile):
    try:
        contents = file.file.read()
        df = pd.read_excel(BytesIO(contents), dtype=str)
        return df
    except Exception as e:
        logger.error(f"Error reading file: {e}")
        return pd.DataFrame([])

def fill_transmittal_template(template_path, group_df):
    try:
        temp_doc = Document(template_path)
        if not temp_doc.tables:
            logger.error("No tables found in the transmittal template.")
            return None
        total_groups = len(group_df) // 4 + (1 if len(group_df) % 4 else 0)
        temp_docs = []
        for group_idx in range(total_groups):
            temp_group_doc = Document(template_path)
            current_group_table = temp_group_doc.tables[0]
            page_rows = group_df.iloc[group_idx * 4:(group_idx + 1) * 4]
            for row_idx, row in enumerate(current_group_table.rows):
                if row_idx < len(page_rows):
                    row_data = page_rows.iloc[row_idx]
                    mapping = {f"«{col.upper()}»": str(row_data[col]) for col in group_df.columns if pd.notnull(row_data[col])}
                    qrcode_buffer = None
                    if dl_code := row_data.get('DL_CODE', ''):
                        qrcode_buffer = generate_qrcode(dl_code)
                    for cell in row.cells:
                        if cell.tables:
                            inner_table = cell.tables[0]
                            fill_inner_table(inner_table, mapping, qrcode_buffer)
                else:
                    for cell in row.cells:
                        if cell.tables:
                            inner_table = cell.tables[0]
                            clear_placeholders(inner_table)
            temp_docs.append(temp_group_doc)
        return temp_docs
    except Exception as e:
        logger.error(f"Failed to fill transmittal template: {e}")
        return None

# API Endpoints
@app.post("/api/login")
async def login(request: LoginRequest):
    if request.username in USERS and USERS[request.username]['password'] == request.password:
        return {"success": True, "role": USERS[request.username]['role'], "username": request.username}
    raise HTTPException(status_code=401, detail="Invalid credentials")

@app.get("/api/folders")
async def get_folders():
    if not all([FTP_CONFIG["hostname"], FTP_CONFIG["username"], FTP_CONFIG["password"]]):
        raise HTTPException(status_code=500, detail="FTP configuration incomplete")
    with FTPConnection(FTP_CONFIG["hostname"], FTP_CONFIG["port"], FTP_CONFIG["username"], FTP_CONFIG["password"]) as ftp_conn:
        ftp = ftp_conn.connect()
        if not ftp:
            raise HTTPException(status_code=500, detail="Failed to connect to FTP")
        folders = get_ftp_folders(ftp)
        return folders

@app.post("/api/dl_types")
async def get_dl_types(request: FolderRequest):
    if not request.folder:
        raise HTTPException(status_code=400, detail="Folder not specified")
    sheet_df = get_sheet_data(SERVICE_ACCOUNT_JSON, SPREADSHEET_ID, SHEET_NAME)
    if sheet_df.empty:
        raise HTTPException(status_code=500, detail="Failed to retrieve Google Sheets data")
    dl_types = sorted(sheet_df[sheet_df["CAMPAIGN"] == request.folder]["DL TYPE"].dropna().unique().tolist())
    return dl_types

@app.post("/api/templates")
async def get_templates(request: FolderRequest):
    if not request.folder:
        raise HTTPException(status_code=400, detail="Folder not specified")
    with FTPConnection(FTP_CONFIG["hostname"], FTP_CONFIG["port"], FTP_CONFIG["username"], FTP_CONFIG["password"]) as ftp_conn:
        ftp = ftp_conn.connect()
        if not ftp:
            raise HTTPException(status_code=500, detail="Failed to connect to FTP")
        templates = get_ftp_templates(ftp, request.folder)
        return templates

@app.post("/api/placeholders")
async def get_placeholders(request: TemplateRequest):
    if not all([request.folder, request.dl_type, request.template]):
        raise HTTPException(status_code=400, detail="Missing parameters")
    with FTPConnection(FTP_CONFIG["hostname"], FTP_CONFIG["port"], FTP_CONFIG["username"], FTP_CONFIG["password"]) as ftp_conn:
        ftp = ftp_conn.connect()
        if not ftp:
            raise HTTPException(status_code=500, detail="Failed to connect to FTP")
        signature_img_path = fetch_signature_from_ftp(ftp)
        if not signature_img_path:
            raise HTTPException(status_code=500, detail="Signature image not found")
        SESSION_STATE['files_to_cleanup'].append(signature_img_path)
        template_path = download_ftp_template(ftp, request.folder, request.template, is_header_footer=False)
        if not template_path:
            raise HTTPException(status_code=500, detail="Failed to download content template")
        SESSION_STATE['files_to_cleanup'].append(template_path)
        SESSION_STATE['template_path'] = template_path
        sheet_df = get_sheet_data(SERVICE_ACCOUNT_JSON, SPREADSHEET_ID, SHEET_NAME)
        matching_row = sheet_df[(sheet_df["CAMPAIGN"] == request.folder) & (sheet_df["DL TYPE"] == request.dl_type)]
        if matching_row.empty:
            raise HTTPException(status_code=404, detail=f"No header/footer template for {request.folder}/{request.dl_type}")
        header_footer_filename = matching_row["FILE"].iloc[0]
        if not header_footer_filename.lower().endswith('.docx'):
            header_footer_filename += '.docx'
        header_footer_template_path = download_ftp_template(ftp, None, header_footer_filename, is_header_footer=True)
        if not header_footer_template_path:
            raise HTTPException(status_code=500, detail=f"Failed to download header/footer template {header_footer_filename}")
        SESSION_STATE['files_to_cleanup'].append(header_footer_template_path)
        SESSION_STATE['header_footer_template_path'] = header_footer_template_path
        pythoncom.CoInitialize()
        word_app = win32com.client.DispatchEx("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = False
        try:
            base_template = combine_templates(header_footer_template_path, template_path, signature_img_path, word_app)
            if not base_template:
                raise HTTPException(status_code=500, detail="Failed to combine templates")
            SESSION_STATE['base_template'] = base_template
            placeholders = extract_placeholders(base_template)
            SESSION_STATE['placeholders'] = placeholders
            return placeholders
        finally:
            word_app.Quit()
            pythoncom.CoUninitialize()
            os.system("taskkill /IM WINWORD.EXE /F >nul 2>&1")

@app.post("/api/upload_excel")
async def upload_excel(file: UploadFile = File(...)):
    if not file.filename.endswith('.xlsx'):
        raise HTTPException(status_code=400, detail="Invalid file format. Please upload an .xlsx file")
    df = get_raw_file(file)
    if df.empty:
        raise HTTPException(status_code=500, detail="Failed to read Excel file")
    if 'FINAL_AREA' not in df.columns or 'DL_CODE' not in df.columns:
        raise HTTPException(status_code=400, detail="Excel file must contain FINAL_AREA and DL_CODE columns")
    return {"data": df.to_dict(orient='records')}


@app.post("/api/set_mode")
async def set_mode(request: ModeRequest):
    if request.mode not in ["DL Only", "DL w/ Transmittal", "Transmittal Only"]:
        raise HTTPException(status_code=400, detail="Invalid mode")
    
    SESSION_STATE['selected_mode'] = request.mode
    template_status = {}

    if request.mode in ["DL Only", "DL w/ Transmittal"]:
        if SESSION_STATE.get('base_template') and SESSION_STATE.get('header_footer_template_path') and SESSION_STATE.get('template_path'):
            template_status['dl_template'] = "DL and header/footer templates are ready"
        else:
            template_status['dl_template'] = "DL or header/footer template not loaded"

    if request.mode in ["DL w/ Transmittal", "Transmittal Only"]:
        try:
            with FTPConnection(FTP_CONFIG["hostname"], FTP_CONFIG["port"], FTP_CONFIG["username"], FTP_CONFIG["password"]) as ftp_conn:
                ftp = ftp_conn.connect()
                if not ftp:
                    raise HTTPException(status_code=500, detail="Failed to connect to FTP server")
                transmittal_template = "Transmittal QRCODE.docx"
                transmittal_template_path = download_ftp_template(ftp, None, transmittal_template, is_transmittal=True)
                if not transmittal_template_path:
                    raise HTTPException(status_code=500, detail=f"Failed to download transmittal template {transmittal_template}")
                SESSION_STATE['files_to_cleanup'].append(transmittal_template_path)
                SESSION_STATE['transmittal_template_path'] = transmittal_template_path
                template_status['transmittal_template'] = "Transmittal template is ready"
        except Exception as e:
            logger.error(f"Failed to set mode {request.mode}: {e}")
            template_status['transmittal_template'] = f"Failed to load transmittal template: {str(e)}"
            raise HTTPException(status_code=500, detail=f"Failed to set mode: {str(e)}")

    return {
        "success": True,
        "mode": request.mode,
        "template_status": template_status
    }

# ... (other imports and code remain unchanged until generate_pdfs_stream)

async def generate_pdfs_stream(uploaded_file: UploadFile, dataframe: pd.DataFrame):
    zipf = None
    try:
        if not SESSION_STATE.get('selected_mode'):
            yield json.dumps({'error': 'No processing mode selected'}) + '\n'
            return

        if not SESSION_STATE.get('base_template') and SESSION_STATE.get('selected_mode') in ["DL Only", "DL w/ Transmittal"]:
            yield json.dumps({'error': 'DL templates not loaded'}) + '\n'
            return
        if not SESSION_STATE.get('transmittal_template_path') and SESSION_STATE.get('selected_mode') in ["DL w/ Transmittal", "Transmittal Only"]:
            yield json.dumps({'error': 'Transmittal template not loaded'}) + '\n'
            return

        logger.debug(f"Excel file loaded with {len(dataframe)} rows")
        today_date = datetime.now().strftime("%B %d, %Y")
        dataframe['DL_ADDRESS'] = dataframe['DL_ADDRESS'].str.upper()
        valid_rows = dataframe[dataframe['LEADS_CHNAME'].notna()]
        total_records = len(valid_rows)
        if total_records == 0:
            yield json.dumps({'error': 'No valid rows found (LEADS_CHNAME missing)'}) + '\n'
            return

        logger.info(f"Processing {total_records} records across {len(valid_rows.groupby('FINAL_AREA'))} FINAL_AREA groups")
        with TemporaryDirectory() as temp_zip_dir:
            zip_path = Path(temp_zip_dir) / "final_area_pdfs.zip"
            zipf = zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED)
            try:
                with TemporaryDirectory() as temp_dir:
                    temp_base_path = Path(temp_dir) / "base_template.docx"
                    if SESSION_STATE.get('selected_mode') in ["DL Only", "DL w/ Transmittal"]:
                        SESSION_STATE['base_template'].save(temp_base_path)
                        SESSION_STATE['files_to_cleanup'].append(temp_base_path)
                    temp_transmittal_path = Path(temp_dir) / "transmittal_template.docx"
                    if SESSION_STATE.get('selected_mode') in ["DL w/ Transmittal", "Transmittal Only"]:
                        Document(SESSION_STATE['transmittal_template_path']).save(temp_transmittal_path)
                        SESSION_STATE['files_to_cleanup'].append(temp_transmittal_path)

                    processed_records = 0
                    start_time = time.time()
                    for final_area, group_df in valid_rows.groupby('FINAL_AREA'):
                        logger.debug(f"Processing FINAL_AREA: {final_area} ({len(group_df)} records)")
                        area_dir = OUTPUT_DIR / f"{final_area}"
                        os.makedirs(area_dir, exist_ok=True)
                        docx_files = []
                        temp_files = []
                        dl_pdf_merger = PdfMerger() if SESSION_STATE.get('selected_mode') in ["DL Only", "DL w/ Transmittal"] else None
                        transmittal_pdf_merger = PdfMerger() if SESSION_STATE.get('selected_mode') in ["DL w/ Transmittal", "Transmittal Only"] else None
                        record_count = len(group_df)
                        try:
                            if SESSION_STATE.get('selected_mode') in ["DL Only", "DL w/ Transmittal"]:
                                for idx, row in group_df.iterrows():
                                    barcode_buffer = None
                                    if barcode_value := row.get('DL_CODE', ''):
                                        barcode_buffer = generate_barcode(barcode_value)
                                    amount_words = amount_to_words(row.get('amount', '0.00'))
                                    mapping = {f"«{col.upper()}»": row[col] for col in dataframe.columns if pd.notnull(row[col])}
                                    mapping.update({
                                        "«IMAGE_BARCODE»": barcode_buffer or "",
                                        "«DL_DATE»": today_date,
                                        "«AMOUNT_ABBR»": amount_words,
                                        "«IMAGE_SIGNATURE»": SESSION_STATE.get('files_to_cleanup', [])[0] or ""
                                    })
                                    filled_doc = Document(temp_base_path)
                                    filled_doc = fill_template(filled_doc, mapping, barcode_buffer)
                                    if filled_doc:
                                        unique_name = f"dl_{final_area}_{idx}_{uuid.uuid4().hex[:6]}"
                                        docx_output = area_dir / f"{unique_name}.docx"
                                        filled_doc.save(docx_output)
                                        docx_files.append(str(docx_output))
                                        temp_files.append(str(docx_output))
                                    processed_records += 1
                                    progress = (processed_records / total_records) * 100
                                    yield json.dumps({
                                        'progress': progress,
                                        'message': f"Processing record {processed_records}/{total_records} (FINAL_AREA: {final_area})"
                                    }) + '\n'
                                    await asyncio.sleep(0)  # Allow event loop to process other tasks

                            if SESSION_STATE.get('selected_mode') in ["DL w/ Transmittal", "Transmittal Only"]:
                                transmittal_docs = fill_transmittal_template(temp_transmittal_path, group_df)
                                if transmittal_docs:
                                    for doc_idx, transmittal_doc in enumerate(transmittal_docs):
                                        unique_name = f"transmittal_{final_area}_{doc_idx}_{uuid.uuid4().hex[:6]}"
                                        transmittal_docx_output = area_dir / f"{unique_name}.docx"
                                        transmittal_doc.save(transmittal_docx_output)
                                        docx_files.append(str(transmittal_docx_output))
                                        temp_files.append(str(transmittal_docx_output))

                            if docx_files:
                                logger.debug(f"Converting {len(docx_files)} DOCX files for {final_area}")
                                yield json.dumps({
                                    'progress': (processed_records / total_records) * 100,
                                    'message': f"Converting {len(docx_files)} documents to PDF for {final_area}"
                                }) + '\n'
                                await asyncio.sleep(0)
                                pdf_files = batch_convert_libreoffice(docx_files, area_dir)
                                yield json.dumps({
                                    'progress': (processed_records / total_records) * 100,
                                    'message': f"Converted {len(pdf_files)}/{len(docx_files)} PDFs for {final_area}"
                                }) + '\n'
                                await asyncio.sleep(0)
                                for pdf_file in pdf_files:
                                    pdf_name = Path(pdf_file).name
                                    if "transmittal_" in pdf_name and SESSION_STATE.get('selected_mode') in ["DL w/ Transmittal", "Transmittal Only"]:
                                        transmittal_pdf_merger.append(pdf_file)
                                        temp_files.append(pdf_file)
                                    elif "dl_" in pdf_name and SESSION_STATE.get('selected_mode') in ["DL Only", "DL w/ Transmittal"]:
                                        dl_pdf_merger.append(pdf_file)
                                        temp_files.append(pdf_file)
                            if SESSION_STATE.get('selected_mode') in ["DL Only", "DL w/ Transmittal"] and dl_pdf_merger:
                                dl_merged_pdf_path = area_dir / f"{final_area}_dl_{record_count}.pdf"
                                with open(dl_merged_pdf_path, "wb") as f:
                                    dl_pdf_merger.write(f)
                                if dl_merged_pdf_path.exists():
                                    zipf.write(dl_merged_pdf_path, f"{final_area}/{final_area}_dl_{record_count}.pdf")
                                    temp_files.append(str(dl_merged_pdf_path))
                            if SESSION_STATE.get('selected_mode') in ["DL w/ Transmittal", "Transmittal Only"] and transmittal_pdf_merger:
                                transmittal_merged_pdf_path = area_dir / f"{final_area}_transmittal_{record_count}.pdf"
                                with open(transmittal_merged_pdf_path, "wb") as f:
                                    transmittal_pdf_merger.write(f)
                                if transmittal_merged_pdf_path.exists():
                                    zipf.write(transmittal_merged_pdf_path, f"{final_area}/{final_area}_transmittal_{record_count}.pdf")
                                    temp_files.append(str(transmittal_merged_pdf_path))
                        except Exception as e:
                            logger.error(f"Error processing FINAL_AREA {final_area}: {e}")
                            yield json.dumps({'error': f"Failed to process {final_area}: {str(e)}"}) + '\n'
                            return
                        finally:
                            if dl_pdf_merger:
                                dl_pdf_merger.close()
                            if transmittal_pdf_merger:
                                transmittal_pdf_merger.close()
                            cleanup_files(temp_files)
            finally:
                if zipf:
                    zipf.close()
                    logger.debug(f"Closed ZipFile: {zip_path}")

            if zip_path.exists() and zip_path.stat().st_size > 0:
                persistent_zip_path = OUTPUT_DIR / "final_area_pdfs.zip"
                shutil.move(str(zip_path), str(persistent_zip_path))
                SESSION_STATE['download_completed'] = True
                SESSION_STATE['zip_path'] = str(persistent_zip_path)
                total_time = time.time() - start_time
                logger.info(f"Completed processing in {total_time:.1f}s")
                yield json.dumps({
                    'progress': 100,
                    'message': f"Completed {total_records} records in {total_time:.1f}s!",
                    'download_ready': True
                }) + '\n'
            else:
                logger.error("No valid PDFs generated")
                yield json.dumps({'error': 'No valid PDFs generated'}) + '\n'
    except Exception as e:
        logger.error(f"Processing failed in generate: {str(e)}")
        yield json.dumps({'error': f"Processing failed: {str(e)}"}) + '\n'
    finally:
        if zipf and not zipf.fp is None:
            zipf.close()
            logger.debug("Ensured ZipFile is closed in finally block")

# ... (rest of the code remains unchanged, including /api/set_mode and other endpoints)

@app.post("/api/generate_pdfs")
async def generate_pdfs(file: UploadFile = File(...)):
    if not file.filename.endswith('.xlsx'):
        raise HTTPException(status_code=400, detail="Invalid file format. Please upload an .xlsx file")
    df = get_raw_file(file)
    if df.empty:
        raise HTTPException(status_code=500, detail="Failed to read Excel file")
    if 'FINAL_AREA' not in df.columns or 'DL_CODE' not in df.columns:
        raise HTTPException(status_code=400, detail="Excel file must contain FINAL_AREA and DL_CODE columns")
    return StreamingResponse(generate_pdfs_stream(file, df), media_type="text/event-stream")

@app.get("/api/download_zip")
async def download_zip():
    if not SESSION_STATE.get('download_completed') or not SESSION_STATE.get('zip_path'):
        raise HTTPException(status_code=404, detail="No ZIP file available")
    if not Path(SESSION_STATE['zip_path']).exists():
        raise HTTPException(status_code=404, detail="ZIP file not found on server")
    return FileResponse(
        SESSION_STATE['zip_path'],
        filename="final_area_pdfs.zip",
        media_type="application/zip"
    )

@app.post("/api/cleanup")
async def cleanup():
    cleanup_files(SESSION_STATE.get('files_to_cleanup', []))
    if SESSION_STATE.get('zip_path') and Path(SESSION_STATE['zip_path']).exists():
        Path(SESSION_STATE['zip_path']).unlink()
        logger.info(f"Deleted ZIP file: {SESSION_STATE['zip_path']}")
    shutil.rmtree(OUTPUT_DIR, ignore_errors=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    SESSION_STATE['download_completed'] = False
    SESSION_STATE['files_to_cleanup'] = []
    SESSION_STATE['base_template'] = None
    SESSION_STATE['template_path'] = None
    SESSION_STATE['header_footer_template_path'] = None
    SESSION_STATE['transmittal_template_path'] = None
    SESSION_STATE['placeholders'] = []
    SESSION_STATE['zip_path'] = None
    return {"success": True, "message": "Files cleaned up successfully"}

@app.post("/api/set_mode")
async def set_mode(request: ModeRequest):
    if request.mode not in ["DL Only", "DL w/ Transmittal", "Transmittal Only"]:
        raise HTTPException(status_code=400, detail="Invalid mode")
    
    SESSION_STATE['selected_mode'] = request.mode
    template_status = {}

    if request.mode in ["DL Only", "DL w/ Transmittal"]:
        if SESSION_STATE.get('base_template') and SESSION_STATE.get('header_footer_template_path') and SESSION_STATE.get('template_path'):
            template_status['dl_template'] = "DL and header/footer templates are ready"
        else:
            template_status['dl_template'] = "DL or header/footer template not loaded"

    if request.mode in ["DL w/ Transmittal", "Transmittal Only"]:
        try:
            with FTPConnection(FTP_CONFIG["hostname"], FTP_CONFIG["port"], FTP_CONFIG["username"], FTP_CONFIG["password"]) as ftp_conn:
                ftp = ftp_conn.connect()
                if not ftp:
                    raise HTTPException(status_code=500, detail="Failed to connect to FTP server")
                transmittal_template = "Transmittal QRCODE.docx"
                transmittal_template_path = download_ftp_template(ftp, None, transmittal_template, is_transmittal=True)
                if not transmittal_template_path:
                    raise HTTPException(status_code=500, detail=f"Failed to download transmittal template {transmittal_template}")
                SESSION_STATE['files_to_cleanup'].append(transmittal_template_path)
                SESSION_STATE['transmittal_template_path'] = transmittal_template_path
                template_status['transmittal_template'] = "Transmittal template is ready"
        except Exception as e:
            logger.error(f"Failed to set mode {request.mode}: {e}")
            template_status['transmittal_template'] = f"Failed to load transmittal template: {str(e)}"
            raise HTTPException(status_code=500, detail=f"Failed to set mode: {str(e)}")

    return {
        "success": True,
        "mode": request.mode,
        "template_status": template_status
    }
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)