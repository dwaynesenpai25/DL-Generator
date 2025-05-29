import streamlit as st
import pandas as pd
from docx import Document
import re
import os
import uuid
from tempfile import NamedTemporaryFile
import subprocess
from docx.shared import Inches, Pt
from io import BytesIO
from PyPDF2 import PdfMerger
from PyPDF2 import PdfReader
from barcode import Code128
from barcode.writer import ImageWriter
import shutil
import logging
from ftplib import FTP
from datetime import datetime
from dotenv import load_dotenv
import gspread
from google.oauth2.service_account import Credentials
from pathlib import Path
import win32com.client
import pythoncom
import sys
import tempfile
import zipfile

# Load the .env file from the config folder
env_path = os.path.join(os.path.dirname(__file__), "config", ".env")
load_dotenv(env_path)
logger = logging.getLogger(__name__)
logger.info(f"Loaded .env file from: {env_path}")

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Create output folder
OUTPUT_DIR = os.path.abspath("output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Create barcode images folder
BARCODE_DIR = os.path.abspath("barcode_images")
os.makedirs(BARCODE_DIR, exist_ok=True)

# FTP connection details from .env
FTP_CONFIG = {
    "hostname": os.getenv("OMKT_FTP_HOSTNAME"),
    "port": int(os.getenv("OMKT_FTP_PORT", 21)),
    "username": os.getenv("OMKT_FTP_USERNAME"),
    "password": os.getenv("OMKT_FTP_PASSWORD")
}

# Validate FTP configuration
if not all([FTP_CONFIG["hostname"], FTP_CONFIG["username"], FTP_CONFIG["password"]]):
    st.error("FTP configuration is incomplete. Check your .env file in the config folder.")
    st.stop()

# Helper function to connect to Google Sheets and get CAMPAIGN, DL TYPE, and FILE data
@st.cache_data
def get_sheet_data(service_account_json_path, spreadsheet_id, sheet_name="LetterHeads"):
    try:
        credentials = Credentials.from_service_account_file(
            service_account_json_path,
            scopes=[
                "https://www.googleapis.com/auth/spreadsheets",
                "https://www.googleapis.com/auth/drive"
            ]
        )
        client = gspread.authorize(credentials)
        logger.info("Authenticated with Google Sheets API")

        spreadsheet = client.open_by_key(spreadsheet_id)
        worksheet = spreadsheet.worksheet(sheet_name)
        
        sheet_data = worksheet.get_all_records()
        df = pd.DataFrame(sheet_data)
        
        required_columns = ["CAMPAIGN", "DL TYPE", "FILE"]
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            st.error(f"Missing columns in spreadsheet: {missing_columns}")
            logger.error(f"Missing columns in spreadsheet: {missing_columns}")
            return pd.DataFrame()

        df = df[["CAMPAIGN", "DL TYPE", "FILE"]].dropna()
        if df.empty:
            st.warning("No valid CAMPAIGN, DL TYPE, or FILE values found in the spreadsheet.")
            logger.warning("No valid CAMPAIGN, DL TYPE, or FILE values found in the spreadsheet.")
            return pd.DataFrame()

        logger.info(f"Retrieved sheet data with {len(df)} valid rows")
        return df
    except Exception as e:
        st.error(f"Failed to access Google Sheets: {e}")
        logger.error(f"Failed to access Google Sheets: {e}")
        return pd.DataFrame()

# Helper function to connect to FTP and retrieve folder names
def get_ftp_folders():
    try:
        ftp = FTP()
        ftp.connect(FTP_CONFIG["hostname"], FTP_CONFIG["port"])
        ftp.login(FTP_CONFIG["username"], FTP_CONFIG["password"])
        logger.info("Connected to FTP server for folder retrieval")

        ftp_path = "/DL AUTOMATION/Template DL/Content"
        ftp.cwd(ftp_path)
        logger.info(f"Navigated to FTP path: {ftp_path}")

        items = ftp.nlst()
        folders = []
        current_dir = ftp.pwd()
        for item in items:
            try:
                ftp.cwd(item)
                folders.append(item)
                ftp.cwd(current_dir)
            except Exception:
                continue

        ftp.quit()
        return sorted(folders)
    except Exception as e:
        logger.error(f"Failed to retrieve folders from FTP: {e}")
        st.error(f"Failed to retrieve folders from FTP: {e}")
        return []

# Helper function to list .docx files in the selected FTP folder
def get_ftp_templates(folder_name):
    try:
        ftp = FTP()
        ftp.connect(FTP_CONFIG["hostname"], FTP_CONFIG["port"])
        ftp.login(FTP_CONFIG["username"], FTP_CONFIG["password"])
        logger.info("Connected to FTP server for template retrieval")

        ftp_path = f"/DL AUTOMATION/Template DL/Content/{folder_name}"
        ftp.cwd(ftp_path)
        logger.info(f"Navigated to FTP path: {ftp_path}")

        templates = [item for item in ftp.nlst() if item.lower().endswith('.docx')]
        ftp.quit()
        return sorted(templates)
    except Exception as e:
        logger.error(f"Failed to retrieve templates from FTP folder {folder_name}: {e}")
        st.error(f"Failed to retrieve templates from FTP folder {folder_name}: {e}")
        return []

# Helper function to download a template from FTP
def download_ftp_template(folder_name, template_name, is_header_footer=False):
    try:
        ftp = FTP()
        ftp.connect(FTP_CONFIG["hostname"], FTP_CONFIG["port"])
        ftp.login(FTP_CONFIG["username"], FTP_CONFIG["password"])
        logger.info("Connected to FTP server for template download")

        if is_header_footer:
            ftp_path = "/DL AUTOMATION/Template DL/Letter Head"
        else:
            ftp_path = f"/DL AUTOMATION/Template DL/Content/{folder_name}"
        ftp.cwd(ftp_path)
        logger.info(f"Navigated to FTP path: {ftp_path}")

        with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            ftp.retrbinary(f"RETR {template_name}", tmp.write)
            tmp_path = tmp.name

        if os.path.exists(tmp_path):
            logger.info(f"Downloaded template {template_name} to: {tmp_path}")
            return tmp_path
        else:
            raise Exception(f"Failed to download template {template_name}")
    except Exception as e:
        logger.error(f"Failed to download template {template_name}: {e}")
        st.error(f"Failed to download template {template_name}: {e}")
        return None

# Helper function to combine content template into header/footer template
def combine_templates(header_footer_path, content_path):
    try:
        header_footer_doc = Document(header_footer_path)
        content_doc = Document(content_path)
        logger.info(f"Loaded header/footer template: {header_footer_path}")
        logger.info(f"Loaded content template: {content_path}")

        while header_footer_doc.paragraphs:
            header_footer_doc.paragraphs[0]._element.getparent().remove(header_footer_doc.paragraphs[0]._element)

        for elem in content_doc.element.body:
            header_footer_doc.element.body.append(elem)

        for para in header_footer_doc.paragraphs:
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)

        logger.info("Combined content into header/footer template")
        return header_footer_doc
    except Exception as e:
        logger.error(f"Failed to combine templates: {e}")
        st.error(f"Failed to combine templates: {e}")
        return None

# Helper function to convert numbers to words
def number_to_words(number):
    if number == 0:
        return "ZERO"

    units = ["", "ONE", "TWO", "THREE", "FOUR", "FIVE", "SIX", "SEVEN", "EIGHT", "NINE"]
    teens = ["", "ELEVEN", "TWELVE", "THIRTEEN", "FOURTEEN", "FIFTEEN", "SIXTEEN","SEVENTEEN", "EIGHTEEN", "NINETEEN"]
    tens = ["", "TEN", "TWENTY", "THIRTY", "FORTY", "FIFTY", "SIXTY", "SEVENTY", "EIGHTY", "NINETY"]
    thousands = ["", "THOUSAND", "MILLION"]

    def convert_hundreds(num):
        if num == 0:
            return ""
        result = []
        if num >= 100:
            result.append(units[num // 100] + " HUNDRED")
            num %= 100
            if num > 0:
                result.append("AND")
        if num >= 20:
            result.append(tens[num // 10])
            num %= 10
            if num > 0:
                result.append(units[num])
        elif num >= 11:
            result.append(teens[num - 10])
        elif num == 10:
            result.append(tens[1])
        elif num > 0:
            result.append(units[num])
        return " ".join(result)

    result = []
    i = 0
    while number > 0:
        chunk = number % 1000
        if chunk > 0:
            chunk_words = convert_hundreds(chunk)
            if thousands[i]:
                chunk_words += f" {thousands[i]}"
            result.insert(0, chunk_words)
        number //= 1000
        i += 1

    return ", ".join(result).replace(", AND", " AND").strip()

# Convert amount to word form
def amount_to_words(amount_str):
    try:
        amount = float(amount_str.replace(",", ""))
        pesos = int(amount)
        cents = int(round((amount - pesos) * 100))

        pesos_words = number_to_words(pesos)
        if pesos == 0:
            pesos_words = "ZERO"

        cents_words = number_to_words(cents) if cents > 0 else "ZERO"

        result = f"{pesos_words} PESOS"
        if cents > 0:
            result += f", AND {cents_words} CENTS"
        else:
            result += ", AND ZERO CENTS"

        return result
    except Exception as e:
        logger.error(f"Failed to convert amount to words: {amount_str}, error: {e}")
        st.error(f"Failed to convert amount to words: {amount_str}")
        return "ERROR CONVERTING AMOUNT"

# Generate barcode image
def generate_barcode(barcode_value):
    try:
        barcode = Code128(barcode_value, writer=ImageWriter())
        buffer = BytesIO()
        barcode.write(buffer, options={"write_text": False, "module_width": 1, "module_height": 8, "quiet_zone": 2.0})
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Failed to generate barcode for {barcode_value}: {e}")
        return None

# Fetch signature image from FTP
def fetch_signature_from_ftp():
    try:
        ftp = FTP()
        ftp.connect(FTP_CONFIG["hostname"], FTP_CONFIG["port"])
        ftp.login(FTP_CONFIG["username"], FTP_CONFIG["password"])
        ftp_path = f"field/DL/ATTY SIGNATURE/{datetime.now().strftime('%m-%d-%Y')}"
        ftp.cwd(ftp_path)

        with NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            ftp.retrbinary("RETR attySignature.PNG", tmp.write)
            tmp_path = tmp.name

        ftp.quit()
        return tmp_path if os.path.exists(tmp_path) else None
    except Exception as e:
        st.error(f"Failed to fetch signature from FTP: {e}")
        return None

# Extract placeholders from .docx
def extract_placeholders(doc):
    placeholders = set()

    for para in doc.paragraphs:
        if "«" in para.text:
            matches = re.findall(r"«(.*?)»", para.text)
            placeholders.update(["«" + m.strip() + "»" for m in matches])

    for node in doc._element.iter():
        if node.tag.endswith("}t") and node.text and "«" in node.text:
            matches = re.findall(r"«(.*?)»", node.text)
            placeholders.update(["«" + m.strip() + "»" for m in matches])

    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header is not None:
                for para in header.paragraphs:
                    if "«" in para.text:
                        matches = re.findall(r"«(.*?)»", para.text)
                        placeholders.update(["«" + m.strip() + "»" for m in matches])
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if "«" in para.text:
                                    matches = re.findall(r"«(.*?)»", para.text)
                                    placeholders.update(["«" + m.strip() + "»" for m in matches])
                for node in header._element.iter():
                    if node.tag.endswith("}t") and node.text and "«" in node.text:
                        matches = re.findall(r"«(.*?)»", node.text)
                        placeholders.update(["«" + m.strip() + "»" for m in matches])

        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer is not None:
                for para in footer.paragraphs:
                    if "«" in para.text:
                        matches = re.findall(r"«(.*?)»", para.text)
                        placeholders.update(["«" + m.strip() + "»" for m in matches])
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if "«" in para.text:
                                    matches = re.findall(r"«(.*?)»", para.text)
                                    placeholders.update(["«" + m.strip() + "»" for m in matches])
                for node in footer._element.iter():
                    if node.tag.endswith("}t") and node.text and "«" in node.text:
                        matches = re.findall(r"«(.*?)»", node.text)
                        placeholders.update(["«" + m.strip() + "»" for m in matches])

    return sorted(placeholders)

# Replace placeholders in text boxes with win32com
def replace_in_text_boxes(doc, find_str, replace_with_image_path):
    pythoncom.CoInitialize()
    word_app = None
    try:
        word_app = win32com.client.DispatchEx("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = False

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_doc_path = Path(temp_dir) / "temp_doc.docx"
            doc.save(temp_doc_path)
            word_doc = word_app.Documents.Open(str(temp_doc_path.absolute()))

            for shape in word_doc.Shapes:
                if shape.TextFrame.HasText:
                    text_range = shape.TextFrame.TextRange
                    find = text_range.Find
                    find.Text = find_str
                    find.Forward = True
                    find.Wrap = 1
                    while find.Execute():
                        found_range = text_range.Duplicate
                        found_range.Find.Execute(FindText=find_str)
                        found_range.Text = ""
                        inline_shape = found_range.InlineShapes.AddPicture(
                            FileName=str(Path(replace_with_image_path).absolute()),
                            LinkToFile=False,
                            SaveWithDocument=True
                        )
                        inline_shape.Width = 110
                        inline_shape.Height = 50

            word_doc.Save()
            word_doc.Close(SaveChanges=False)
            return Document(temp_doc_path)
    except Exception as e:
        st.error(f"Error replacing text in text boxes: {e}")
        return None
    finally:
        if word_app:
            word_app.Quit()
        pythoncom.CoUninitialize()

def fill_template(doc, mapping, barcode_buffer=None, signature_img_path=None):
    def replace_in_text(text):
        for k, v in mapping.items():
            if k in ["«IMAGE_SIGNATURE»", "«IMAGE_BARCODE»"] and v:
                continue
            text = text.replace(k, str(v))
        return text

    # Process document body
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = replace_in_text(run.text)

    # Process headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = replace_in_text(run.text)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if "«IMAGE_BARCODE»" in para.text and barcode_buffer:
                                    para.clear()
                                    cell.paragraphs[0].paragraph_format.left_indent = Pt(-20)
                                    run = para.add_run()
                                    run.add_picture(barcode_buffer, width=Inches(3.0), height=Inches(0.35))
                                else:
                                    for run in para.runs:
                                        if run.text:
                                            run.text = replace_in_text(run.text)

        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = replace_in_text(run.text)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    if run.text:
                                        run.text = replace_in_text(run.text)

    # Handle signature in text boxes
    if signature_img_path and os.path.exists(signature_img_path) and "«IMAGE_SIGNATURE»" in mapping:
        doc = replace_in_text_boxes(doc, "«IMAGE_SIGNATURE»", signature_img_path)

    return doc

# Batch convert DOCX files to PDF
def batch_convert_to_pdf(docx_files, output_dir):
    pythoncom.CoInitialize()
    word_app = None
    pdf_files = []
    try:
        word_app = win32com.client.DispatchEx("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = False

        for docx_path in docx_files:
            output_path = os.path.join(output_dir, f"{os.path.splitext(os.path.basename(docx_path))[0]}.pdf")
            if not os.path.exists(docx_path):
                st.warning(f"Document not found: {docx_path}")
                continue

            doc = word_app.Documents.Open(os.path.abspath(docx_path))
            doc.SaveAs(os.path.abspath(output_path), FileFormat=17)
            doc.Close(SaveChanges=False)
            if os.path.exists(output_path):
                pdf_files.append(output_path)
            else:
                st.warning(f"PDF not generated for {docx_path}")

        return pdf_files
    except Exception as e:
        st.error(f"Failed to convert DOCX files to PDF: {e}")
        return []
    finally:
        if word_app:
            word_app.Quit()
        pythoncom.CoUninitialize()

# Read Excel file
def get_raw_file(file, sheet_name=None, engine=None):
    """Read an Excel file into a pandas DataFrame."""
    try:
        if sheet_name is None:
            return pd.read_excel(file, dtype=str, engine=engine)
        else:
            return pd.read_excel(file, sheet_name=sheet_name, engine=engine, dtype=str)
    except Exception as e:
        st.error(f"Error reading file: {e}")
        logger.error(f"Error reading file: {e}")
        return pd.DataFrame([])

# Streamlit UI
st.title("📄 DL GENERATOR")

# Google Sheets configuration
SERVICE_ACCOUNT_JSON = "config/dl_automation_sheet.json"
SPREADSHEET_ID = "1M0Vmmf9HfPRB0oSeJR_xUAZPPpTJ4xsR3gYDQMvnu5k"
SHEET_NAME = "LetterHeads"

# Fetch data from Google Sheets
sheet_df = get_sheet_data(SERVICE_ACCOUNT_JSON, SPREADSHEET_ID, SHEET_NAME)
if sheet_df.empty:
    st.error("Failed to retrieve data from Google Sheets.")
    st.stop()

# Fetch signature image
signature_img_path = fetch_signature_from_ftp()
if not signature_img_path:
    st.error("Signature image not found. Check FTP connection.")
    st.stop()

# Fetch folder names from FTP
st.subheader("Select Template Folder")
folders = get_ftp_folders()
if not folders:
    st.error("No folders found in FTP directory.")
    st.stop()

# Pre-select folder
campaign_values = sheet_df["CAMPAIGN"].dropna().unique().tolist()
default_folder = campaign_values[0] if campaign_values else None
default_index = folders.index(default_folder) + 1 if default_folder in folders else 0
if default_folder and default_index == 0:
    st.warning(f"CAMPAIGN folder '{default_folder}' not found in FTP.")

selected_folder = st.selectbox("Choose a folder:", [""] + folders, index=default_index)

# Select DL TYPE
selected_dl_type = None
if selected_folder:
    dl_types = sorted(sheet_df[sheet_df["CAMPAIGN"] == selected_folder]["DL TYPE"].dropna().unique().tolist())
    if not dl_types:
        st.warning(f"No DL TYPE values found for CAMPAIGN '{selected_folder}'.")
    else:
        st.subheader("Select DL TYPE")
        selected_dl_type = st.selectbox("Choose a DL TYPE:", [""] + dl_types)

# Template selection
template_path = None
header_footer_template_path = None
base_template = None
if selected_folder and selected_dl_type:
    templates = get_ftp_templates(selected_folder)
    if not templates:
        st.error(f"No .docx templates found in folder '{selected_folder}'.")
    else:
        st.subheader("Select Content File")
        selected_template = st.selectbox("Choose a template:", [""] + templates)
        if selected_template:
            template_path = download_ftp_template(selected_folder, selected_template, is_header_footer=False)
            if not template_path:
                st.error("Failed to download content template.")
                st.stop()

            matching_row = sheet_df[(sheet_df["CAMPAIGN"] == selected_folder) & (sheet_df["DL TYPE"] == selected_dl_type)]
            if matching_row.empty:
                st.error(f"No header/footer template found for CAMPAIGN '{selected_folder}' and DL TYPE '{selected_dl_type}'.")
                st.stop()

            header_footer_filename = matching_row["FILE"].iloc[0]
            if not header_footer_filename.lower().endswith('.docx'):
                header_footer_filename += '.docx'

            header_footer_template_path = download_ftp_template(None, header_footer_filename, is_header_footer=True)
            if not header_footer_template_path:
                st.error(f"Failed to download header/footer template '{header_footer_filename}'.")
                st.stop()

            # Combine templates once
            base_template = combine_templates(header_footer_template_path, template_path)
            if not base_template:
                st.error("Failed to combine templates.")
                st.stop()

            st.success(f"Loaded content template '{selected_template}' and header/footer template '{header_footer_filename}'.")
            placeholders = extract_placeholders(base_template)
            st.write("🔍 Detected placeholders:", placeholders)

# Process Excel file and generate PDFs
if template_path and header_footer_template_path and base_template:
    uploaded_excel = st.file_uploader("Upload Excel file with matching column headers", type=["xlsx"])
    if uploaded_excel:
        df = get_raw_file(uploaded_excel)
        st.write("📋 Data preview:")
        st.write(df)

        if 'FINAL_AREA' not in df.columns:
            st.error("Excel file must contain a 'FINAL_AREA' column.")
            st.stop()

        if st.button("🔄 Generate PDFs by FINAL_AREA"):
            st.info("Processing...")
            today_date = datetime.now().strftime("%B %d, %Y")
            valid_rows = df[df['LEADS_CHNAME'].notna()]
            total_rows = len(valid_rows)
            if total_rows == 0:
                st.error("No valid rows found (LEADS_CHNAME missing).")
                st.stop()

            progress_bar = st.progress(0)
            progress_text = st.empty()
            row_counter = 0
            pdf_paths = []

            # Process by FINAL_AREA
            for final_area, group_df in valid_rows.groupby('FINAL_AREA'):
                pdf_merger = PdfMerger()
                docx_files = []

                for _, row in group_df.iterrows():
                    row_counter += 1
                    progress_bar.progress(row_counter / total_rows)
                    progress_text.text(f"Processing row {row_counter} of {total_rows} (FINAL_AREA: {final_area})")

                    # Generate barcode in memory
                    barcode_buffer = None
                    if barcode_value := row.get('DL_CODE', ''):
                        barcode_buffer = generate_barcode(barcode_value)

                    # Prepare data mapping
                    amount_words = amount_to_words(row.get('amount', '0.00'))
                    mapping = {f"«{col.upper()}»": row[col] for col in df.columns if pd.notnull(row[col])}
                    mapping.update({
                        "«IMAGE_BARCODE»": barcode_buffer or "",
                        "«DL_DATE»": today_date,
                        "«AMOUNT_ABBR»": amount_words,
                        "«IMAGE_SIGNATURE»": signature_img_path or ""
                    })

                    # Create a copy of the base template
                    with tempfile.TemporaryDirectory() as temp_dir:
                        temp_doc_path = Path(temp_dir) / "temp_base.docx"
                        base_template.save(temp_doc_path)
                        filled_doc = Document(temp_doc_path)

                    # Fill template
                    filled_doc = fill_template(filled_doc, mapping, barcode_buffer, signature_img_path)
                    if not filled_doc:
                        st.warning(f"Failed to fill template for row {row_counter} in FINAL_AREA '{final_area}'.")
                        continue

                    # Save DOCX
                    unique_name = f"document_{final_area}_{row_counter}_{uuid.uuid4().hex[:8]}"
                    docx_output = os.path.join(OUTPUT_DIR, f"{unique_name}.docx")
                    filled_doc.save(docx_output)
                    docx_files.append(docx_output)

                # Batch convert DOCX to PDF
                pdf_files = batch_convert_to_pdf(docx_files, OUTPUT_DIR)
                for pdf_file in pdf_files:
                    pdf_merger.append(pdf_file)

                # Save merged PDF
                merged_pdf_path = os.path.join(OUTPUT_DIR, f"{final_area}.pdf")
                with open(merged_pdf_path, "wb") as f:
                    pdf_merger.write(f)

                if os.path.exists(merged_pdf_path):
                    with open(merged_pdf_path, "rb") as f:
                        pdf_reader = PdfReader(f)
                        st.write(f"PDF for FINAL_AREA '{final_area}' generated with {len(pdf_reader.pages)} pages")
                    pdf_paths.append(merged_pdf_path)
                else:
                    st.error(f"Merged PDF not generated for FINAL_AREA '{final_area}'.")

                # Clean up individual DOCX and PDF files
                for file_path in docx_files + pdf_files:
                    try:
                        os.remove(file_path)
                    except:
                        pass

            # Create and offer ZIP file
            zip_path = os.path.join(OUTPUT_DIR, "final_area_pdfs.zip")
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for pdf_path in pdf_paths:
                    zipf.write(pdf_path, os.path.basename(pdf_path))

            if os.path.exists(zip_path):
                with open(zip_path, "rb") as f:
                    st.download_button(
                        label="⬇️ Download ZIP with all FINAL_AREA PDFs",
                        data=f.read(),
                        file_name="final_area_pdfs.zip",
                        mime="application/zip"
                    )
            else:
                st.error("Failed to create ZIP file.")

            # Clean up all temporary files
            for file_path in [signature_img_path, template_path, header_footer_template_path, zip_path] + pdf_paths:
                try:
                    if file_path and os.path.exists(file_path):
                        os.remove(file_path)
                except:
                    pass

            progress_bar.empty()
            progress_text.empty()
            st.success("✅ ZIP file with PDFs generated for each FINAL_AREA!")