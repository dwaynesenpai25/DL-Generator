
import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
from copy import deepcopy
from PyPDF2 import PdfMerger
import os
import subprocess
import stat
import zipfile

# Title of the app
st.title("Template Placeholder Replacer")

# Upload Excel file
uploaded_excel = st.file_uploader("Upload your Excel file", type=["xlsx", "xls"])

# Upload Word template
uploaded_template = st.file_uploader("Upload your Word template (.docx)", type=["docx"])

# Output folder
OUTPUT_FOLDER = r"C:\Users\SPM\Desktop\ONLY SAVE FILE HERE\MAIL\output"
# Path to LibreOffice soffice.exe
LIBREOFFICE_PATH = r"C:\Program Files\LibreOffice\program\soffice.exe"

def clear_placeholders(inner_table):
    """Clear all placeholders in an inner table by replacing them with empty strings."""
    def replace_in_text(text):
        # Replace any placeholder (text between « and ») with empty string
        import re
        return re.sub(r'«[^»]+»', "", text)

    for row in inner_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text:
                        run.text = replace_in_text(run.text)

def fill_inner_table(inner_table, mapping):
    """Fill inner table with data, replacing placeholders."""
    def replace_in_text(text):
        for k, v in mapping.items():
            if k == "«IMAGE_BARCODE»" and v:  # Skip barcode if present
                continue
            text = text.replace(k, str(v))
        return text

    for row in inner_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text:
                        run.text = replace_in_text(run.text)

def convert_to_pdf(input_path, output_path):
    """Convert a docx file to PDF using LibreOffice."""
    if not os.path.exists(LIBREOFFICE_PATH):
        st.error(f"LibreOffice not found at {LIBREOFFICE_PATH}. Please ensure it is installed and the path is correct.")
        raise FileNotFoundError(f"LibreOffice not found at {LIBREOFFICE_PATH}")
    
    # Ensure input file is readable and output directory is writable
    os.chmod(input_path, stat.S_IRWXU)
    output_dir = os.path.dirname(output_path)
    if not os.access(output_dir, os.W_OK):
        st.error(f"Output directory {output_dir} is not writable.")
        raise PermissionError(f"Output directory {output_dir} is not writable")

    try:
        subprocess.run([
            LIBREOFFICE_PATH,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            output_dir,
            input_path
        ], check=True, capture_output=True, text=True)
        # Check if the output PDF exists and is non-empty
        if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
            st.error(f"PDF conversion failed: {output_path} is empty or not created.")
            raise ValueError(f"PDF conversion failed: {output_path} is empty or not created")
    except subprocess.CalledProcessError as e:
        st.error(f"Failed to convert {input_path} to PDF: {e.stderr}")
        raise

if uploaded_excel and uploaded_template:
    # Create output folder if it doesn't exist
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)

    # Read the Excel file
    df = pd.read_excel(uploaded_excel)

    # Create a new Word document to store all results
    output_doc = Document()

    # Load the initial template
    template_doc = Document(uploaded_template)

    # Check for outer tables in the template
    num_outer_tables = len(template_doc.tables)
    if num_outer_tables == 0:
        st.error("No tables found in the template. Please check the template structure.")
        st.stop()

    # Calculate total steps for progress bar (one step per group of 4 rows)
    total_groups = sum(len(area_df) // 4 + (1 if len(area_df) % 4 else 0) for _, area_df in df.groupby("FINAL_AREA"))
    progress_bar = st.progress(0)
    current_step = 0

    # Store merged PDF filenames for zipping
    merged_pdf_filenames = []

    # Group by FINAL_AREA
    for area, area_df in df.groupby("FINAL_AREA"):
        temp_docs = []  # Store docx paths for this area
        pdf_merger = PdfMerger()  # One merger per FINAL_AREA

        # Process rows in groups of 4 (based on template capacity)
        for group_idx, page_start in enumerate(range(0, len(area_df), 4)):
            # Update progress bar
            current_step += 1
            progress_bar.progress(min(current_step / total_groups, 1.0))

            # Create a new document instance using a fresh template copy
            temp_doc = Document(uploaded_template)

            # Access the outer table in the temporary document
            if not temp_doc.tables:
                st.error("Failed to load tables in the template. Please check the template structure.")
                break
            current_outer_table = temp_doc.tables[0]

            # Get the rows for this group (up to 4 rows)
            page_rows = area_df.iloc[page_start:page_start + 4]

            # Iterate over the rows of the outer table
            for row_idx, row in enumerate(current_outer_table.rows):
                if row_idx < len(page_rows):
                    # Fill inner table with data
                    row_data = page_rows.iloc[row_idx]
                    mapping = {f"«{col.upper()}»": str(row_data[col]) for col in df.columns if pd.notnull(row_data[col])}
                    for cell in row.cells:
                        if cell.tables:  # Check if the cell contains an inner table
                            inner_table = cell.tables[0]
                            fill_inner_table(inner_table, mapping)
                else:
                    # Clear placeholders in unused inner tables
                    for cell in row.cells:
                        if cell.tables:
                            inner_table = cell.tables[0]
                            clear_placeholders(inner_table)

            # Save the temporary document to the output folder
            doc_filename = os.path.join(OUTPUT_FOLDER, f"area_{area}_group_{group_idx}.docx")
            temp_doc.save(doc_filename)
            temp_docs.append(doc_filename)

            # Append the temporary document content to the output document
            for element in temp_doc.element.body:
                output_doc.element.body.append(deepcopy(element))

        # Convert all temporary documents for this FINAL_AREA to PDF
        for doc_path in temp_docs:
            pdf_filename = os.path.splitext(doc_path)[0] + ".pdf"
            convert_to_pdf(doc_path, pdf_filename)
            # Read the PDF into BytesIO for merging
            pdf_io = BytesIO()
            with open(pdf_filename, "rb") as f:
                pdf_io.write(f.read())
            pdf_io.seek(0)
            # Verify the PDF is non-empty before merging
            if pdf_io.getbuffer().nbytes == 0:
                st.error(f"PDF file {pdf_filename} is empty and will be skipped.")
                os.unlink(pdf_filename)
                continue
            # Append the PDF to the merger
            pdf_merger.append(pdf_io)
            # Keep the PDF file in the output folder for inspection

        # Save the merged PDF for this FINAL_AREA
        merged_pdf_filename = os.path.join(OUTPUT_FOLDER, f"merged_{area}.pdf")
        merged_pdf_io = BytesIO()
        try:
            pdf_merger.write(merged_pdf_io)
            merged_pdf_io.seek(0)
            if merged_pdf_io.getbuffer().nbytes == 0:
                st.error(f"Merged PDF for area {area} is empty. No valid PDFs were generated.")
            else:
                # Save merged PDF to output folder
                with open(merged_pdf_filename, "wb") as f:
                    f.write(merged_pdf_io.getbuffer())
                merged_pdf_filenames.append(merged_pdf_filename)
        except Exception as e:
            st.error(f"Failed to merge PDFs for area {area}: {str(e)}")
        finally:
            pdf_merger.close()

    # Save the final Word document
    output_doc_filename = os.path.join(OUTPUT_FOLDER, "filled_template.docx")
    output_doc.save(output_doc_filename)
    output_doc_io = BytesIO()
    with open(output_doc_filename, "rb") as f:
        output_doc_io.write(f.read())
    output_doc_io.seek(0)

    # Provide download link for the combined Word document
    st.download_button(
        label="Download Filled Word Document",
        data=output_doc_io,
        file_name="filled_template.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Create a zip file containing all merged PDFs
    zip_io = BytesIO()
    with zipfile.ZipFile(zip_io, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for pdf_filename in merged_pdf_filenames:
            if os.path.exists(pdf_filename):
                zip_file.write(pdf_filename, os.path.basename(pdf_filename))
    zip_io.seek(0)

    # Provide download button for the zip file
    if zip_io.getbuffer().nbytes > 0:
        st.download_button(
            label="Download All Merged PDFs as Zip",
            data=zip_io,
            file_name="merged_pdfs_by_area.zip",
            mime="application/zip"
        )
    else:
        st.error("No valid merged PDFs were generated to include in the zip file.")

    st.write(f"Note: Each group of up to 4 Excel rows within a FINAL_AREA is filled into a fresh copy of the template's inner tables. Unused inner tables have placeholders replaced with empty strings. Temporary docx files and PDFs are saved to {OUTPUT_FOLDER} for inspection. Merged PDFs are created per FINAL_AREA after processing all rows in each area, and all merged PDFs are included in a zip file.")
else:
    st.write("Please upload both the Excel file and the Word template to proceed.")
