from fastapi import FastAPI, HTTPException, UploadFile, File, Depends, Response, Request
from fastapi.responses import FileResponse, StreamingResponse, RedirectResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional, Dict
import pandas as pd
from docx import Document
import re
import os
import uuid
from pathlib import Path
from tempfile import NamedTemporaryFile, TemporaryDirectory
from fastapi.staticfiles import StaticFiles
from docx.shared import Inches, Pt
from io import BytesIO
from PyPDF2 import PdfMerger
from barcode import Code128
from barcode.writer import ImageWriter
import shutil
import logging
from datetime import datetime, timedelta
from dotenv import load_dotenv
import gspread
from google.oauth2.service_account import Credentials
import win32com.client
import pythoncom
import zipfile
from ftplib import FTP
import time
import psutil
import qrcode
import json
import asyncio
import subprocess
import requests
import sqlite3
import win32print

# Initialize FastAPI app
app = FastAPI(title="DL Generator API")
app.mount("/static", StaticFiles(directory="static"), name="static")
origins = [
    "http://localhost:8000",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load environment variables
env_path = Path(__file__).parent / "config" / ".env"
load_dotenv(env_path)

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Create output and barcode directories
OUTPUT_DIR = Path("output").absolute()
BARCODE_DIR = Path("barcode_images").absolute()
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(BARCODE_DIR, exist_ok=True)

# Database setup
DB_PATH = "dl_generator.db"

def init_database():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Create users table with multiple clients support
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE NOT NULL,
            clients TEXT NOT NULL,
            access TEXT NOT NULL DEFAULT 'user',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Create audit_trail table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS audit_trail (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            client TEXT NOT NULL,
            processed_by TEXT NOT NULL,
            processed_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            total_accounts INTEGER NOT NULL,
            mode TEXT NOT NULL,
            template_folder TEXT,
            dl_type TEXT
        )
    ''')
    
    # Create processed_accounts table to store actual account data
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS processed_accounts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            audit_id INTEGER NOT NULL,
            dl_code TEXT,
            leads_chname TEXT,
            dl_address TEXT,
            final_area TEXT,
            processed_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (audit_id) REFERENCES audit_trail (id) ON DELETE CASCADE
        )
    ''')
    
    conn.commit()
    conn.close()

# Initialize database
init_database()

FTP_CONFIG = {
    "hostname": os.getenv("OMKT_FTP_HOSTNAME"),
    "port": int(os.getenv("OMKT_FTP_PORT", 21)),
    "username": os.getenv("OMKT_FTP_USERNAME"),
    "password": os.getenv("OMKT_FTP_PASSWORD")
}
APP_ID = os.getenv("APP_ID")
APP_SECRET = os.getenv("APP_SECRET")
REDIRECT_URI = os.getenv("REDIRECT_URI", "http://localhost:5000/api/lark_callback")
AUTH_URL = f"https://open.larksuite.com/open-apis/authen/v1/authorize?app_id={APP_ID}&redirect_uri={REDIRECT_URI}"
TOKEN_URL = "https://open.larksuite.com/open-apis/auth/v3/tenant_access_token/internal"
USER_ACCESS_TOKEN_URL = "https://open.larksuite.com/open-apis/authen/v1/oidc/access_token"
USER_INFO_URL = "https://open.larksuite.com/open-apis/authen/v1/user_info"
REFRESH_TOKEN_URL = "https://open.larksuite.com/open-apis/authen/v1/oidc/refresh_access_token"
SERVICE_ACCOUNT_JSON = "config/dl_automation_sheet.json"
SPREADSHEET_ID = "1M0Vmmf9HfPRB0oSeJR_xUAZPPpTJ4xsR3gYDQMvnu5k"
SHEET_NAME = "LetterHeads"
SESSION_FILE = Path("sessions.json")

# Pydantic Models
class LoginRequest(BaseModel):
    username: str
    password: str

class User(BaseModel):
    email: str
    clients: List[str]
    access: str

class UserCreate(BaseModel):
    email: str
    clients: List[str]
    access: str

class FolderRequest(BaseModel):
    folder: str

class TemplateRequest(BaseModel):
    folder: str
    dl_type: str
    template: str

class ModeRequest(BaseModel):
    mode: str

class OutputFormatRequest(BaseModel):
    format: str  # "zip" or "print"

class AuditEntry(BaseModel):
    client: str
    processed_by: str
    total_accounts: int
    mode: str
    template_folder: Optional[str] = None
    dl_type: Optional[str] = None

SESSION_STORE = {}

# Session state simulation - Use a function to get fresh state
def get_fresh_session_state():
    return {
        'selected_mode': '',
        'selected_folder': '',
        'selected_dl_type': '',
        'selected_template': '',
        'base_template': None,
        'template_path': None,
        'header_footer_template_path': None,
        'transmittal_template_path': None,
        'placeholders': [],
        'download_completed': False,
        'files_to_cleanup': [],
        'zip_path': None,
        'template_combined': False,
        'output_format': 'zip'
    }

# Initialize with fresh state
SESSION_STATE = get_fresh_session_state()

class FTPConnection:
    def __init__(self, hostname, port, username, password):
        self.ftp = None
        self.hostname = hostname
        self.port = port
        self.username = username
        self.password = password

    def connect(self):
        try:
            self.ftp = FTP()
            self.ftp.connect(self.hostname, self.port)
            self.ftp.login(self.username, self.password)
            logger.info("Connected to FTP server")
            return self.ftp
        except Exception as e:
            logger.error(f"Failed to connect to FTP: {e}")
            return None

    def close(self):
        if self.ftp:
            try:
                self.ftp.quit()
                logger.info("Closed FTP connection")
            except Exception as e:
                logger.warning(f"Error closing FTP connection: {e}")
            self.ftp = None

    def __enter__(self):
        self.connect()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()

def get_app_access_token():
    headers = {"Content-Type": "application/json; charset=utf-8"}
    data = {"app_id": APP_ID, "app_secret": APP_SECRET}
    try:
        response = requests.post(TOKEN_URL, json=data, headers=headers)
        response.raise_for_status()
        return response.json().get("tenant_access_token")
    except Exception as e:
        logger.error(f"Failed to get app access token: {e}")
        raise HTTPException(status_code=500, detail="Failed to get app access token")

def get_user_access_token(auth_code: str, tenant_access_token: str):
    headers = {
        "Content-Type": "application/json; charset=utf-8",
        "Authorization": f"Bearer {tenant_access_token}"
    }
    data = {"grant_type": "authorization_code", "code": auth_code}
    try:
        response = requests.post(USER_ACCESS_TOKEN_URL, json=data, headers=headers)
        response_data = response.json()
        if "data" not in response_data or "access_token" not in response_data["data"]:
            logger.error(f"Failed to get user access token: {response_data}")
            raise HTTPException(status_code=401, detail="Failed to get user access token")
        expires_in = response_data["data"]["expires_in"]
        expires_at = datetime.now() + timedelta(seconds=expires_in)
        return (
            response_data["data"]["access_token"],
            response_data["data"]["refresh_token"],
            expires_at
        )
    except Exception as e:
        logger.error(f"Error getting user access token: {e}")
        raise HTTPException(status_code=500, detail="Error getting user access token")

def refresh_user_access_token(refresh_token: str, tenant_access_token: str):
    headers = {
        "Content-Type": "application/json; charset=utf-8",
        "Authorization": f"Bearer {tenant_access_token}"
    }
    data = {"grant_type": "refresh_token", "refresh_token": refresh_token}
    try:
        response = requests.post(REFRESH_TOKEN_URL, json=data, headers=headers)
        response_data = response.json()
        if response_data.get("code") != 0:
            logger.error(f"Failed to refresh access token: {response_data.get('message', 'Unknown error')}")
            raise HTTPException(status_code=401, detail="Failed to refresh access token")
        expires_in = response_data["data"]["expires_in"]
        expires_at = datetime.now() + timedelta(seconds=expires_in)
        return (
            response_data["data"]["access_token"],
            response_data["data"]["refresh_token"],
            expires_at
        )
    except Exception as e:
        logger.error(f"Error refreshing access token: {e}")
        raise HTTPException(status_code=500, detail="Error refreshing access token")

def get_user_info(user_access_token: str):
    headers = {"Authorization": f"Bearer {user_access_token}"}
    try:
        response = requests.get(USER_INFO_URL, headers=headers)
        response.raise_for_status()
        response_data = response.json()
        if "data" not in response_data:
            logger.error(f"Failed to get user information: {response_data}")
            raise HTTPException(status_code=401, detail="Failed to get user information")
        return response_data["data"]
    except Exception as e:
        logger.error(f"Error getting user info: {e}")
        raise HTTPException(status_code=500, detail="Error getting user info")

def get_user_clients_and_access(email: str):
    """Get user's assigned clients and access level from database"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT clients, access FROM users WHERE email = ?", (email,))
    result = cursor.fetchone()
    conn.close()
    if result:
        clients_str, access = result
        clients = clients_str.split(',') if clients_str else []
        return clients, access
    return [], None

# Dependency to get current user
async def get_current_user(request: Request):
    session_id = request.cookies.get("session_id")
    if not session_id or session_id not in SESSION_STORE:
        raise HTTPException(status_code=401, detail="Invalid or missing session")
    session = SESSION_STORE[session_id]
    expires_at = datetime.fromisoformat(session["expires_at"])
    if datetime.now() > expires_at:
        tenant_access_token = get_app_access_token()
        user_access_token, refresh_token, expires_at = refresh_user_access_token(session["refresh_token"], tenant_access_token)
        if not user_access_token:
            raise HTTPException(status_code=401, detail="Session expired")
        SESSION_STORE[session_id]["user_access_token"] = user_access_token
        SESSION_STORE[session_id]["refresh_token"] = refresh_token
        SESSION_STORE[session_id]["expires_at"] = expires_at.isoformat()
        save_sessions()
        session["user_access_token"] = user_access_token
        session["expires_at"] = expires_at
    return session["user_info"]

def load_sessions():
    global SESSION_STORE
    if SESSION_FILE.exists():
        try:
            with SESSION_FILE.open("r") as f:
                SESSION_STORE = json.load(f)
        except Exception as e:
            logger.warning(f"Failed to load sessions: {e}")

# Save sessions
def save_sessions():
    try:
        with SESSION_FILE.open("w") as f:
            json.dump(SESSION_STORE, f)
    except Exception as e:
        logger.error(f"Failed to save sessions: {e}")

def add_audit_entry(client: str, processed_by: str, total_accounts: int, mode: str, template_folder: str = None, dl_type: str = None):
    """Add entry to audit trail"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO audit_trail (client, processed_by, total_accounts, mode, template_folder, dl_type)
        VALUES (?, ?, ?, ?, ?, ?)
    ''', (client, processed_by, total_accounts, mode, template_folder, dl_type))
    conn.commit()
    conn.close()

# Call at startup
load_sessions()

# Enhanced cleanup function
def cleanup_files(file_paths):
    """Enhanced cleanup function that handles all types of files"""
    for file_path in file_paths:
        try:
            if file_path and Path(file_path).exists():
                path_obj = Path(file_path)
                if path_obj.is_file():
                    path_obj.unlink()
                    logger.info(f"Deleted file: {file_path}")
                elif path_obj.is_dir():
                    shutil.rmtree(path_obj)
                    logger.info(f"Deleted directory: {file_path}")
        except Exception as e:
            logger.warning(f"Failed to delete {file_path}: {e}")

def cleanup_output_directory():
    """Clean up all files in the output directory"""
    try:
        if OUTPUT_DIR.exists():
            for item in OUTPUT_DIR.iterdir():
                try:
                    if item.is_file():
                        item.unlink()
                        logger.info(f"Deleted output file: {item}")
                    elif item.is_dir():
                        shutil.rmtree(item)
                        logger.info(f"Deleted output directory: {item}")
                except Exception as e:
                    logger.warning(f"Failed to delete output item {item}: {e}")
    except Exception as e:
        logger.error(f"Error cleaning output directory: {e}")

def reset_session_state():
    """Reset the session state to fresh values"""
    global SESSION_STATE
    # Clean up any existing files first
    cleanup_files(SESSION_STATE.get('files_to_cleanup', []))
    cleanup_output_directory()
    
    # Reset to fresh state
    SESSION_STATE = get_fresh_session_state()
    logger.info("Session state has been reset")

# Add new endpoint to get audit details
@app.get("/api/audit_details/{audit_id}")
async def get_audit_details(audit_id: int, page: int = 1, limit: int = 50, user_info: dict = Depends(get_current_user)):
    """Get detailed information about processed accounts for a specific audit entry"""
    user_email = user_info.get("email", "")
    user_clients, user_access = get_user_clients_and_access(user_email)
    
    # Only admin can access audit details
    if user_access != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # First, verify the audit entry exists and get its details
    cursor.execute("SELECT id, client, processed_by, processed_at, total_accounts, mode FROM audit_trail WHERE id = ?", (audit_id,))
    audit_entry = cursor.fetchone()
    
    if not audit_entry:
        conn.close()
        raise HTTPException(status_code=404, detail="Audit entry not found")
    
    # Get total count of accounts
    cursor.execute("SELECT COUNT(*) as total FROM processed_accounts WHERE audit_id = ?", (audit_id,))
    total_count = cursor.fetchone()["total"]
    
    # Get paginated accounts
    offset = (page - 1) * limit
    cursor.execute("""
        SELECT dl_code, leads_chname, dl_address, final_area
        FROM processed_accounts
        WHERE audit_id = ?
        ORDER BY id
        LIMIT ? OFFSET ?
    """, (audit_id, limit, offset))
    
    accounts = []
    for row in cursor.fetchall():
        accounts.append({
            "dl_code": row["dl_code"],
            "name": row["leads_chname"],
            "address": row["dl_address"],
            "area": row["final_area"]
        })
    
    conn.close()
    
    total_pages = (total_count + limit - 1) // limit
    
    return {
        "audit_id": audit_id,
        "client": audit_entry["client"],
        "processed_by": audit_entry["processed_by"],
        "processed_at": audit_entry["processed_at"],
        "total_accounts": audit_entry["total_accounts"],
        "mode": audit_entry["mode"],
        "accounts": accounts,
        "pagination": {
            "current_page": page,
            "total_pages": total_pages,
            "total_count": total_count,
            "limit": limit,
            "has_next": page < total_pages,
            "has_prev": page > 1
        }
    }

@app.get("/api/check_session")
async def check_session(user_info: dict = Depends(get_current_user)):
    user_email = user_info.get("email", "")
    user_clients, user_access = get_user_clients_and_access(user_email)
    
    # Determine access level
    if user_access == "admin":
        access_level = "admin"
    elif user_access == "user":
        access_level = "user"
    else:
        # Default logic for users not in database
        access_level = "admin" if user_email.endswith("@spmadridlaw.com") else "user"
    
    return {
        "success": True,
        "username": user_info.get("name", user_info.get("email", "Unknown")),
        "role": access_level.title(),
        "access": access_level,
        "clients": user_clients,
        "avatar": user_info
    }
    
# API Endpoints
@app.get("/api/login")
async def login():
    return RedirectResponse(url=AUTH_URL)

@app.get("/api/lark_callback")
async def lark_callback(code: str, response: Response):
    try:
        tenant_access_token = get_app_access_token()
        user_access_token, refresh_token, expires_at = get_user_access_token(code, tenant_access_token)
        user_info = get_user_info(user_access_token)
        if not user_info:
            raise HTTPException(status_code=401, detail="Failed to retrieve user info")
        
        # Generate session ID
        session_id = str(uuid.uuid4())
        SESSION_STORE[session_id] = {
            "user_access_token": user_access_token,
            "refresh_token": refresh_token,
            "expires_at": expires_at.isoformat(),
            "user_info": user_info
        }
        save_sessions()
        
        # Get user access level from database
        user_email = user_info.get("email", "")
        user_clients, user_access = get_user_clients_and_access(user_email)
        
        if user_access == "admin":
            role = "Admin"
        elif user_access == "user":
            role = "User"
        else:
            # Default logic for users not in database
            role = "Admin" if user_email.endswith("@spmadridlaw.com") else "User"
        
        # Set session cookie
        response.set_cookie(key="session_id", value=session_id, httponly=True, secure=True, samesite="lax")
        
        return {
            "success": True,
            "role": role,
            "username": user_info.get("name", user_info.get("email", "Unknown"))
        }
    except Exception as e:
        logger.error(f"Error in Lark callback: {e}")
        raise HTTPException(status_code=500, detail="Authentication failed")

@app.get("/api/logout")
async def logout(request: Request, response: Response):
    session_id = request.cookies.get("session_id")
    if session_id in SESSION_STORE:
        del SESSION_STORE[session_id]
        save_sessions()  # Save after deleting
    response.delete_cookie("session_id")
    
    # Reset session state on logout
    reset_session_state()
    
    return {"success": True, "message": "Logged out successfully"}

def get_ftp_folders(ftp):
    try:
        ftp_path = "/DL AUTOMATION/Template DL V2/Content"
        ftp.cwd(ftp_path)
        items = ftp.nlst()
        folders = []
        current_dir = ftp.pwd()
        for item in items:
            try:
                ftp.cwd(item)
                folders.append(item)
                ftp.cwd(current_dir)
            except:
                continue
        return sorted(folders)
    except Exception as e:
        logger.error(f"Failed to retrieve folders from FTP: {e}")
        return []

def get_ftp_templates(ftp, folder_name):
    try:
        ftp_path = f"/DL AUTOMATION/Template DL V2/Content/{folder_name}"
        ftp.cwd(ftp_path)
        templates = [item for item in ftp.nlst() if item.lower().endswith('.docx')]
        return sorted(templates)
    except Exception as e:
        logger.error(f"Failed to retrieve templates from FTP folder {folder_name}: {e}")
        return []

def download_ftp_template(ftp, folder_name, template_name, is_header_footer=False, is_transmittal=False):
    try:
        if is_transmittal:
            ftp_path = "/DL AUTOMATION/Template Transmittal V2"
        elif is_header_footer:
            ftp_path = "/DL AUTOMATION/Template DL V2/Letter Head"
        else:
            ftp_path = f"/DL AUTOMATION/Template DL V2/Content/{folder_name}"
        ftp.cwd(ftp_path)
        with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            ftp.retrbinary(f"RETR {template_name}", tmp.write)
            tmp_path = tmp.name
        if os.path.exists(tmp_path):
            logger.info(f"Downloaded template {template_name} to: {tmp_path}")
            return tmp_path
        raise Exception(f"Failed to download template {template_name}")
    except Exception as e:
        logger.error(f"Failed to download template {template_name}: {e}")
        return None

def fetch_signature_from_ftp(ftp):
    # Try today and yesterday's date
    for days_ago in [0, 1]:
        date_str = (datetime.today() - timedelta(days=days_ago)).strftime('%m-%d-%Y')
        # ftp_path = f"field/DL/ATTY SIGNATURE/{date_str}"
        ftp_path = f"field/DL/ATTY SIGNATURE/05-29-2025"
        try:
            ftp.cwd(ftp_path)
            with NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                ftp.retrbinary("RETR attySignature.PNG", tmp.write)
                tmp_path = tmp.name
            if os.path.exists(tmp_path):
                return tmp_path
        except Exception as e:
            logger.warning(f"Failed to fetch signature for {date_str}: {e}")
    logger.error("Failed to fetch signature from FTP for today or yesterday.")
    return None

def get_sheet_data(service_account_json_path, spreadsheet_id, sheet_name="LetterHeads"):
    try:
        # Updated scopes to match exactly what gspread expects
        scopes = [
            'https://spreadsheets.google.com/feeds',
            'https://www.googleapis.com/auth/spreadsheets',
            'https://www.googleapis.com/auth/spreadsheets',
            'https://www.googleapis.com/auth/drive'
        ]
        
        credentials = Credentials.from_service_account_file(
            service_account_json_path,
            scopes=scopes
        )
        
        client = gspread.authorize(credentials)
        spreadsheet = client.open_by_key(spreadsheet_id)
        worksheet = spreadsheet.worksheet(sheet_name)
        sheet_data = worksheet.get_all_records()
        df = pd.DataFrame(sheet_data)
        required_columns = ["CAMPAIGN", "DL TYPE", "FILE"]
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            logger.error(f"Missing columns in spreadsheet: {missing_columns}")
            return pd.DataFrame()
        df = df[["CAMPAIGN", "DL TYPE", "FILE"]].dropna()
        logger.info(f"Retrieved sheet data with {len(df)} valid rows")
        return df
    except Exception as e:
        logger.error(f"Failed to access Google Sheets: {e}")
        return pd.DataFrame()

def combine_templates(header_footer_path, content_path, signature_img_path, word_app):
    try:
        header_footer_doc = Document(header_footer_path)
        content_doc = Document(content_path)
        # while header_footer_doc.paragraphs:
        #     header_footer_doc.paragraphs[0]._element.getparent().remove(header_footer_doc.paragraphs[0]._element)
        for elem in content_doc.element.body:
            header_footer_doc.element.body.append(elem)
        # for para in header_footer_doc.paragraphs:
        #     para.paragraph_format.space_before = Pt(0)
        #     para.paragraph_format.space_after = Pt(0)
         # ✅ Set top margin to 1.04 inches
        for section in header_footer_doc.sections:
            section.header_distance = Inches(0.2)
            section.footer_distance = Inches(0.2)

        if signature_img_path and Path(signature_img_path).exists():
            with TemporaryDirectory() as temp_dir:
                temp_doc_path = Path(temp_dir) / "temp_combined_doc.docx"
                header_footer_doc.save(temp_doc_path)
                temp_doc_path = replace_in_text_boxes(header_footer_doc, "«IMAGE_SIGNATURE»", signature_img_path, word_app, temp_doc_path)
                header_footer_doc = Document(temp_doc_path)
        
        # Set flag that template has been combined
        SESSION_STATE['template_combined'] = True
        
        return header_footer_doc
    except Exception as e:
        logger.error(f"Failed to combine templates or replace signature: {e}")
        return None

def number_to_words(number):
    if number == 0:
        return "ZERO"
    units = ["", "ONE", "TWO", "THREE", "FOUR", "FIVE", "SIX", "SEVEN", "EIGHT", "NINE"]
    teens = ["", "ELEVEN", "TWELVE", "THIRTEEN", "FOURTEEN", "FIFTEEN", "SIXTEEN", "SEVENTEEN", "EIGHTEEN", "NINETEEN"]
    tens = ["", "TEN", "TWENTY", "THIRTY", "FORTY", "FIFTY", "SIXTY", "SEVENTY", "EIGHTY", "NINETY"]
    thousands = ["", "THOUSAND", "MILLION"]
    def convert_hundreds(num):
        if num == 0:
            return ""
        result = []
        if num >= 100:
            result.append(units[num // 100] + " HUNDRED")
            num %= 100
            if num > 0:
                result.append("AND")
        if num >= 20:
            result.append(tens[num // 10])
            num %= 10
            if num > 0:
                result.append(units[num])
        elif num >= 11:
            result.append(teens[num - 10])
        elif num == 10:
            result.append(tens[1])
        elif num > 0:
            result.append(units[num])
        return " ".join(result)
    result = []
    i = 0
    while number > 0:
        chunk = number % 1000
        if chunk > 0:
            chunk_words = convert_hundreds(chunk)
            if thousands[i]:
                chunk_words += f" {thousands[i]}"
            result.insert(0, chunk_words)
        number //= 1000
        i += 1
    return ", ".join(result).replace(", AND", " AND").strip()

def amount_to_words(amount_str):
    try:
        amount = float(amount_str.replace(",", ""))
        pesos = int(amount)
        cents = int(round((amount - pesos) * 100))
        pesos_words = number_to_words(pesos)
        if pesos == 0:
            pesos_words = "ZERO"
        cents_words = number_to_words(cents) if cents > 0 else "ZERO"
        result = f"{pesos_words} PESOS"
        if cents > 0:
            result += f", AND {cents_words} CENTS"
        else:
            result += ", AND ZERO CENTS"
        return result
    except Exception as e:
        logger.error(f"Failed to convert amount to words: {amount_str}, error: {e}")
        return "ERROR CONVERTING AMOUNT"

def generate_barcode(barcode_value):
    try:
        barcode = Code128(barcode_value, writer=ImageWriter())
        buffer = BytesIO()
        barcode.write(buffer, options={"write_text": False, "module_width": 1, "module_height": 8, "quiet_zone": 2.0})
        buffer.seek(0)
        return buffer
    except Exception as e:
        logger.error(f"Failed to generate barcode for {barcode_value}: {e}")
        return None

def generate_qrcode(data):
    try:
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        qr.add_data(data)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        buffer = BytesIO()
        img.save(buffer, format="PNG")
        buffer.seek(0)
        return buffer
    except Exception as e:
        logger.error(f"Failed to generate QR code for {data}: {e}")
        return None

def extract_placeholders(doc):
    placeholders = set()
    for para in doc.paragraphs:
        if "«" in para.text:
            matches = re.findall(r"«(.*?)»", para.text)
            placeholders.update(["«" + m.strip() + "»" for m in matches])
    for node in doc._element.iter():
        if node.tag.endswith("}t") and node.text and "«" in node.text:
            matches = re.findall(r"«(.*?)»", node.text)
            placeholders.update(["«" + m.strip() + "»" for m in matches])
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header is not None:
                for para in header.paragraphs:
                    if "«" in para.text:
                        matches = re.findall(r"«(.*?)»", para.text)
                        placeholders.update(["«" + m.strip() + "»" for m in matches])
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if "«" in para.text:
                                    matches = re.findall(r"«(.*?)»", para.text)
                                    placeholders.update(["«" + m.strip() + "»" for m in matches])
                for node in header._element.iter():
                    if node.tag.endswith("}t") and node.text and "«" in node.text:
                        matches = re.findall(r"«(.*?)»", node.text)
                        placeholders.update(["«" + m.strip() + "»" for m in matches])
        for footer in [section.footer, section.first_page_footer, section.even_page_header]:
            if footer is not None:
                for para in footer.paragraphs:
                    if "«" in para.text:
                        matches = re.findall(r"«(.*?)»", para.text)
                        placeholders.update(["«" + m.strip() + "»" for m in matches])
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if "«" in para.text:
                                    matches = re.findall(r"«(.*?)»", para.text)
                                    placeholders.update(["«" + m.strip() + "»" for m in matches])
                for node in footer._element.iter():
                    if node.tag.endswith("}t") and node.text and "«" in node.text:
                        matches = re.findall(r"«(.*?)»", node.text)
                        placeholders.update(["«" + m.strip() + "»" for m in matches])
    return sorted(placeholders)

def replace_in_text_boxes(doc, find_str, replace_with_image_path, word_app, temp_doc_path):
    try:
        word_doc = word_app.Documents.Open(str(temp_doc_path.absolute()))
        modified = False
        for shape in word_doc.Shapes:
            if shape.TextFrame.HasText:
                text_range = shape.TextFrame.TextRange
                find = text_range.Find
                find.Text = find_str
                find.Forward = True
                find.Wrap = 1
                while find.Execute():
                    found_range = text_range.Duplicate
                    found_range.Find.Execute(FindText=find_str)
                    found_range.Text = ""
                    inline_shape = found_range.InlineShapes.AddPicture(
                        FileName=str(Path(replace_with_image_path).absolute()),
                        LinkToFile=False,
                        SaveWithDocument=True
                    )
                    inline_shape.Width = 110
                    inline_shape.Height = 50
                    modified = True
        word_doc.Save()
        word_doc.Close(SaveChanges=True)
        if modified:
            logger.info(f"Replaced '{find_str}' with image in shapes.")
        return temp_doc_path
    except Exception as e:
        logger.error(f"Error replacing text in shapes: {e}")
        try:
            word_doc.Close(SaveChanges=False)
        except:
            pass
        return temp_doc_path

def fill_template(doc, mapping, barcode_buffer=None):
    def replace_in_text(text):
        for k, v in mapping.items():
            if k == "«IMAGE_BARCODE»" and v:
                continue
            text = text.replace(k, str(v))
        return text
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = replace_in_text(run.text)
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = replace_in_text(run.text)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if "«IMAGE_BARCODE»" in para.text and barcode_buffer:
                                    para.clear()
                                    cell.paragraphs[0].paragraph_format.left_indent = Pt(-20)
                                    run = para.add_run()
                                    run.add_picture(barcode_buffer, width=Inches(3.0), height=Inches(0.35))
                                else:
                                    for run in para.runs:
                                        if run.text:
                                            run.text = replace_in_text(run.text)
        for footer in [section.footer, section.first_page_footer, section.even_page_header]:
            if footer:
                for para in footer.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = replace_in_text(run.text)
    return doc

def clear_placeholders(inner_table):
    def replace_in_text(text):
        return re.sub(r'«[^»]+»', "", text)
    for row in inner_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text:
                        run.text = replace_in_text(run.text)

def fill_inner_table(inner_table, mapping, qrcode_buffer=None):
    def replace_in_text(text):
        for k, v in mapping.items():
            if k == "«IMAGE_QRCODE»":
                continue
            text = text.replace(k, str(v))
        return text
    for row in inner_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if "«IMAGE_QRCODE»" in para.text and qrcode_buffer:
                    para.clear()
                    para.paragraph_format.left_indent = Pt(-5)
                    run = para.add_run()
                    run.add_picture(qrcode_buffer, width=Inches(1), height=Inches(1))
                else:
                    for run in para.runs:
                        if run.text:
                            run.text = replace_in_text(run.text)

def kill_libreoffice_processes():
    killed_count = 0
    try:
        for proc in psutil.process_iter(['pid', 'name']):
            if 'soffice' in proc.info['name'].lower():
                try:
                    proc.kill()
                    killed_count += 1
                except psutil.NoSuchProcess:
                    continue
        if killed_count > 0:
            logger.info(f"Killed {killed_count} existing LibreOffice processes")
        time.sleep(2)
    except Exception as e:
        logger.error(f"Error killing LibreOffice processes: {e}")

def convert_batch_with_retry(batch_files, output_dir, batch_id, timeout=180):
    max_retries = 3
    output_dir = Path(output_dir)
    for attempt in range(max_retries):
        try:
            with TemporaryDirectory() as temp_batch_dir:
                temp_output = Path(temp_batch_dir)
                logger.debug(f"Batch {batch_id} (Attempt {attempt + 1}): Converting {len(batch_files)} files...")
                cmd = [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    "--headless",
                    "--invisible",
                    "--nodefault",
                    "--nolockcheck",
                    "--nologo",
                    "--norestore",
                    "--convert-to", "pdf",
                    "--outdir", str(temp_output)
                ] + [str(Path(f)) for f in batch_files]
                process = subprocess.Popen(
                    cmd,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    creationflags=subprocess.CREATE_NO_WINDOW if os.name == 'nt' else 0,
                    text=True
                )
                try:
                    stdout, stderr = process.communicate(timeout=timeout)
                    if process.returncode != 0:
                        logger.error(f"Batch {batch_id} LibreOffice error (code {process.returncode}): {stderr}")
                        if attempt < max_retries - 1:
                            logger.info(f"Retrying batch {batch_id}...")
                            kill_libreoffice_processes()
                            time.sleep(2)
                            continue
                        return [], batch_files
                except subprocess.TimeoutExpired:
                    process.kill()
                    logger.error(f"Batch {batch_id} timeout after {timeout} seconds (attempt {attempt + 1})")
                    if attempt < max_retries - 1:
                        logger.info(f"Retrying batch {batch_id}...")
                        kill_libreoffice_processes()
                        time.sleep(2)
                        continue
                    return [], batch_files
                batch_pdfs = []
                failed_files = []
                for docx_path in batch_files:
                    docx_name = Path(docx_path).stem
                    temp_pdf = temp_output / f"{docx_name}.pdf"
                    final_pdf = output_dir / f"{docx_name}.pdf"
                    if temp_pdf.exists():
                        shutil.move(str(temp_pdf), str(final_pdf))
                        batch_pdfs.append(str(final_pdf))
                    else:
                        failed_files.append(docx_path)
                        logger.warning(f"Batch {batch_id}: Failed to convert {Path(docx_path).name}")
                success_rate = len(batch_pdfs) / len(batch_files) * 100 if batch_files else 0
                logger.info(f"Batch {batch_id} result: {len(batch_pdfs)}/{len(batch_files)} successful ({success_rate:.1f}%)")
                return batch_pdfs, failed_files
        except Exception as e:
            logger.error(f"Batch {batch_id} conversion error (attempt {attempt + 1}): {e}")
            if attempt < max_retries - 1:
                kill_libreoffice_processes()
                time.sleep(2)
                continue
            return [], batch_files
    return [], batch_files

def batch_convert_libreoffice(docx_files, output_dir, batch_size=350):
    if not docx_files:
        return []
    pdf_files = []
    total_failed = []
    output_dir = Path(output_dir)
    kill_libreoffice_processes()
    batches = [docx_files[i:i + batch_size] for i in range(0, len(docx_files), batch_size)]
    logger.info(f"Converting {len(docx_files)} DOCX files in {len(batches)} batches (size: {batch_size})...")
    start_time = time.time()
    for batch_id, batch in enumerate(batches, 1):
        batch_pdfs, batch_failed = convert_batch_with_retry(batch, output_dir, batch_id)
        pdf_files.extend(batch_pdfs)
        total_failed.extend(batch_failed)
    kill_libreoffice_processes()
    total_time = time.time() - start_time
    success_rate = len(pdf_files) / len(docx_files) * 100 if docx_files else 0
    logger.info(f"\n=== CONVERSION SUMMARY ===")
    logger.info(f"Total files: {len(docx_files)}")
    logger.info(f"Successful: {len(pdf_files)} ({success_rate:.1f}%)")
    logger.info(f"Failed: {len(total_failed)} ({100-success_rate:.1f}%)")
    logger.info(f"Time: {total_time:.1f}s | Rate: {len(pdf_files)/total_time:.1f} PDFs/sec" if total_time > 0 else "Time: 0s")
    return pdf_files

def get_raw_file(file: UploadFile):
    try:
        contents = file.file.read()
        df = pd.read_excel(BytesIO(contents), dtype=str)
        return df
    except Exception as e:
        logger.error(f"Error reading file: {e}")
        return pd.DataFrame([])

def fill_transmittal_template(template_path, group_df):
    try:
        temp_doc = Document(template_path)
        if not temp_doc.tables:
            logger.error("No tables found in the transmittal template.")
            return None
        total_groups = len(group_df) // 4 + (1 if len(group_df) % 4 else 0)
        temp_docs = []
        for group_idx in range(total_groups):
            temp_group_doc = Document(template_path)
            current_group_table = temp_group_doc.tables[0]
            page_rows = group_df.iloc[group_idx * 4:(group_idx + 1) * 4]
            for row_idx, row in enumerate(current_group_table.rows):
                if row_idx < len(page_rows):
                    row_data = page_rows.iloc[row_idx]
                    mapping = {f"«{col.upper()}»": str(row_data[col]) for col in group_df.columns if pd.notnull(row_data[col])}
                    qrcode_buffer = None
                    if dl_code := row_data.get('DL_CODE', ''):
                        qrcode_buffer = generate_qrcode(dl_code)
                    for cell in row.cells:
                        if cell.tables:
                            inner_table = cell.tables[0]
                            fill_inner_table(inner_table, mapping, qrcode_buffer)
                else:
                    for cell in row.cells:
                        if cell.tables:
                            inner_table = cell.tables[0]
                            clear_placeholders(inner_table)
            temp_docs.append(temp_group_doc)
        return temp_docs
    except Exception as e:
        logger.error(f"Failed to fill transmittal template: {e}")
        return None

# User Management API Endpoints
@app.get("/api/users")
async def get_users(user_info: dict = Depends(get_current_user)):
    user_email = user_info.get("email", "")
    user_clients, user_access = get_user_clients_and_access(user_email)
    
    if user_access != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT email, clients, access FROM users ORDER BY email")
    users = []
    for row in cursor.fetchall():
        email, clients_str, access = row
        clients = clients_str.split(',') if clients_str else []
        users.append({"email": email, "clients": clients, "access": access})
    conn.close()
    return users

@app.post("/api/users")
async def create_user(user: UserCreate, user_info: dict = Depends(get_current_user)):
    user_email = user_info.get("email", "")
    user_clients, user_access = get_user_clients_and_access(user_email)
    
    if user_access != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    try:
        clients_str = ','.join(user.clients)
        cursor.execute("INSERT INTO users (email, clients, access) VALUES (?, ?, ?)", 
                      (user.email, clients_str, user.access))
        conn.commit()
        return {"success": True, "message": "User created successfully"}
    except sqlite3.IntegrityError:
        raise HTTPException(status_code=400, detail="User with this email already exists")
    finally:
        conn.close()

@app.delete("/api/users/{email}")
async def delete_user(email: str, user_info: dict = Depends(get_current_user)):
    user_email = user_info.get("email", "")
    user_clients, user_access = get_user_clients_and_access(user_email)
    
    if user_access != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM users WHERE email = ?", (email,))
    if cursor.rowcount == 0:
        conn.close()
        raise HTTPException(status_code=404, detail="User not found")
    conn.commit()
    conn.close()
    return {"success": True, "message": "User deleted successfully"}

@app.put("/api/users/{email}")
async def update_user(email: str, user: UserCreate, user_info: dict = Depends(get_current_user)):
    user_email = user_info.get("email", "")
    user_clients, user_access = get_user_clients_and_access(user_email)
    
    if user_access != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    try:
        clients_str = ','.join(user.clients)
        cursor.execute("UPDATE users SET email = ?, clients = ?, access = ? WHERE email = ?", 
                      (user.email, clients_str, user.access, email))
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="User not found")
        conn.commit()
        return {"success": True, "message": "User updated successfully"}
    except sqlite3.IntegrityError:
        raise HTTPException(status_code=400, detail="User with this email already exists")
    finally:
        conn.close()
        
# Audit Trail API Endpoints
@app.get("/api/audit_trail")
async def get_audit_trail(page: int = 1, limit: int = 10, user_info: dict = Depends(get_current_user)):
    offset = (page - 1) * limit
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # Get total count
    cursor.execute('SELECT COUNT(*) as total FROM audit_trail')
    total_count = cursor.fetchone()["total"]
    
    # Get paginated results
    cursor.execute('''
        SELECT id, client, processed_by, processed_at, total_accounts, mode 
        FROM audit_trail 
        ORDER BY processed_at DESC 
        LIMIT ? OFFSET ?
    ''', (limit, offset))
    
    audit_entries = []
    for row in cursor.fetchall():
        audit_entries.append({
            "id": row["id"],
            "client": row["client"],
            "processed_by": row["processed_by"],
            "processed_at": row["processed_at"],
            "total_accounts": row["total_accounts"],
            "mode": row["mode"]
        })
    conn.close()
    
    total_pages = (total_count + limit - 1) // limit
    
    return {
        "entries": audit_entries,
        "pagination": {
            "current_page": page,
            "total_pages": total_pages,
            "total_count": total_count,
            "limit": limit,
            "has_next": page < total_pages,
            "has_prev": page > 1
        }
    }

@app.get("/api/folders")
async def get_folders(user_info: dict = Depends(get_current_user)):
    if not all([FTP_CONFIG["hostname"], FTP_CONFIG["username"], FTP_CONFIG["password"]]):
        raise HTTPException(status_code=500, detail="FTP configuration incomplete")
    
    user_email = user_info.get("email", "")
    user_clients, user_access = get_user_clients_and_access(user_email)
    
    with FTPConnection(FTP_CONFIG["hostname"], FTP_CONFIG["port"], FTP_CONFIG["username"], FTP_CONFIG["password"]) as ftp_conn:
        ftp = ftp_conn.connect()
        if not ftp:
            raise HTTPException(status_code=500, detail="Failed to connect to FTP")
        
        all_folders = get_ftp_folders(ftp)
        
        # If user is admin, return all folders
        if user_access == "admin":
            return all_folders
        
        # If user has assigned clients, return only those folders
        if user_clients:
            available_folders = [folder for folder in all_folders if folder in user_clients]
            return available_folders
        
        # If user has no assigned clients, return empty
        return []

@app.get("/api/all_folders")
async def get_all_folders(user_info: dict = Depends(get_current_user)):
    """Get all available folders for admin users (used in user management modal)"""
    user_email = user_info.get("email", "")
    user_clients, user_access = get_user_clients_and_access(user_email)
    
    if user_access != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    if not all([FTP_CONFIG["hostname"], FTP_CONFIG["username"], FTP_CONFIG["password"]]):
        raise HTTPException(status_code=500, detail="FTP configuration incomplete")
    
    with FTPConnection(FTP_CONFIG["hostname"], FTP_CONFIG["port"], FTP_CONFIG["username"], FTP_CONFIG["password"]) as ftp_conn:
        ftp = ftp_conn.connect()
        if not ftp:
            raise HTTPException(status_code=500, detail="Failed to connect to FTP")
        
        all_folders = get_ftp_folders(ftp)
        return all_folders

@app.post("/api/dl_types")
async def get_dl_types(request: FolderRequest, user_info: dict = Depends(get_current_user)):
    if not request.folder:
        raise HTTPException(status_code=400, detail="Folder not specified")
    
    # Check if user has access to this folder
    user_email = user_info.get("email", "")
    user_clients, user_access = get_user_clients_and_access(user_email)
    
    if user_access != "admin":  # Not admin
        if not user_clients or request.folder not in user_clients:
            raise HTTPException(status_code=403, detail="Access denied to this template folder")
    
    sheet_df = get_sheet_data(SERVICE_ACCOUNT_JSON, SPREADSHEET_ID, SHEET_NAME)
    if sheet_df.empty:
        raise HTTPException(status_code=500, detail="Failed to retrieve Google Sheets data")
    dl_types = sorted(sheet_df[sheet_df["CAMPAIGN"] == request.folder]["DL TYPE"].dropna().unique().tolist())
    return dl_types

@app.post("/api/templates")
async def get_templates(request: FolderRequest, user_info: dict = Depends(get_current_user)):
    if not request.folder:
        raise HTTPException(status_code=400, detail="Folder not specified")
    
    # Check if user has access to this folder
    user_email = user_info.get("email", "")
    user_clients, user_access = get_user_clients_and_access(user_email)
    
    if user_access != "admin":  # Not admin
        if not user_clients or request.folder not in user_clients:
            raise HTTPException(status_code=403, detail="Access denied to this template folder")
    
    with FTPConnection(FTP_CONFIG["hostname"], FTP_CONFIG["port"], FTP_CONFIG["username"], FTP_CONFIG["password"]) as ftp_conn:
        ftp = ftp_conn.connect()
        if not ftp:
            raise HTTPException(status_code=500, detail="Failed to connect to FTP")
        templates = get_ftp_templates(ftp, request.folder)
        return {"message": "Content/Letter Head retrieved successfully", "templates": templates}

@app.post("/api/placeholders")
async def get_placeholders(request: TemplateRequest, user_info: dict = Depends(get_current_user)):
    if not all([request.folder, request.dl_type, request.template]):
        raise HTTPException(status_code=400, detail="Missing parameters")
    
    # Check if user has access to this folder
    user_email = user_info.get("email", "")
    user_clients, user_access = get_user_clients_and_access(user_email)
    
    if user_access != "admin":  # Not admin
        if not user_clients or request.folder not in user_clients:
            raise HTTPException(status_code=403, detail="Access denied to this template folder")
    
    # Store selected folder and dl_type for audit trail
    SESSION_STATE['selected_folder'] = request.folder
    SESSION_STATE['selected_dl_type'] = request.dl_type
    
    with FTPConnection(FTP_CONFIG["hostname"], FTP_CONFIG["port"], FTP_CONFIG["username"], FTP_CONFIG["password"]) as ftp_conn:
        ftp = ftp_conn.connect()
        if not ftp:
            raise HTTPException(status_code=500, detail="Failed to connect to FTP")
        signature_img_path = fetch_signature_from_ftp(ftp)
        if not signature_img_path:
            raise HTTPException(status_code=500, detail="Failed to fetch signature from file server. Folder or file might not be created yet.")
        SESSION_STATE['files_to_cleanup'].append(signature_img_path)
        template_path = download_ftp_template(ftp, request.folder, request.template, is_header_footer=False)
        if not template_path:
            raise HTTPException(status_code=500, detail="Failed to download content template")
        SESSION_STATE['files_to_cleanup'].append(template_path)
        SESSION_STATE['template_path'] = template_path
        sheet_df = get_sheet_data(SERVICE_ACCOUNT_JSON, SPREADSHEET_ID, SHEET_NAME)
        matching_row = sheet_df[(sheet_df["CAMPAIGN"] == request.folder) & (sheet_df["DL TYPE"] == request.dl_type)]
        if matching_row.empty:
            raise HTTPException(status_code=404, detail=f"No header/footer template for {request.folder}/{request.dl_type}")
        header_footer_filename = matching_row["FILE"].iloc[0]
        if not header_footer_filename.lower().endswith('.docx'):
            header_footer_filename += '.docx'
        header_footer_template_path = download_ftp_template(ftp, None, header_footer_filename, is_header_footer=True)
        if not header_footer_template_path:
            raise HTTPException(status_code=500, detail=f"Failed to download header/footer template {header_footer_filename}")
        SESSION_STATE['files_to_cleanup'].append(header_footer_template_path)
        SESSION_STATE['header_footer_template_path'] = header_footer_template_path
        pythoncom.CoInitialize()
        word_app = win32com.client.DispatchEx("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = False
        try:
            base_template = combine_templates(header_footer_template_path, template_path, signature_img_path, word_app)
            if not base_template:
                raise HTTPException(status_code=500, detail="Failed to combine templates")
            SESSION_STATE['base_template'] = base_template
            placeholders = extract_placeholders(base_template)
            SESSION_STATE['placeholders'] = placeholders
            
            # Check if template is already combined
            template_combined_message = ""
            if SESSION_STATE.get('template_combined', False):
                template_combined_message = " (Template already)"
            
            return {
                "message": f"Final template retrieved successfully{template_combined_message}", 
                "placeholders": placeholders,
                "template_combined": SESSION_STATE.get('template_combined', False)
            }
        finally:
            word_app.Quit()
            pythoncom.CoUninitialize()
            os.system("taskkill /IM WINWORD.EXE /F >nul 2>&1")

@app.post("/api/upload_excel")
async def upload_excel(file: UploadFile = File(...)):
    if not file.filename.endswith('.xlsx'):
        raise HTTPException(status_code=400, detail="Invalid file format. Please upload an .xlsx file")
    df = get_raw_file(file)
    if df.empty:
        raise HTTPException(status_code=500, detail="Failed to read Excel file")
    if 'FINAL_AREA' not in df.columns or 'DL_CODE' not in df.columns:
        raise HTTPException(status_code=400, detail="Excel file must contain FINAL_AREA and DL_CODE columns")
    return {"data": df.to_dict(orient='records')}

@app.post("/api/set_mode")
async def set_mode(request: ModeRequest):
    if request.mode not in ["DL Only", "DL w/ Transmittal", "Transmittal Only"]:
        raise HTTPException(status_code=400, detail="Invalid mode")
    
    # Reset session state before setting new mode to prevent conflicts
    if SESSION_STATE.get('selected_mode') and SESSION_STATE.get('selected_mode') != request.mode:
        logger.info(f"Mode changing from {SESSION_STATE.get('selected_mode')} to {request.mode}, resetting state")
        reset_session_state()
    
    SESSION_STATE['selected_mode'] = request.mode
    template_status = {}

    if request.mode in ["DL Only", "DL w/ Transmittal"]:
        if SESSION_STATE.get('base_template') and SESSION_STATE.get('header_footer_template_path') and SESSION_STATE.get('template_path'):
            template_status['dl_template'] = "DL and header/footer templates are ready"
        else:
            template_status['dl_template'] = "DL or header/footer template not loaded"

    if request.mode in ["DL w/ Transmittal", "Transmittal Only"]:
        try:
            with FTPConnection(FTP_CONFIG["hostname"], FTP_CONFIG["port"], FTP_CONFIG["username"], FTP_CONFIG["password"]) as ftp_conn:
                ftp = ftp_conn.connect()
                if not ftp:
                    raise HTTPException(status_code=500, detail="Failed to connect to FTP server")
                transmittal_template = "Transmittal QRCODE.docx"
                transmittal_template_path = download_ftp_template(ftp, None, transmittal_template, is_transmittal=True)
                if not transmittal_template_path:
                    raise HTTPException(status_code=500, detail=f"Failed to download transmittal template {transmittal_template}")
                SESSION_STATE['files_to_cleanup'].append(transmittal_template_path)
                SESSION_STATE['transmittal_template_path'] = transmittal_template_path
                template_status['transmittal_template'] = "Transmittal template is ready"
        except Exception as e:
            logger.error(f"Failed to set mode {request.mode}: {e}")
            template_status['transmittal_template'] = f"Failed to load transmittal template: {str(e)}"
            raise HTTPException(status_code=500, detail=f"Failed to set mode: {str(e)}")

    return {
        "success": True,
        "mode": request.mode,
        "template_status": template_status
    }

@app.post("/api/set_output_format")
async def set_output_format(request: OutputFormatRequest):
    if request.format not in ["zip", "print"]:
        raise HTTPException(status_code=400, detail="Invalid output format")
    
    SESSION_STATE['output_format'] = request.format
    return {
        "success": True,
        "format": request.format
    }

async def generate_pdfs_stream(uploaded_file: UploadFile, dataframe: pd.DataFrame, user_info: dict):
    zipf = None
    try:
        if not SESSION_STATE.get('selected_mode'):
            yield json.dumps({'error': 'No processing mode selected'}) + '\n'
            return

        if not SESSION_STATE.get('base_template') and SESSION_STATE.get('selected_mode') in ["DL Only", "DL w/ Transmittal"]:
            yield json.dumps({'error': 'DL templates not loaded'}) + '\n'
            return
        if not SESSION_STATE.get('transmittal_template_path') and SESSION_STATE.get('selected_mode') in ["DL w/ Transmittal", "Transmittal Only"]:
            yield json.dumps({'error': 'Transmittal template not loaded'}) + '\n'
            return

        logger.debug(f"Excel file loaded with {len(dataframe)} rows")
        today_date = datetime.now().strftime("%B %d, %Y")
        dataframe['DL_ADDRESS'] = dataframe['DL_ADDRESS'].str.upper()
        valid_rows = dataframe[dataframe['LEADS_CHNAME'].notna()]
        total_records = len(valid_rows)
        if total_records == 0:
            yield json.dumps({'error': 'No valid rows found (LEADS_CHNAME missing)'}) + '\n'
            return

        logger.info(f"Processing {total_records} records across {len(valid_rows.groupby('FINAL_AREA'))} FINAL_AREA groups")
        
        # Get user info for audit trail
        user_email = user_info.get("email", "Unknown")
        user_name = user_info.get("name", user_email)
        selected_folder = SESSION_STATE.get('selected_folder', 'Unknown')
        selected_mode = SESSION_STATE.get('selected_mode', 'Unknown')
        
        # Create audit trail entry and get its ID for linking processed accounts
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO audit_trail (client, processed_by, total_accounts, mode, template_folder, dl_type)
            VALUES (?, ?, ?, ?, ?, ?)
        ''', (selected_folder, user_name, total_records, selected_mode, selected_folder, SESSION_STATE.get('selected_dl_type', 'Unknown')))
        audit_id = cursor.lastrowid
        conn.commit()

        # Get selected output format
        output_format = SESSION_STATE.get('output_format', 'zip')  # Default to zip if not set

        with TemporaryDirectory() as temp_dir:
            temp_base_path = Path(temp_dir) / "base_template.docx"
            if SESSION_STATE.get('selected_mode') in ["DL Only", "DL w/ Transmittal"]:
                SESSION_STATE['base_template'].save(temp_base_path)
                SESSION_STATE['files_to_cleanup'].append(temp_base_path)
            temp_transmittal_path = Path(temp_dir) / "transmittal_template.docx"
            if SESSION_STATE.get('selected_mode') in ["DL w/ Transmittal", "Transmittal Only"]:
                Document(SESSION_STATE['transmittal_template_path']).save(temp_transmittal_path)
                SESSION_STATE['files_to_cleanup'].append(temp_transmittal_path)

            processed_records = 0
            start_time = time.time()
            
            # Store processed accounts in batches
            account_batch = []
            batch_size = 100

            # Initialize ZIP file for zip output format
            if output_format == "zip":
                zip_path = Path(temp_dir) / "final_area_pdfs.zip"
                zipf = zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED)

            for final_area, group_df in valid_rows.groupby('FINAL_AREA'):
                logger.debug(f"Processing FINAL_AREA: {final_area} ({len(group_df)} records)")
                area_dir = OUTPUT_DIR / f"{final_area}"
                os.makedirs(area_dir, exist_ok=True)
                docx_files = []
                temp_files = []
                dl_pdf_merger = PdfMerger() if output_format == "zip" and SESSION_STATE.get('selected_mode') in ["DL Only", "DL w/ Transmittal"] else None
                transmittal_pdf_merger = PdfMerger() if output_format == "zip" and SESSION_STATE.get('selected_mode') in ["DL w/ Transmittal", "Transmittal Only"] else None

                try:
                    if SESSION_STATE.get('selected_mode') in ["DL Only", "DL w/ Transmittal"]:
                        for idx, row in group_df.iterrows():
                            # Store account data
                            account_batch.append((
                                audit_id,
                                row.get('DL_CODE', ''),
                                row.get('LEADS_CHNAME', ''),
                                row.get('DL_ADDRESS', ''),
                                row.get('FINAL_AREA', '')
                            ))
                            
                            # Process batch if it reaches the batch size
                            if len(account_batch) >= batch_size:
                                cursor.executemany(
                                    "INSERT INTO processed_accounts (audit_id, dl_code, leads_chname, dl_address, final_area) VALUES (?, ?, ?, ?, ?)",
                                    account_batch
                                )
                                conn.commit()
                                account_batch = []
                            
                            barcode_buffer = None
                            if barcode_value := row.get('DL_CODE', ''):
                                barcode_buffer = generate_barcode(barcode_value)
                            amount_words = amount_to_words(row.get('amount', '0.00'))
                            mapping = {f"«{col.upper()}»": row[col] for col in dataframe.columns if pd.notnull(row[col])}
                            mapping.update({
                                "«IMAGE_BARCODE»": barcode_buffer or "",
                                "«DL_DATE»": today_date,
                                "«AMOUNT_ABBR»": amount_words,
                                "«IMAGE_SIGNATURE»": SESSION_STATE.get('files_to_cleanup', [])[0] or ""
                            })
                            filled_doc = Document(temp_base_path)
                            filled_doc = fill_template(filled_doc, mapping, barcode_buffer)
                            if filled_doc:
                                unique_name = f"dl_{final_area}_{idx}_{uuid.uuid4().hex[:6]}"
                                docx_output = area_dir / f"{unique_name}.docx"
                                filled_doc.save(docx_output)
                                docx_files.append(str(docx_output))
                                temp_files.append(str(docx_output))
                            processed_records += 1
                            progress = (processed_records / total_records) * 100
                            yield json.dumps({
                                'progress': progress,
                                'message': f"Processing record {processed_records}/{total_records} (FINAL_AREA: {final_area})"
                            }) + '\n'
                            await asyncio.sleep(0)

                    if SESSION_STATE.get('selected_mode') in ["DL w/ Transmittal", "Transmittal Only"]:
                        # If we're only doing transmittal, we need to store the account data here
                        if SESSION_STATE.get('selected_mode') == "Transmittal Only":
                            for _, row in group_df.iterrows():
                                account_batch.append((
                                    audit_id,
                                    row.get('DL_CODE', ''),
                                    row.get('LEADS_CHNAME', ''),
                                    row.get('DL_ADDRESS', ''),
                                    row.get('FINAL_AREA', '')
                                ))
                                
                                # Process batch if it reaches the batch size
                                if len(account_batch) >= batch_size:
                                    cursor.executemany(
                                        "INSERT INTO processed_accounts (audit_id, dl_code, leads_chname, dl_address, final_area) VALUES (?, ?, ?, ?, ?)",
                                        account_batch
                                    )
                                    conn.commit()
                                    account_batch = []
                    
                        transmittal_docs = fill_transmittal_template(temp_transmittal_path, group_df)
                        if transmittal_docs:
                            for doc_idx, transmittal_doc in enumerate(transmittal_docs):
                                unique_name = f"transmittal_{final_area}_{doc_idx}_{uuid.uuid4().hex[:6]}"
                                transmittal_docx_output = area_dir / f"{unique_name}.docx"
                                transmittal_doc.save(transmittal_docx_output)
                                docx_files.append(str(transmittal_docx_output))
                                temp_files.append(str(transmittal_docx_output))

                    if docx_files:
                        if output_format == "zip":
                            logger.debug(f"Converting {len(docx_files)} DOCX files to PDF for {final_area}")
                            yield json.dumps({
                                'progress': (processed_records / total_records) * 100,
                                'message': f"Converting {len(docx_files)} documents to PDF for {final_area}..."
                            }) + '\n'
                            await asyncio.sleep(0)
                            pdf_files = batch_convert_libreoffice(docx_files, area_dir, batch_size=300)
                            if pdf_files:
                                for pdf_file in pdf_files:
                                    if "transmittal" in Path(pdf_file).name and transmittal_pdf_merger:
                                        transmittal_pdf_merger.append(pdf_file)
                                    elif "dl" in Path(pdf_file).name and dl_pdf_merger:
                                        dl_pdf_merger.append(pdf_file)
                                if dl_pdf_merger and SESSION_STATE.get('selected_mode') in ["DL Only", "DL w/ Transmittal"]:
                                    dl_merged_path = area_dir / f"{final_area}_DL.pdf"
                                    with open(dl_merged_path, 'wb') as output_file:
                                        dl_pdf_merger.write(output_file)
                                    zipf.write(dl_merged_path, f"{final_area}_DL.pdf")
                                    logger.info(f"Created DL PDF for {final_area}: {dl_merged_path}")
                                if transmittal_pdf_merger and SESSION_STATE.get('selected_mode') in ["DL w/ Transmittal", "Transmittal Only"]:
                                    transmittal_merged_path = area_dir / f"{final_area}_Transmittal.pdf"
                                    with open(transmittal_merged_path, 'wb') as output_file:
                                        transmittal_pdf_merger.write(output_file)
                                    zipf.write(transmittal_merged_path, f"{final_area}_Transmittal.pdf")
                                    logger.info(f"Created Transmittal PDF for {final_area}: {transmittal_merged_path}")
                            else:
                                logger.warning(f"No PDFs generated for {final_area}")
                        else:
                            logger.info(f"Keeping {len(docx_files)} DOCX files for direct printing in {final_area}")
                            yield json.dumps({
                                'progress': (processed_records / total_records) * 100,
                                'message': f"Prepared {len(docx_files)} documents for printing in {final_area}..."
                            }) + '\n'
                            await asyncio.sleep(0)
                    cleanup_files(temp_files)
                except Exception as e:
                    logger.error(f"Error during document processing for {final_area}: {e}")
                    yield json.dumps({'error': f'Document processing failed for {final_area}: {str(e)}'}) + '\n'
                finally:
                    if dl_pdf_merger:
                        dl_pdf_merger.close()
                    if transmittal_pdf_merger:
                        transmittal_pdf_merger.close()

            # Process any remaining accounts in the batch
            if account_batch:
                cursor.executemany(
                    "INSERT INTO processed_accounts (audit_id, dl_code, leads_chname, dl_address, final_area) VALUES (?, ?, ?, ?, ?)",
                    account_batch
                )
                conn.commit()

            conn.close()

            # Finalize ZIP file if applicable
            if output_format == "zip":
                zipf.close()
                zipf = None
                final_zip_path = OUTPUT_DIR / "final_area_pdfs.zip"
                shutil.move(str(zip_path), str(final_zip_path))
                SESSION_STATE['zip_path'] = str(final_zip_path)
                SESSION_STATE['files_to_cleanup'].append(str(final_zip_path))

            total_time = time.time() - start_time
            yield json.dumps({
                'progress': 100,
                'message': f"Processing complete! Generated {'PDFs' if output_format == 'zip' else 'DOCX files'} for {len(valid_rows.groupby('FINAL_AREA'))} FINAL_AREA groups in {total_time:.1f}s",
                'download_ready': output_format == "zip",
                'print_ready': output_format == "print",
                'areas': list(valid_rows.groupby('FINAL_AREA').groups.keys())
            }) + '\n'

    except Exception as e:
        logger.error(f"Error in generate_pdfs_stream: {e}")
        yield json.dumps({'error': f'Processing failed: {str(e)}'}) + '\n'
    finally:
        if zipf:
            zipf.close()

def print_docx_to_specific_printer(docx_files, printer_name=None):
    """Print DOCX files to a specific printer using Word"""
    try:
        pythoncom.CoInitialize()
        word_app = win32com.client.DispatchEx("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = False
        
        try:
            for docx_file in docx_files:
                doc = word_app.Documents.Open(str(Path(docx_file).absolute()))
                try:
                    if printer_name:
                        # Set specific printer
                        word_app.ActivePrinter = printer_name
                    # Print the document
                    doc.PrintOut()
                    logger.info(f"Printed {docx_file} to {printer_name or 'default printer'}")
                finally:
                    doc.Close(SaveChanges=False)
        finally:
            word_app.Quit()
            pythoncom.CoUninitialize()
        return True
    except Exception as e:
        logger.error(f"Failed to print DOCX files to {printer_name}: {e}")
        return False

@app.get("/api/print_files/{area}")
async def print_files(area: str, printer: str = None):
    area_dir = OUTPUT_DIR / f"{area}"
    if not area_dir.exists():
        raise HTTPException(status_code=404, detail=f"Files for area {area} not found")
    
    # Check for PDF files first (ZIP format), then DOCX files (print format)
    pdf_files = list(area_dir.glob("*.pdf"))
    docx_files = list(area_dir.glob("*.docx"))
    
    if pdf_files:
        # Use existing PDF printing logic
        try:
            if printer:
                success = print_to_specific_printer(pdf_files, printer)
                if not success:
                    raise HTTPException(status_code=500, detail=f"Failed to print to {printer}")
            else:
                for pdf_file in pdf_files:
                    subprocess.run(["SumatraPDF", "-print-to-default", str(pdf_file)], 
                                  check=True, creationflags=subprocess.CREATE_NO_WINDOW)
            
            return {"success": True, "message": f"Printed {len(pdf_files)} PDF files for area {area}" + (f" to {printer}" if printer else " to default printer")}
        except Exception as e:
            logger.error(f"Failed to print PDF files for area {area}: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to print PDF files: {str(e)}")
    
    elif docx_files:
        # Use new DOCX printing logic
        try:
            success = print_docx_to_specific_printer(docx_files, printer)
            if not success:
                raise HTTPException(status_code=500, detail=f"Failed to print DOCX files to {printer or 'default printer'}")
            
            return {"success": True, "message": f"Printed {len(docx_files)} DOCX files for area {area}" + (f" to {printer}" if printer else " to default printer")}
        except Exception as e:
            logger.error(f"Failed to print DOCX files for area {area}: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to print DOCX files: {str(e)}")
    
    else:
        raise HTTPException(status_code=404, detail=f"No files found for area {area}")

@app.post("/api/cleanup")
async def cleanup():
    try:
        # Enhanced cleanup that resets everything
        reset_session_state()
        return {"success": True, "message": "Files cleaned up and session reset successfully"}
    except Exception as e:
        logger.error(f"Cleanup failed: {e}")
        return {"success": False, "detail": str(e)}

@app.post("/api/generate_pdfs")
async def generate_pdfs(file: UploadFile = File(...), user_info: dict = Depends(get_current_user)):
    if not file.filename.endswith('.xlsx'):
        raise HTTPException(status_code=400, detail="Invalid file format. Please upload an .xlsx file")
    df = get_raw_file(file)
    if df.empty:
        raise HTTPException(status_code=500, detail="Failed to read Excel file")
    return StreamingResponse(
        generate_pdfs_stream(file, df, user_info),
        media_type="application/json"
    )

@app.get("/api/download_zip")
async def download_zip():
    zip_path = SESSION_STATE.get('zip_path')
    if not zip_path or not Path(zip_path).exists():
        raise HTTPException(status_code=404, detail="ZIP file not found")
    return FileResponse(zip_path, filename="final_area_pdfs.zip", media_type="application/zip")

def get_available_printers():
    """Get list of available printers on the system"""
    try:
        import win32print
        printers = []
        printer_enum = win32print.EnumPrinters(win32print.PRINTER_ENUM_LOCAL | win32print.PRINTER_ENUM_CONNECTIONS)
        for printer in printer_enum:
            printers.append({
                "name": printer[2],  # Printer name
                "is_default": printer[2] == win32print.GetDefaultPrinter()
            })
        return printers
    except Exception as e:
        logger.error(f"Failed to get available printers: {e}")
        return []

def print_to_specific_printer(pdf_files, printer_name):
    """Print PDF files to a specific printer"""
    try:
        for pdf_file in pdf_files:
            # Use SumatraPDF with specific printer
            subprocess.run([
                "SumatraPDF", 
                "-print-to", printer_name, 
                "-silent",
                str(pdf_file)
            ], check=True, creationflags=subprocess.CREATE_NO_WINDOW)
        return True
    except Exception as e:
        logger.error(f"Failed to print to {printer_name}: {e}")
        return False

@app.get("/api/printers")
async def get_printers(user_info: dict = Depends(get_current_user)):
    """Get list of available printers"""
    printers = get_available_printers()
    return {"printers": printers}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)
